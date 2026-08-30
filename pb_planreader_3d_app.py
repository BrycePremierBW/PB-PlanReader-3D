from __future__ import annotations

import base64
import csv
import gc
import hashlib
import io
import json
import math
import mimetypes
import os
import queue
import re
import shlex
import shutil
import sqlite3
import subprocess
import sys
import tempfile
import textwrap
import threading
import time
import zipfile
from contextlib import contextmanager
from dataclasses import dataclass
from datetime import datetime
from functools import lru_cache
from pathlib import Path
from typing import Any, Callable, Dict, Iterable, List, Optional, Sequence, Tuple

import numpy as np
import pandas as pd
import plotly.graph_objects as go
import requests
import streamlit as st
from PIL import Image, ImageDraw, ImageFont

try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None

try:
    import cv2  # opencv-python-headless for building-envelope detection
except Exception:
    cv2 = None

try:
    from docx import Document as DocxDocument
except Exception:
    DocxDocument = None

def _patch_image_to_url_compat() -> bool:
    """Re-add ``streamlit.elements.image.image_to_url`` for the canvas package.

    ``streamlit-drawable-canvas`` 0.9.3 calls the legacy
    ``st_image.image_to_url(image, width, clamp, channels, output_format,
    image_id)`` signature. Modern Streamlit moved ``image_to_url`` into
    ``streamlit.elements.lib.image_utils`` and replaced the ``width`` argument
    with a ``LayoutConfig``, so the attribute no longer exists on the public
    module. This shim re-adds the legacy signature and forwards it to the
    current implementation.
    """
    try:
        import streamlit.elements.image as st_image_module
        from streamlit.elements.lib.image_utils import image_to_url as _modern_image_to_url
        from streamlit.elements.lib.layout_utils import LayoutConfig

        def _image_to_url(image, width, clamp, channels, output_format, image_id):
            return _modern_image_to_url(
                image, LayoutConfig(width=width), clamp, channels, output_format, image_id
            )

        st_image_module.image_to_url = _image_to_url
        return True
    except Exception:
        return False


try:
    _patch_image_to_url_compat()
    from streamlit_drawable_canvas import st_canvas
    CANVAS_AVAILABLE = True
except Exception:
    st_canvas = None
    CANVAS_AVAILABLE = False

try:
    from planreader_line_mapper import plan_line_editor
except Exception:
    plan_line_editor = None

try:
    from openai import OpenAI
except Exception:
    OpenAI = None

# Offline plan reader (no AI needed)
try:
    from pb_planreader_offline import (
        analyze_page_offline,
        generate_takeoff_offline,
        extract_text_offline,
        detect_walls,
        detect_dimensions,
        detect_scale,
        detect_rooms,
        detect_materials,
        detect_colours,
        classify_page_offline,
        generate_report,
        wall_length_real_m,
        real_metres_per_page_mm,
        PDF_PT_TO_MM,
    )
    OFFLINE_READER_AVAILABLE = True
except Exception:
    OFFLINE_READER_AVAILABLE = False

try:
    import psycopg2
    import psycopg2.extras
except Exception:
    psycopg2 = None

APP_NAME = "Premier Brushworks Plan Reader & 3D Take-off"
APP_VERSION = "1.0.0"
PB_DARK = "#171717"
PB_GOLD = "#D7A21B"
PB_CREAM = "#F7F2E8"
PB_GREY = "#E6E1D8"
PB_BLUE = "#276FBF"
PB_GREEN = "#2E8B57"
PB_RED = "#B33A3A"

DATA_DIR = Path(os.environ.get("PLANREADER_DATA_DIR", Path.cwd() / "planreader_data")).resolve()
DB_PATH = DATA_DIR / "planreader.db"
WORKSPACE_DIR = DATA_DIR / "workspaces"
DATA_DIR.mkdir(parents=True, exist_ok=True)
WORKSPACE_DIR.mkdir(parents=True, exist_ok=True)

JOBHUB_DATABASE_URL = os.environ.get("JOBHUB_DATABASE_URL") or os.environ.get("DATABASE_URL") or ""
JOBHUB_DB_PATH = os.environ.get("JOBHUB_DB_PATH", "")
DEFAULT_OPENAI_MODEL = os.environ.get("OPENAI_MODEL", "gpt-5.6")
# gemini-2.5-flash / 2.5-flash-lite / 2.5-pro were retired for new keys; a
# GEMINI_MODEL env var naming a retired model is ignored rather than reused.
_RETIRED_GEMINI_MODELS = {"gemini-2.5-flash", "gemini-2.5-flash-lite", "gemini-2.5-pro"}
_candidate_gemini = os.environ.get("GEMINI_MODEL", "").strip()
if _candidate_gemini and _candidate_gemini not in _RETIRED_GEMINI_MODELS:
    DEFAULT_GEMINI_MODEL = _candidate_gemini
else:
    DEFAULT_GEMINI_MODEL = "gemini-3.5-flash"
DEFAULT_AI_PROVIDER = os.environ.get("AI_PROVIDER", "OpenAI")
AI_PROVIDERS = ["OpenAI", "Google Gemini"]

TAKEOFF_COLUMNS = [
    "section",
    "element",
    "location",
    "substrate",
    "finish_system",
    "quantity",
    "unit",
    "quantity_status",
    "source_page",
    "source_reference",
    "inclusion_status",
    "coats",
    "coverage_m2_per_litre",
    "productivity_m2_per_hour",
    "rate_per_unit",
    "confidence",
    "notes",
]

REGISTER_NAMES = [
    "source_basis",
    "drawing_register",
    "inclusions",
    "exclusions",
    "clarifications",
    "assumptions",
    "rfis",
    "door_schedule",
    "colour_finish_schedule",
    "access_constraints",
    "risks",
]

PAGE_TYPES = [
    "Title / Drawing Register",
    "Floor Plan",
    "Reflected Ceiling Plan",
    "Roof Plan",
    "Elevation",
    "Section",
    "Render / Artist's Impression",
    "Door / Window Schedule",
    "Finishes Schedule",
    "Specification",
    "Structural",
    "Services",
    "Landscape / Civil",
    "Other",
]

SUBSTRATES = [
    "Plasterboard",
    "Wet-area plasterboard",
    "Fibre cement",
    "Precast concrete",
    "Masonry / blockwork",
    "Render",
    "Acrylic render",
    "Cement render",
    "Timber door",
    "Timber trim / joinery",
    "Structural steel",
    "Metalwork",
    "Concrete floor",
    "Soffit",
    "Previously painted substrate",
    "Other",
]

FINISH_SYSTEMS = [
    "Interior coatings",
    "Ceiling flat",
    "Low sheen wall system",
    "Semi-gloss / enamel",
    "Exterior acrylic",
    "Elastomeric / membrane",
    "Concrete coating",
    "Specialist floor coating",
    "Metal primer + topcoats",
    "Clear / stain system",
    "To be confirmed",
]

UNIT_OPTIONS = ["m²", "lm", "No.", "item", "L", "allowance"]
STATUS_OPTIONS = ["Measured", "Provisional measured", "To measure", "Allowance", "Excluded", "Not applicable"]
INCLUSION_OPTIONS = ["INCLUSION", "SEPARATE ITEM", "PROVISIONAL", "EXCLUSION", "CLARIFICATION"]


def now_stamp() -> str:
    return datetime.now().isoformat(timespec="seconds")


def safe_name(value: Any, fallback: str = "item") -> str:
    text = re.sub(r"[^A-Za-z0-9._-]+", "_", str(value or "").strip()).strip("._-")
    return text or fallback


def sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def sha256_text(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def verify_password(password: str, stored_hash: Any) -> bool:
    """Verify a JobHub password hash (PBKDF2-SHA256, with legacy SHA-256 fallback)."""
    import hmac as _hmac
    stored_hash = str(stored_hash or "")
    if not stored_hash or not password:
        return False
    if stored_hash.startswith("pbkdf2_sha256$"):
        try:
            _, iterations_text, salt_hex, expected = stored_hash.split("$", 3)
            actual = hashlib.pbkdf2_hmac(
                "sha256",
                str(password).encode("utf-8"),
                bytes.fromhex(salt_hex),
                int(iterations_text),
            ).hex()
            return _hmac.compare_digest(actual, expected)
        except (ValueError, TypeError):
            return False
    legacy = hashlib.sha256(str(password).encode("utf-8")).hexdigest()
    return _hmac.compare_digest(legacy, stored_hash)


def to_float(value: Any, default: float = 0.0) -> float:
    try:
        if value is None or value == "":
            return default
        result = float(value)
        if result != result or result in (float("inf"), float("-inf")):
            return default
        return result
    except Exception:
        return default


def to_int(value: Any, default: int = 0) -> int:
    try:
        if value is None or value == "":
            return default
        return int(float(value))
    except Exception:
        return default


def normalise_whitespace(text: Any) -> str:
    return re.sub(r"\s+", " ", str(text or "")).strip()


def app_css() -> None:
    st.markdown(
        f"""
        <style>
        .stApp {{ background: {PB_CREAM}; }}
        [data-testid="stSidebar"] {{ background: {PB_DARK}; }}
        [data-testid="stSidebar"] * {{ color: #fff; }}
        .block-container {{ padding-top: 1.1rem; padding-bottom: 3rem; max-width: 1650px; }}
        .pb-hero {{ background: linear-gradient(135deg, #171717 0%, #303030 100%); color: white;
                   border-left: 7px solid {PB_GOLD}; border-radius: 14px; padding: 1.2rem 1.35rem; margin-bottom: 1rem; }}
        .pb-hero h1 {{ margin: 0; font-size: 1.75rem; }}
        .pb-hero p {{ margin: .35rem 0 0 0; color: #ddd; }}
        .pb-card {{ background: white; border: 1px solid #ddd4c7; border-radius: 12px; padding: 1rem; margin-bottom: .9rem;
                   box-shadow: 0 2px 10px rgba(0,0,0,.035); }}
        .pb-note {{ background: #fff8df; border-left: 5px solid {PB_GOLD}; border-radius: 8px; padding: .8rem 1rem; }}
        .pb-warning {{ background: #fff0f0; border-left: 5px solid {PB_RED}; border-radius: 8px; padding: .8rem 1rem; }}
        .pb-good {{ background: #eef9f1; border-left: 5px solid {PB_GREEN}; border-radius: 8px; padding: .8rem 1rem; }}
        .pb-badge {{ display:inline-block; background:{PB_GOLD}; color:#171717; font-weight:700; border-radius:999px;
                    padding:.22rem .55rem; margin-right:.35rem; font-size:.78rem; }}
        div[data-testid="stMetric"] {{ background:white; border:1px solid #ded8cd; padding:.7rem; border-radius:10px; }}
        .small-muted {{ color:#6e675e; font-size:.85rem; }}
        </style>
        """,
        unsafe_allow_html=True,
    )


# -----------------------------------------------------------------------------
# Local PlanReader database
# -----------------------------------------------------------------------------


def local_connect() -> sqlite3.Connection:
    conn = sqlite3.connect(DB_PATH, check_same_thread=False)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA foreign_keys=ON")
    return conn


def init_local_db() -> None:
    conn = local_connect()
    cur = conn.cursor()
    cur.executescript(
        """
        CREATE TABLE IF NOT EXISTS workspaces (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            jobhub_job_id INTEGER,
            job_no TEXT,
            job_name TEXT,
            builder_client TEXT,
            site_address TEXT,
            drawing_issue TEXT,
            estimator TEXT,
            status TEXT DEFAULT 'Draft',
            executive_summary TEXT,
            created_at TEXT,
            updated_at TEXT
        );
        CREATE UNIQUE INDEX IF NOT EXISTS idx_workspace_jobhub ON workspaces(jobhub_job_id) WHERE jobhub_job_id IS NOT NULL;

        CREATE TABLE IF NOT EXISTS workspace_settings (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            workspace_id INTEGER NOT NULL,
            key TEXT NOT NULL,
            value TEXT,
            updated_at TEXT,
            UNIQUE(workspace_id, key)
        );

        CREATE TABLE IF NOT EXISTS documents (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            workspace_id INTEGER NOT NULL,
            source_type TEXT,
            jobhub_table TEXT,
            jobhub_record_id TEXT,
            file_name TEXT,
            mime_type TEXT,
            path TEXT,
            sha256 TEXT,
            category TEXT,
            page_count INTEGER DEFAULT 0,
            extracted_text TEXT,
            uploaded_at TEXT,
            FOREIGN KEY(workspace_id) REFERENCES workspaces(id) ON DELETE CASCADE
        );
        CREATE UNIQUE INDEX IF NOT EXISTS idx_doc_hash_workspace ON documents(workspace_id, sha256);

        CREATE TABLE IF NOT EXISTS pages (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            document_id INTEGER NOT NULL,
            workspace_id INTEGER NOT NULL,
            page_no INTEGER,
            page_label TEXT,
            page_type TEXT,
            scale_text TEXT,
            px_per_m REAL,
            image_path TEXT,
            width_px INTEGER,
            height_px INTEGER,
            render_zoom REAL,
            extracted_text TEXT,
            selected INTEGER DEFAULT 1,
            created_at TEXT,
            FOREIGN KEY(document_id) REFERENCES documents(id) ON DELETE CASCADE,
            FOREIGN KEY(workspace_id) REFERENCES workspaces(id) ON DELETE CASCADE
        );

        CREATE TABLE IF NOT EXISTS takeoff_rows (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            workspace_id INTEGER NOT NULL,
            section TEXT,
            element TEXT,
            location TEXT,
            substrate TEXT,
            finish_system TEXT,
            quantity REAL DEFAULT 0,
            unit TEXT,
            quantity_status TEXT,
            source_page TEXT,
            source_reference TEXT,
            inclusion_status TEXT,
            coats REAL DEFAULT 2,
            coverage_m2_per_litre REAL DEFAULT 12,
            productivity_m2_per_hour REAL DEFAULT 8,
            rate_per_unit REAL DEFAULT 0,
            confidence TEXT,
            notes TEXT,
            row_role TEXT DEFAULT '',
            created_at TEXT,
            updated_at TEXT,
            FOREIGN KEY(workspace_id) REFERENCES workspaces(id) ON DELETE CASCADE
        );

        CREATE TABLE IF NOT EXISTS register_items (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            workspace_id INTEGER NOT NULL,
            register_name TEXT NOT NULL,
            item_no TEXT,
            title TEXT,
            detail TEXT,
            priority TEXT,
            source_reference TEXT,
            status TEXT,
            created_at TEXT,
            FOREIGN KEY(workspace_id) REFERENCES workspaces(id) ON DELETE CASCADE
        );

        CREATE TABLE IF NOT EXISTS mapped_zones (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            workspace_id INTEGER NOT NULL,
            page_id INTEGER NOT NULL,
            name TEXT,
            view_type TEXT,
            polygon_json TEXT,
            x_px REAL,
            y_px REAL,
            w_px REAL,
            h_px REAL,
            px_per_m REAL,
            wall_height_m REAL DEFAULT 2.7,
            area_m2 REAL DEFAULT 0,
            substrate TEXT,
            finish_system TEXT,
            quantity_status TEXT,
            source_reference TEXT,
            created_at TEXT,
            FOREIGN KEY(workspace_id) REFERENCES workspaces(id) ON DELETE CASCADE,
            FOREIGN KEY(page_id) REFERENCES pages(id) ON DELETE CASCADE
        );

        CREATE TABLE IF NOT EXISTS model_masses (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            workspace_id INTEGER NOT NULL,
            label TEXT,
            level_name TEXT,
            x REAL DEFAULT 0,
            y REAL DEFAULT 0,
            z REAL DEFAULT 0,
            width REAL DEFAULT 1,
            depth REAL DEFAULT 1,
            height REAL DEFAULT 2.7,
            finish TEXT,
            source_reference TEXT,
            confidence TEXT,
            notes TEXT,
            created_at TEXT,
            FOREIGN KEY(workspace_id) REFERENCES workspaces(id) ON DELETE CASCADE
        );

        CREATE TABLE IF NOT EXISTS measurement_lines (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            workspace_id INTEGER NOT NULL,
            page_id INTEGER NOT NULL,
            takeoff_row_id INTEGER,
            label TEXT,
            unit TEXT,
            colour TEXT,
            kind TEXT DEFAULT 'line',
            x1 REAL DEFAULT 0,
            y1 REAL DEFAULT 0,
            x2 REAL DEFAULT 0,
            y2 REAL DEFAULT 0,
            points TEXT,
            length_m REAL DEFAULT 0,
            area_m2 REAL DEFAULT 0,
            perimeter_m REAL DEFAULT 0,
            quantity_status TEXT,
            moved INTEGER DEFAULT 0,
            notes TEXT,
            created_at TEXT,
            FOREIGN KEY(workspace_id) REFERENCES workspaces(id) ON DELETE CASCADE,
            FOREIGN KEY(page_id) REFERENCES pages(id) ON DELETE CASCADE
        );

        CREATE TABLE IF NOT EXISTS model_openings (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            workspace_id INTEGER NOT NULL,
            mass_id INTEGER,
            label TEXT,
            opening_type TEXT,
            face TEXT,
            offset_x REAL DEFAULT 0,
            offset_z REAL DEFAULT 0,
            width REAL DEFAULT 0.9,
            height REAL DEFAULT 2.1,
            count INTEGER DEFAULT 1,
            notes TEXT,
            source_reference TEXT,
            created_at TEXT,
            FOREIGN KEY(workspace_id) REFERENCES workspaces(id) ON DELETE CASCADE,
            FOREIGN KEY(mass_id) REFERENCES model_masses(id) ON DELETE SET NULL
        );

        CREATE TABLE IF NOT EXISTS ai_runs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            workspace_id INTEGER NOT NULL,
            run_type TEXT,
            model TEXT,
            source_pages TEXT,
            status TEXT,
            response_json TEXT,
            error_message TEXT,
            created_at TEXT,
            FOREIGN KEY(workspace_id) REFERENCES workspaces(id) ON DELETE CASCADE
        );
        """
    )
    _ensure_measurement_columns(conn)
    _ensure_takeoff_columns(conn)
    _ensure_pages_columns(conn)
    conn.commit()
    conn.close()


def _ensure_measurement_columns(conn: sqlite3.Connection) -> None:
    """Migrate existing ``measurement_lines`` tables to the draw-first schema.

    Older databases were created with only single-segment line columns
    (x1/y1/x2/y2/length_m). The draw-first mapper adds polygon outlines, so
    add the missing columns idempotently.
    """
    existing = {row[1] for row in conn.execute("PRAGMA table_info(measurement_lines)").fetchall()}
    wanted = {
        "kind": "TEXT DEFAULT 'line'",
        "points": "TEXT",
        "area_m2": "REAL DEFAULT 0",
        "perimeter_m": "REAL DEFAULT 0",
    }
    for name, ddl in wanted.items():
        if name not in existing:
            conn.execute(f"ALTER TABLE measurement_lines ADD COLUMN {name} {ddl}")


def _ensure_takeoff_columns(conn: sqlite3.Connection) -> None:
    """Migrate existing ``takeoff_rows`` tables to carry the row role column.

    ``row_role`` marks measurement rows such as per-level internal floor areas
    (``floor_area``) that drive floor-m² pricing but are not themselves priced.
    """
    existing = {row[1] for row in conn.execute("PRAGMA table_info(takeoff_rows)").fetchall()}
    if "row_role" not in existing:
        conn.execute("ALTER TABLE takeoff_rows ADD COLUMN row_role TEXT DEFAULT ''")


def _ensure_pages_columns(conn: sqlite3.Connection) -> None:
    """Migrate existing ``pages`` tables to carry the PDF render zoom.

    ``render_zoom`` records the zoom a page image was rasterised at, so the
    detected drawing scale can be converted to pixels-per-metre exactly
    (large sheets are capped below the historical 1.7x to bound memory).
    Legacy rows predating the column were all rendered at 1.7x and fall back
    to that in ``auto_detect_scale``.
    """
    existing = {row[1] for row in conn.execute("PRAGMA table_info(pages)").fetchall()}
    if "render_zoom" not in existing:
        conn.execute("ALTER TABLE pages ADD COLUMN render_zoom REAL")


def lquery(sql: str, params: Sequence[Any] = ()) -> List[Dict[str, Any]]:
    conn = local_connect()
    try:
        rows = conn.execute(sql, tuple(params)).fetchall()
        return [dict(r) for r in rows]
    finally:
        conn.close()


def ldf(sql: str, params: Sequence[Any] = ()) -> pd.DataFrame:
    conn = local_connect()
    try:
        return pd.read_sql_query(sql, conn, params=tuple(params))
    finally:
        conn.close()


def lexecute(sql: str, params: Sequence[Any] = ()) -> int:
    conn = local_connect()
    try:
        cur = conn.execute(sql, tuple(params))
        conn.commit()
        return int(cur.lastrowid or 0)
    finally:
        conn.close()


def lexecutemany(sql: str, rows: Iterable[Sequence[Any]]) -> None:
    conn = local_connect()
    try:
        conn.executemany(sql, list(rows))
        conn.commit()
    finally:
        conn.close()


def workspace_path(workspace_id: int) -> Path:
    path = WORKSPACE_DIR / str(workspace_id)
    (path / "documents").mkdir(parents=True, exist_ok=True)
    (path / "pages").mkdir(parents=True, exist_ok=True)
    (path / "exports").mkdir(parents=True, exist_ok=True)
    return path


# -----------------------------------------------------------------------------
# JobHub bridge
# -----------------------------------------------------------------------------

# Column names that can hold a document's file bytes. Discovery and per-file
# fetch must agree on this list so a blob stored under any common name is read.
_DOCUMENT_BLOB_COLUMNS = (
    "blob_data", "file_data", "file_content", "filedata", "pdf_data", "pdf_file",
    "document_data", "doc_data", "attachment", "attachments", "file", "content",
    "data", "blob", "bytes", "binary_data", "raw_data", "object_data",
)


@dataclass
class JobHubBridge:
    kind: str
    source: str

    @contextmanager
    def connect(self):
        if self.kind == "postgres":
            if psycopg2 is None:
                raise RuntimeError("psycopg2-binary is not installed.")
            conn = psycopg2.connect(self.source)
            try:
                yield conn
            finally:
                conn.close()
        else:
            conn = sqlite3.connect(self.source)
            conn.row_factory = sqlite3.Row
            try:
                yield conn
            finally:
                conn.close()

    def table_names(self) -> List[str]:
        with self.connect() as conn:
            cur = conn.cursor()
            if self.kind == "postgres":
                cur.execute("SELECT table_name FROM information_schema.tables WHERE table_schema='public'")
                return [str(r[0]) for r in cur.fetchall()]
            cur.execute("SELECT name FROM sqlite_master WHERE type='table'")
            return [str(r[0]) for r in cur.fetchall()]

    def columns(self, table: str) -> List[str]:
        safe = re.sub(r"[^A-Za-z0-9_]", "", table)
        with self.connect() as conn:
            cur = conn.cursor()
            if self.kind == "postgres":
                cur.execute(
                    "SELECT column_name FROM information_schema.columns WHERE table_schema='public' AND table_name=%s ORDER BY ordinal_position",
                    (safe,),
                )
                return [str(r[0]) for r in cur.fetchall()]
            cur.execute(f"PRAGMA table_info({safe})")
            return [str(r[1]) for r in cur.fetchall()]

    def query(self, sql: str, params: Sequence[Any] = ()) -> List[Dict[str, Any]]:
        with self.connect() as conn:
            if self.kind == "postgres":
                cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
                cur.execute(sql.replace("?", "%s"), tuple(params))
                return [dict(r) for r in cur.fetchall()]
            cur = conn.execute(sql, tuple(params))
            return [dict(r) for r in cur.fetchall()]

    def execute(self, sql: str, params: Sequence[Any] = (), returning: bool = False) -> Any:
        with self.connect() as conn:
            cur = conn.cursor()
            cur.execute(sql.replace("?", "%s") if self.kind == "postgres" else sql, tuple(params))
            result = None
            if returning:
                row = cur.fetchone()
                result = row[0] if row else None
            conn.commit()
            return result

    def discover_documents_for_job(self, job_id: int) -> List[Dict[str, Any]]:
        """Scan the common JobHub document/attachment tables in a single connection.

        Opening a separate connection per table used to spike connection counts
        on shared Postgres (which can itself trigger SSL drops). One connection
        also keeps the whole discovery fast and atomic.
        """
        candidates = [
            "planreader_documents",
            "job_document_blobs",
            "job_documents",
            "documents",
            "job_files",
            "job_attachments",
            "attachments",
            "files",
        ]
        records: List[Dict[str, Any]] = []
        with self.connect() as conn:
            if self.kind == "postgres":
                cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
                cur.execute("SELECT table_name FROM information_schema.tables WHERE table_schema='public'")
                tables = {str(r["table_name"]) for r in cur.fetchall()}
                placeholder = "%s"
            else:
                cur = conn.cursor()
                cur.execute("SELECT name FROM sqlite_master WHERE type='table'")
                tables = {str(r[0]) for r in cur.fetchall()}
                placeholder = "?"
            for table in candidates:
                if table not in tables:
                    continue
                cols = self._columns_for_connection(cur, table)
                job_col = next((c for c in ["job_id", "project_id", "jobhub_job_id"] if c in cols), None)
                if not job_col:
                    continue
                id_col = "id" if "id" in cols else None
                name_col = next((c for c in ["file_name", "filename", "name", "original_name", "title"] if c in cols), None)
                mime_col = next((c for c in ["mime_type", "content_type", "file_type"] if c in cols), None)
                path_col = next((c for c in ["storage_path", "file_path", "path", "local_path"] if c in cols), None)
                url_col = next((c for c in ["download_url", "file_url", "url", "public_url"] if c in cols), None)
                blob_col = next((c for c in _DOCUMENT_BLOB_COLUMNS if c in cols), None)
                date_col = next((c for c in ["uploaded_at", "created_at", "date_uploaded"] if c in cols), None)
                select_parts = [f"{id_col} AS record_id" if id_col else "NULL AS record_id"]
                select_parts.append(f"{name_col} AS file_name" if name_col else "'' AS file_name")
                select_parts.append(f"{mime_col} AS mime_type" if mime_col else "'' AS mime_type")
                select_parts.append(f"{path_col} AS storage_path" if path_col else "'' AS storage_path")
                select_parts.append(f"{url_col} AS file_url" if url_col else "'' AS file_url")
                select_parts.append(f"{date_col} AS uploaded_at" if date_col else "'' AS uploaded_at")
                try:
                    cur.execute(
                        f"SELECT {', '.join(select_parts)} FROM {table} WHERE {job_col}={placeholder}",
                        (job_id,),
                    )
                    rows = cur.fetchall()
                except Exception:
                    continue
                for row in rows:
                    record = dict(row)
                    record["source_table"] = table
                    record["has_blob"] = bool(blob_col)
                    records.append(record)
        return records

    def fetch_document_blob(self, table: str, record_id: int) -> Optional[Any]:
        """Return a single document's file bytes (or None if absent).

        Binary columns come back as ``bytes``; text columns (e.g. base64 stored
        in ``job_document_blobs``) come back as ``str`` so the caller can decode
        them. Fetched one document at a time so discovery never holds every PDF
        in memory at once — the render free tier only has ~512 MB to work with.
        """
        if not record_id:
            return None
        safe = re.sub(r"[^A-Za-z0-9_]", "", table)
        with self.connect() as conn:
            cur = conn.cursor()
            cols = self._columns_for_connection(cur, safe)
            blob_col = next((c for c in _DOCUMENT_BLOB_COLUMNS if c in cols), None)
            if not blob_col:
                return None
            placeholder = "%s" if self.kind == "postgres" else "?"
            cur.execute(f"SELECT {blob_col} FROM {safe} WHERE id={placeholder}", (int(record_id),))
            row = cur.fetchone()
            if row is None:
                return None
            value = row[0]
            if isinstance(value, (bytes, bytearray, memoryview)):
                return bytes(value)
            if isinstance(value, str) and value.strip():
                return value
            return None

    def _columns_for_connection(self, cur, table: str) -> List[str]:
        safe = re.sub(r"[^A-Za-z0-9_]", "", table)
        if self.kind == "postgres":
            cur.execute(
                "SELECT column_name FROM information_schema.columns WHERE table_schema='public' AND table_name=%s ORDER BY ordinal_position",
                (safe,),
            )
            return [str(r["column_name"]) for r in cur.fetchall()]
        cur.execute(f"PRAGMA table_info({safe})")
        return [str(r[1]) for r in cur.fetchall()]


def get_jobhub_bridge() -> Optional[JobHubBridge]:
    if JOBHUB_DATABASE_URL:
        return JobHubBridge("postgres", JOBHUB_DATABASE_URL)
    if JOBHUB_DB_PATH and Path(JOBHUB_DB_PATH).exists():
        return JobHubBridge("sqlite", JOBHUB_DB_PATH)
    return None


def fetch_jobhub_jobs(bridge: JobHubBridge) -> List[Dict[str, Any]]:
    tables = set(bridge.table_names())
    if "jobs" not in tables:
        return []
    job_cols = set(bridge.columns("jobs"))
    builder_expr = "'' AS builder_client"
    join = ""
    if "builder_client_id" in job_cols and "builders_clients" in tables:
        join = " LEFT JOIN builders_clients b ON b.id=j.builder_client_id "
        builder_expr = "COALESCE(b.name,'') AS builder_client"
    elif "builder_client" in job_cols:
        builder_expr = "COALESCE(j.builder_client,'') AS builder_client"
    fields = [
        "j.id",
        "COALESCE(j.job_no,'') AS job_no" if "job_no" in job_cols else "CAST(j.id AS TEXT) AS job_no",
        "COALESCE(j.job_name,'') AS job_name" if "job_name" in job_cols else "'' AS job_name",
        builder_expr,
        "COALESCE(j.site_address,'') AS site_address" if "site_address" in job_cols else "'' AS site_address",
        "COALESCE(j.status,'') AS status" if "status" in job_cols else "'' AS status",
    ]
    return bridge.query(f"SELECT {', '.join(fields)} FROM jobs j {join} ORDER BY j.id DESC")


def authenticate_jobhub_user(bridge: JobHubBridge, username: str, password: str) -> Optional[Dict[str, Any]]:
    if "app_users" not in set(bridge.table_names()):
        return None
    cols = set(bridge.columns("app_users"))
    if not {"username", "password_hash"}.issubset(cols):
        return None
    active_filter = "AND COALESCE(active,1)=1" if "active" in cols else ""
    role_expr = "COALESCE(role,'employee') AS role" if "role" in cols else "'employee' AS role"
    employee_expr = "employee_id" if "employee_id" in cols else "NULL AS employee_id"
    rows = bridge.query(
        f"SELECT id,username,password_hash,{role_expr},{employee_expr} FROM app_users WHERE lower(username)=lower(?) {active_filter}",
        (username.strip(),),
    )
    if not rows:
        return None
    user = rows[0]
    if not verify_password(password, user.get("password_hash", "")):
        return None
    return user


def ensure_planreader_document_table(bridge: JobHubBridge) -> None:
    if bridge.kind == "postgres":
        sql = """
        CREATE TABLE IF NOT EXISTS planreader_documents (
            id SERIAL PRIMARY KEY,
            job_id INTEGER NOT NULL,
            file_name TEXT,
            mime_type TEXT,
            storage_path TEXT,
            source_app TEXT DEFAULT 'PlanReader',
            uploaded_by TEXT,
            uploaded_at TEXT,
            notes TEXT
        )
        """
    else:
        sql = """
        CREATE TABLE IF NOT EXISTS planreader_documents (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            job_id INTEGER NOT NULL,
            file_name TEXT,
            mime_type TEXT,
            storage_path TEXT,
            source_app TEXT DEFAULT 'PlanReader',
            uploaded_by TEXT,
            uploaded_at TEXT,
            notes TEXT
        )
        """
    bridge.execute(sql)


def discover_jobhub_document_records(bridge: JobHubBridge, job_id: int) -> List[Dict[str, Any]]:
    """Discover JobHub document records in a single connection, retrying once.

    The old implementation opened a new connection for the table probe plus one
    per candidate table, which hammered the shared Render Postgres and let its
    SSL layer drop connections. One connection keeps discovery fast and lets a
    transient SSL drop be absorbed by a single retry.
    """
    try:
        return bridge.discover_documents_for_job(job_id)
    except Exception:
        time.sleep(0.4)
        return bridge.discover_documents_for_job(job_id)


def copy_jobhub_document_to_workspace(record: Dict[str, Any], workspace_id: int) -> Tuple[bool, str]:
    out_dir = workspace_path(workspace_id) / "documents"
    file_name = safe_name(record.get("file_name") or f"jobhub_document_{record.get('record_id')}")
    mime_type = str(record.get("mime_type") or mimetypes.guess_type(file_name)[0] or "application/octet-stream")
    data: Optional[bytes] = None
    source_path = str(record.get("storage_path") or "").strip()
    source_url = str(record.get("file_url") or "").strip()
    blob = record.get("file_blob")
    try:
        if blob is not None:
            if isinstance(blob, (bytes, bytearray, memoryview)):
                data = bytes(blob)
            elif str(record.get("source_table") or "") == "job_document_blobs":
                data = base64.b64decode(str(blob))
            else:
                data = str(blob).encode("utf-8", errors="ignore")
        elif source_url.startswith("http://") or source_url.startswith("https://"):
            response = requests.get(source_url, timeout=30)
            response.raise_for_status()
            data = response.content
        elif source_path and Path(source_path).exists():
            data = Path(source_path).read_bytes()
        else:
            reasons = []
            if record.get("has_blob"):
                reasons.append("the JobHub record has a file column but no stored bytes")
            if source_path:
                reasons.append(f"stored path not reachable from this app: {source_path}")
            if source_url:
                reasons.append(f"stored URL not reachable from this app: {source_url}")
            detail = " · ".join(reasons) if reasons else "no reachable file data (no blob, URL or local path)"
            return False, f"{file_name}: metadata found, but its file is not reachable from this app — {detail}"
    except Exception as exc:
        return False, f"{file_name}: {exc}"
    digest = sha256_bytes(data)
    existing = lquery("SELECT id FROM documents WHERE workspace_id=? AND sha256=?", (workspace_id, digest))
    if existing:
        return True, f"{file_name}: already imported"
    target = out_dir / f"{digest[:12]}_{file_name}"
    target.write_bytes(data)
    lexecute(
        """
        INSERT INTO documents(workspace_id,source_type,jobhub_table,jobhub_record_id,file_name,mime_type,path,sha256,category,page_count,extracted_text,uploaded_at)
        VALUES(?,?,?,?,?,?,?,?,?,?,?,?)
        """,
        (
            workspace_id,
            "JobHub linked document",
            str(record.get("source_table") or ""),
            str(record.get("record_id") or ""),
            file_name,
            mime_type,
            str(target),
            digest,
            "Linked",
            0,
            "",
            str(record.get("uploaded_at") or now_stamp()),
        ),
    )
    return True, f"{file_name}: imported"


# -----------------------------------------------------------------------------
# Workspace and document processing
# -----------------------------------------------------------------------------


def open_jobhub_workspace(job: Dict[str, Any]) -> int:
    existing = lquery("SELECT id FROM workspaces WHERE jobhub_job_id=?", (int(job["id"]),))
    if existing:
        workspace_id = int(existing[0]["id"])
        lexecute(
            """UPDATE workspaces SET job_no=?,job_name=?,builder_client=?,site_address=?,updated_at=? WHERE id=?""",
            (
                job.get("job_no", ""),
                job.get("job_name", ""),
                job.get("builder_client", ""),
                job.get("site_address", ""),
                now_stamp(),
                workspace_id,
            ),
        )
        workspace_path(workspace_id)
        return workspace_id
    workspace_id = lexecute(
        """
        INSERT INTO workspaces(jobhub_job_id,job_no,job_name,builder_client,site_address,drawing_issue,estimator,status,executive_summary,created_at,updated_at)
        VALUES(?,?,?,?,?,?,?,?,?,?,?)
        """,
        (
            int(job["id"]),
            job.get("job_no", ""),
            job.get("job_name", ""),
            job.get("builder_client", ""),
            job.get("site_address", ""),
            "",
            "",
            "Draft",
            "",
            now_stamp(),
            now_stamp(),
        ),
    )
    workspace_path(workspace_id)
    return workspace_id


def create_standalone_workspace(job_no: str, job_name: str, builder: str, address: str) -> int:
    workspace_id = lexecute(
        """
        INSERT INTO workspaces(jobhub_job_id,job_no,job_name,builder_client,site_address,drawing_issue,estimator,status,executive_summary,created_at,updated_at)
        VALUES(NULL,?,?,?,?,?,?,?,?,?,?)
        """,
        (job_no, job_name, builder, address, "", "", "Draft", "", now_stamp(), now_stamp()),
    )
    workspace_path(workspace_id)
    return workspace_id


def classify_page(text: str, file_name: str, page_no: int) -> Tuple[str, str]:
    lower = f"{file_name} {text}".lower()
    page_type = "Other"
    patterns = [
        ("Title / Drawing Register", ["drawing register", "drawing schedule", "title sheet"]),
        ("Reflected Ceiling Plan", ["reflected ceiling", "rcp", "ceiling plan"]),
        ("Floor Plan", ["floor plan", "proposed plan", "general arrangement"]),
        ("Roof Plan", ["roof plan"]),
        ("Elevation", ["elevation", "north elev", "south elev", "east elev", "west elev"]),
        ("Render / Artist's Impression", ["artist's impression", "artists impression", "artists rendering", "artist rendering", "3d view", "3d render", "concept image", "perspective render", "concept render", "visualisation", "visualization", "render", "impression"]),
        ("Section", ["section", "cross section"]),
        ("Door / Window Schedule", ["door schedule", "window schedule", "door elevations"]),
        ("Finishes Schedule", ["finish schedule", "finishes schedule", "colour schedule", "paint schedule"]),
        ("Specification", ["specification", "painting specification", "architectural specification"]),
        ("Structural", ["structural", "steel framing", "footing"]),
        ("Services", ["mechanical", "electrical", "hydraulic", "fire services"]),
        ("Landscape / Civil", ["civil", "landscape", "line marking", "pavement"]),
    ]
    for candidate, words in patterns:
        if any(word in lower for word in words):
            page_type = candidate
            break
    drawing_match = re.search(r"\b([A-Z]{1,3}\d{2,4}(?:[-/.][A-Z0-9]+)?)\b", text)
    label = drawing_match.group(1) if drawing_match else f"Page {page_no}"
    return page_type, label


def placeholder_image(title: str, body: str, target: Path) -> Tuple[int, int]:
    width, height = 1400, 1000
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    try:
        title_font = ImageFont.truetype("arial.ttf", 42)
        body_font = ImageFont.truetype("arial.ttf", 24)
    except Exception:
        title_font = ImageFont.load_default()
        body_font = ImageFont.load_default()
    draw.rectangle((0, 0, width, 95), fill=PB_DARK)
    draw.text((35, 25), title[:80], fill="white", font=title_font)
    wrapped = textwrap.wrap(normalise_whitespace(body), width=95)[:30]
    y = 135
    for line in wrapped:
        draw.text((40, y), line, fill="black", font=body_font)
        y += 29
    img.save(target)
    return width, height


# -----------------------------------------------------------------------------
# Plan measurement mapping: auto-map take-off rows as draggable measurement lines
# -----------------------------------------------------------------------------

LINE_COLOUR_PALETTE = [
    "#1f6fb2", "#1f7a4d", "#b33a3a", "#b26a00", "#7a3fb2",
    "#0e7c86", "#a03b6f", "#5b7a1f", "#8a4b08", "#33477a",
]

_LINE_COLOUR_KEYWORDS = [
    ("ceiling", "#1f6fb2"),
    ("wall", "#1f7a4d"),
    ("door", "#b26a00"),
    ("skirt", "#7a3fb2"),
    ("trim", "#7a3fb2"),
    ("fascia", "#0e7c86"),
    ("soffit", "#a03b6f"),
    ("cladding", "#5b7a1f"),
    ("fence", "#8a4b08"),
    ("metal", "#33477a"),
    ("steel", "#33477a"),
]

_LINEAR_UNITS = {
    "lm", "m", "ml", "lin m", "lineal m", "linear metre", "lineal metre",
    "mtr", "metres", "meters", "lf", "lfm",
}

_COUNT_UNITS = {"no", "no.", "ea", "each", "count", "1", "door", "doors"}


def line_colour_for(section: Any, element: Any) -> str:
    """Pick a stable colour for a take-off row's measurement line."""
    key = f"{str(section or '')} {str(element or '')}".lower()
    for word, colour in _LINE_COLOUR_KEYWORDS:
        if word in key:
            return colour
    if "external" in str(section or "").lower():
        return "#b33a3a"
    if "internal" in str(section or "").lower():
        return "#1f7a4d"
    digest = hashlib.md5(key.encode("utf-8", "ignore")).digest()
    return LINE_COLOUR_PALETTE[digest[0] % len(LINE_COLOUR_PALETTE)]


def normalise_line_unit(unit: Any) -> str:
    u = str(unit or "").strip().lower()
    if not u:
        return ""
    if u in {"m2", "sqm", "sq m", "m²", "square metre", "square metres"}:
        return "m2"
    if u in _LINEAR_UNITS:
        return "m"
    if u in _COUNT_UNITS:
        return "count"
    return u


_SCALE_RATIO_RE = re.compile(r"(?:^|[^\d])1\s*[:/]\s*(\d{2,4})(?![:\d])")
_SCALE_IN_RE = re.compile(r"\b1\s*in\s*(\d{2,4})\b", re.IGNORECASE)


def auto_detect_scale(page: Dict[str, Any]) -> Optional[Dict[str, Any]]:
    """Estimate pixels-per-metre from a drawing scale annotation (e.g. 1:100).

    Pages are rasterised from PDF points at ``render_zoom`` (capped so a page's
    long edge never exceeds 2400 px; small sheets stay at 1.7x). At that zoom
    1 px = 1/zoom pt and 1 pt = 25.4/72 mm. At scale 1:N, 1 mm on the drawing
    equals N/1000 m real, giving px_per_m = 2834.646 * render_zoom / N. Rows
    recorded before the zoom was stored assume 1.7x, the historical default.
    This is only a starting estimate - the user confirms it in the mapper
    before measured quantities are trusted.
    """
    source = str(page.get("scale_text") or "")
    text = str(page.get("extracted_text") or "")
    match = _SCALE_RATIO_RE.search(source) or _SCALE_RATIO_RE.search(text) or _SCALE_IN_RE.search(text)
    if not match:
        return None
    ratio = int(match.group(1))
    if not (10 <= ratio <= 2000):
        return None
    zoom = to_float(page.get("render_zoom")) or 1.7
    px_per_m = 2834.6458 * zoom / ratio
    return {"ratio": ratio, "px_per_m": round(px_per_m, 3), "source": match.group(0).strip()}


def scale_gate_issues(workspace_id: int) -> List[Dict[str, Any]]:
    """Pages that feed the take-off but do not have a calibrated scale yet.

    The scale gate protects measured quantities: a row is only trustworthy when
    every drawing page it was measured from has a confirmed ``px_per_m``. This
    lists every selected page that is referenced by take-off rows or mapped
    zones while still lacking a scale.
    """
    row_sources: set = set()
    rows = ldf("SELECT DISTINCT source_page FROM takeoff_rows WHERE workspace_id=?", (workspace_id,))
    if not rows.empty:
        row_sources = {str(x).strip() for x in rows["source_page"].tolist() if str(x).strip()}
    zone_page_ids: set = set()
    zones = ldf("SELECT DISTINCT page_id FROM mapped_zones WHERE workspace_id=?", (workspace_id,))
    if not zones.empty:
        zone_page_ids = {int(x) for x in zones["page_id"].tolist()}
    issues: List[Dict[str, Any]] = []
    pages = ldf("SELECT id,page_label,page_type,px_per_m FROM pages WHERE workspace_id=? AND selected=1", (workspace_id,))
    for p in pages.itertuples():
        if to_float(p.px_per_m) > 0:
            continue
        label = str(p.page_label or "")
        if label.strip() in row_sources or int(p.id) in zone_page_ids:
            issues.append({
                "page_id": int(p.id),
                "page_label": label,
                "page_type": str(p.page_type or ""),
                "px_per_m": to_float(p.px_per_m),
            })
    return issues


def scale_gate_blocked(workspace_id: int) -> bool:
    return bool(scale_gate_issues(workspace_id))


def _line_grid_positions(count: int) -> List[Tuple[float, float, float, float]]:
    """Spread ``count`` horizontal line segments across the page in a grid."""
    if count <= 0:
        return []
    cols = 2 if count <= 6 else (3 if count <= 12 else 4)
    rows = max(1, math.ceil(count / cols))
    margin = 6.0
    cell_w = (100.0 - 2 * margin) / cols
    cell_h = (100.0 - 2 * margin) / rows
    out = []
    for i in range(count):
        col = i // rows
        row = i % rows
        cx = margin + col * cell_w + cell_w * 0.5
        cy = margin + row * cell_h + cell_h * 0.5
        width_pct = min(16.0, cell_w * 0.6)
        out.append((max(0.0, cx - width_pct / 2), cy, min(100.0, cx + width_pct / 2), cy))
    return out


def takeoff_rows_for_mapper(workspace_id: int) -> List[Dict[str, Any]]:
    """Take-off rows offered as draw targets in the line mapper.

    Each row becomes a draw target: a lineal-metre row wants line(s), an area
    row wants closed outline(s). Rows whose unit is neither lineal nor area
    (e.g. counts) are excluded from drawing.
    """
    rows = ldf("SELECT * FROM takeoff_rows WHERE workspace_id=? ORDER BY section, element, location, id", (workspace_id,))
    out: List[Dict[str, Any]] = []
    for r in rows.itertuples(index=False):
        unit = normalise_line_unit(r.unit)
        if unit not in {"m", "m2"}:
            continue
        label = f"{r.section or ''} · {r.element or ''} · {r.location or ''}".strip(" ·")
        detail = " · ".join(x for x in [str(r.substrate or ""), str(r.finish_system or "")] if str(x).strip())
        if detail:
            label = f"{label} ({detail})"
        group = f"{r.section or 'Unassigned'} · {r.element or 'Items'}"
        out.append({
            "id": int(r.id),
            "label": label,
            "group": group,
            "unit": unit,
            "colour": line_colour_for(r.section, r.element),
            "quantity": to_float(r.quantity),
            "status": str(r.quantity_status or ""),
            "substrate": str(r.substrate or ""),
            "finish_system": str(r.finish_system or ""),
        })
    return out


def auto_map_measurements(workspace_id: int, page_id: int, px_per_m: float) -> List[Dict[str, Any]]:
    """Build a measurement shape for every take-off row on this page.

    Shapes are created empty and centred - the user clicks on the plan to draw
    the real footprint/ceiling/wall/doors lines. A lineal-metre row gets a
    short horizontal line; an area row gets a closed square outline; both are
    only starting points the user replaces with a real drawing.
    """
    rows = ldf("SELECT * FROM takeoff_rows WHERE workspace_id=? ORDER BY id", (workspace_id,))
    pages = lquery("SELECT width_px,height_px FROM pages WHERE id=?", (page_id,))
    img_w = int(pages[0].get("width_px") or 1000) if pages else 1000
    lines: List[Dict[str, Any]] = []
    positions = _line_grid_positions(len(rows))
    pxpm = to_float(px_per_m)
    for i, (r, (x1, y1, x2, y2)) in enumerate(zip(rows.itertuples(index=False), positions)):
        unit = normalise_line_unit(r.unit)
        if unit == "m2":
            label = f"{r.location or ''} · {r.element or ''}".strip(" ·")
            half = 8.0
            cx, cy = (x1 + x2) / 2, (y1 + y2) / 2
            lines.append({
                "id": f"auto_{i}",
                "takeoff_row_id": int(r.id),
                "label": label,
                "unit": unit,
                "colour": line_colour_for(r.section, r.element),
                "kind": "polygon",
                "points": [
                    [max(0.0, cx - half), max(0.0, cy - half)],
                    [min(100.0, cx + half), max(0.0, cy - half)],
                    [min(100.0, cx + half), min(100.0, cy + half)],
                    [max(0.0, cx - half), min(100.0, cy + half)],
                ],
                "area_m2": 0.0, "perimeter_m": 0.0, "length_m": 0.0,
                "quantity_status": "Placeholder", "moved": 0,
                "notes": "Placeholder outline - click points on the plan to draw the real area.",
            })
            continue
        label = f"{r.location or ''} · {r.element or ''}".strip(" ·")
        length_m = 0.0
        if pxpm > 0:
            pct_len = min(20.0, max(4.0, 2.0 * pxpm / img_w * 100.0))
            cx = (x1 + x2) / 2
            x1 = max(0.0, cx - pct_len / 2)
            x2 = min(100.0, cx + pct_len / 2)
        lines.append({
            "id": f"auto_{i}",
            "takeoff_row_id": int(r.id),
            "label": label,
            "unit": unit,
            "colour": line_colour_for(r.section, r.element),
            "kind": "line",
            "x1": round(x1, 3), "y1": round(y1, 3),
            "x2": round(x2, 3), "y2": round(y2, 3),
            "points": [],
            "length_m": round(length_m, 3), "area_m2": 0.0, "perimeter_m": 0.0,
            "quantity_status": "Placeholder", "moved": 0,
            "notes": "Placeholder line - click two points on the plan to draw the real length.",
        })
    return lines


def save_measurement_lines(workspace_id: int, page_id: int, lines: Sequence[Dict[str, Any]]) -> Dict[str, Any]:
    """Persist drawn measurement shapes; sync take-off quantities.

    Shapes are stored as-is in ``measurement_lines``. After saving, each
    take-off row that has at least one drawn shape gets its quantity recomputed
    from the drawing: lineal-metre rows sum line lengths, area rows sum polygon
    areas. Rows without shapes keep their original values.
    """
    lexecute("DELETE FROM measurement_lines WHERE page_id=?", (page_id,))
    saved = 0
    synced = 0
    now = now_stamp()
    for ln in lines or []:
        kind = "polygon" if str(ln.get("kind") or "line") == "polygon" else "line"
        row_id = ln.get("takeoff_row_id")
        if row_id is not None:
            try:
                row_id = int(row_id)
            except (TypeError, ValueError):
                row_id = None
            if row_id == 0:
                row_id = None
        points = ln.get("points") or []
        if isinstance(points, (list, tuple)):
            import json as _json
            points_json = _json.dumps([[round(to_float(p[0]), 3), round(to_float(p[1]), 3)] for p in points])
        else:
            points_json = str(points or "")
        length_m = to_float(ln.get("length_m"))
        area_m2 = to_float(ln.get("area_m2"))
        perimeter_m = to_float(ln.get("perimeter_m"))
        lexecute(
            """INSERT INTO measurement_lines(workspace_id,page_id,takeoff_row_id,label,unit,colour,kind,x1,y1,x2,y2,points,length_m,area_m2,perimeter_m,quantity_status,moved,notes,created_at)
               VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
            (
                workspace_id, page_id, row_id,
                str(ln.get("label") or ""), normalise_line_unit(ln.get("unit")), str(ln.get("colour") or ""),
                kind,
                to_float(ln.get("x1")), to_float(ln.get("y1")),
                to_float(ln.get("x2")), to_float(ln.get("y2")),
                points_json,
                length_m, area_m2, perimeter_m,
                str(ln.get("quantity_status") or "Mapped"), 1 if ln.get("moved") else 0,
                str(ln.get("notes") or ""), now,
            ),
        )
        saved += 1

    # Aggregate per take-off row so multiple shapes on the same row SUM to the
    # row quantity (a wall drawn as several segments, or a floor as several
    # polygons), matching the documented "sum line lengths / sum polygon areas".
    agg: Dict[int, Dict[str, float]] = {}
    for ln in lines or []:
        row_id = ln.get("takeoff_row_id")
        if row_id is not None:
            try:
                row_id = int(row_id)
            except (TypeError, ValueError):
                row_id = None
            if row_id == 0:
                row_id = None
        if row_id is None:
            continue
        length_m = to_float(ln.get("length_m"))
        area_m2 = to_float(ln.get("area_m2"))
        perimeter_m = to_float(ln.get("perimeter_m"))
        if length_m <= 0 and area_m2 <= 0:
            continue
        bucket = agg.setdefault(row_id, {"unit": normalise_line_unit(ln.get("unit")), "length_m": 0.0, "area_m2": 0.0, "perimeter_m": 0.0})
        bucket["length_m"] += length_m
        bucket["area_m2"] += area_m2
        bucket["perimeter_m"] += perimeter_m
    for row_id, b in agg.items():
        unit = b["unit"]
        if unit == "m2" and b["area_m2"] > 0:
            value = round(b["area_m2"], 3)
            qty_for = ldf("SELECT section, element FROM takeoff_rows WHERE id=? AND workspace_id=?", (row_id, workspace_id))
            if not qty_for.empty:
                el_key = f"{qty_for.iloc[0]['section'] or ''} {qty_for.iloc[0]['element'] or ''}".lower()
                if b["perimeter_m"] > 0 and ("external" in el_key or "cladding" in el_key or "render" in el_key):
                    value = round(b["perimeter_m"] * WALL_HEIGHT_M, 3)
            lexecute(
                "UPDATE takeoff_rows SET quantity=?, quantity_status='Mapped', updated_at=? WHERE id=? AND workspace_id=?",
                (value, now, row_id, workspace_id),
            )
            synced += 1
        elif unit == "m" and b["length_m"] > 0:
            lexecute(
                "UPDATE takeoff_rows SET quantity=?, quantity_status='Mapped', updated_at=? WHERE id=? AND workspace_id=?",
                (round(b["length_m"], 3), now, row_id, workspace_id),
            )
            synced += 1
    return {"saved": saved, "synced": synced}


# Assumed wall height (m) used to convert an external building-envelope outline
# (a footprint polygon) into external wall/cladding area: perimeter * height.
WALL_HEIGHT_M = 2.7


def auto_detect_building_envelope(image_path: Any, min_area_pct: float = 0.4, max_contours: int = 3) -> List[Dict[str, Any]]:
    """Detect the outer building envelope(s) on a plan page.

    The detector is iterative: it tests several dark-line thresholds and a
    ladder of morphological closing sizes so gaps in the wall lines (door
    openings, dashed partitions) are bridged before the exterior boundary is
    traced. The largest stable interior region is taken as the building
    envelope. The result is returned as a polygon in percentage image
    coordinates (matching the mapper's coordinate space) with its area as a
    percentage of the page, so the user can place measurement boxes around the
    cladding / external substrate envelope and the floor area of each level.

    The core is pure numpy so it works whether or not OpenCV is installed;
    OpenCV is used only to speed up the morphology when it is available.
    """
    path = Path(str(image_path or ""))
    if not path.exists():
        return []
    try:
        from PIL import Image as _PILImage
        img = _PILImage.open(path).convert("L")
    except Exception:
        return []
    w, h = img.size
    if w < 16 or h < 16:
        return []
    gray = np.asarray(img, dtype=np.uint8)
    min_px = (min_area_pct / 100.0) * w * h
    results: List[Dict[str, Any]] = []
    tried: List[Tuple[int, int]] = []
    thresholds = _envelope_thresholds(gray)
    for thr_val in thresholds:
        dark = gray < thr_val
        # skip degenerate thresholds that catch nothing or nearly everything
        frac = float(dark.sum()) / dark.size
        if frac <= 0.002 or frac >= 0.75:
            continue
        for k in _ENVELOPE_CLOSE_LADDER:
            if (thr_val, k) in tried:
                continue
            tried.append((thr_val, k))
            closed = _morph_close_binary(dark, k)
            outside = _flood_background(closed)
            interior = ~outside
            area = int(interior.sum())
            if area < min_px:
                continue
            outline = _moore_outline(interior)
            if len(outline) < 4:
                continue
            eps = 0.006 * max(w, h)
            simplified = _rdp_simplify(outline, eps)
            if len(simplified) < 3:
                simplified = _rough_polygon(outline)
            pct_pts = [
                [round(float(x) / w * 100.0, 3), round(float(y) / h * 100.0, 3)]
                for x, y in simplified
            ]
            perim = sum(
                math.hypot(pct_pts[j][0] - pct_pts[(j + 1) % len(pct_pts)][0],
                           pct_pts[j][1] - pct_pts[(j + 1) % len(pct_pts)][1])
                for j in range(len(pct_pts))
            ) / 100.0 * max(w, h)
            results.append({
                "points": pct_pts,
                "area_pct": round(polygon_area(pct_pts) / 100.0, 2),
                "perimeter_pct": round(perim / max(w, h) * 100.0, 2),
            })
            break
    results.sort(key=lambda d: -d["area_pct"])
    # de-duplicate near-identical outlines produced by adjacent thresholds
    deduped: List[Dict[str, Any]] = []
    for r in results:
        if deduped and abs(deduped[-1]["area_pct"] - r["area_pct"]) <= 2.0:
            continue
        deduped.append(r)
    return deduped[:max_contours]


_ENVELOPE_CLOSE_LADDER = (9, 21, 45, 81, 121, 181)


def _envelope_thresholds(gray: np.ndarray) -> List[int]:
    """A small ladder of dark-line thresholds, from conservative to permissive."""
    vals = [float(np.percentile(gray, p)) for p in (55, 70, 82, 90)]
    out: List[int] = []
    for v in vals:
        if v < 6 or v > 245:
            continue
        t = int(v)
        if t not in out:
            out.append(t)
    if not out:
        hist = np.bincount(gray.ravel(), minlength=256)
        total = int(gray.size)
        acc = 0
        for i in range(256):
            acc += int(hist[i])
            if acc >= total * 0.8:
                out.append(max(2, i))
                break
    return out or [60, 90]


def _morph_close_binary(mask: np.ndarray, k: int) -> np.ndarray:
    """Binary morphological close (dilate then erode) of a bool mask."""
    if cv2 is not None:
        kernel = np.ones((k, k), np.uint8)
        arr = mask.astype(np.uint8) * 255
        closed = cv2.morphologyEx(arr, cv2.MORPH_CLOSE, kernel)
        return closed > 127
    from PIL import Image, ImageFilter
    img = Image.fromarray((mask * 255).astype(np.uint8))
    dil = img.filter(ImageFilter.MaxFilter(k))
    ero = dil.filter(ImageFilter.MinFilter(k))
    return np.asarray(ero, dtype=np.uint8) > 127


def _dilate4(mask: np.ndarray) -> np.ndarray:
    out = np.zeros_like(mask)
    out[1:, :] |= mask[:-1, :]
    out[:-1, :] |= mask[1:, :]
    out[:, 1:] |= mask[:, :-1]
    out[:, :-1] |= mask[:, 1:]
    return out


def _flood_background(fg: np.ndarray, max_iter: int = 4000) -> np.ndarray:
    """Flood-fill the paper background from the image borders (4-connected)."""
    h, w = fg.shape
    bg = ~fg
    out = np.zeros_like(bg)
    if h <= 2 or w <= 2:
        return out
    out[0, :] = True
    out[h - 1, :] = True
    out[:, 0] = True
    out[:, w - 1] = True
    out &= bg
    for _ in range(max_iter):
        nxt = _dilate4(out) & bg
        if nxt.sum() == out.sum():
            return out
        out = nxt
    return out


_DIR8 = ((0, 1), (-1, 1), (-1, 0), (-1, -1), (0, -1), (1, -1), (1, 0), (1, 1))


def _moore_outline(mask: np.ndarray) -> List[Tuple[float, float]]:
    """Trace the outer boundary of the foreground region (Moore neighbour trace).

    Returns a closed polygon of (x, y) pixel coordinates.
    """
    ys, xs = np.nonzero(mask)
    if len(ys) == 0:
        return []
    order = np.lexsort((xs, ys))
    start = (int(ys[order[0]]), int(xs[order[0]]))
    boundary: List[Tuple[float, float]] = [(float(start[1]), float(start[0]))]
    prev = (start[0], start[1] - 1)
    cur = start
    moved = False
    for _ in range(400000):
        dy = prev[0] - cur[0]
        dx = prev[1] - cur[1]
        try:
            base = _DIR8.index((dy, dx))
        except ValueError:
            break
        nxt = None
        for i in range(1, 9):
            dy2, dx2 = _DIR8[(base + i) % 8]
            ny, nx = cur[0] + dy2, cur[1] + dx2
            if 0 <= ny < mask.shape[0] and 0 <= nx < mask.shape[1] and mask[ny, nx]:
                nxt = (ny, nx)
                break
        if nxt is None:
            break
        if moved and nxt == start:
            boundary.append((float(start[1]), float(start[0])))
            break
        moved = True
        prev = cur
        cur = nxt
        boundary.append((float(cur[1]), float(cur[0])))
        if len(boundary) >= 2 and cur == start:
            break
    return boundary


def _point_line_dist(p: Sequence[float], a: Sequence[float], b: Sequence[float]) -> float:
    x0, y0 = p
    x1, y1 = a
    x2, y2 = b
    dx = x2 - x1
    dy = y2 - y1
    if dx == 0 and dy == 0:
        return math.hypot(x0 - x1, y0 - y1)
    return abs(dy * x0 - dx * y0 + x2 * y1 - y2 * x1) / math.hypot(dx, dy)


def _rdp_simplify(points: Sequence[Sequence[float]], epsilon: float) -> List[Tuple[float, float]]:
    """Ramer-Douglas-Peucker line simplification on an open polygon chain."""
    pts = [(float(p[0]), float(p[1])) for p in points]
    if len(pts) <= 2:
        return pts
    start, end = pts[0], pts[-1]
    dmax, idx = 0.0, 0
    for i in range(1, len(pts) - 1):
        d = _point_line_dist(pts[i], start, end)
        if d > dmax:
            dmax, idx = d, i
    if dmax > epsilon:
        left = _rdp_simplify(pts[:idx + 1], epsilon)
        right = _rdp_simplify(pts[idx:], epsilon)
        return left[:-1] + right
    return [start, end]


def _rough_polygon(outline: List[Tuple[float, float]]) -> List[Tuple[float, float]]:
    """Fallback: axis-aligned bounding box of a traced outline."""
    if not outline:
        return []
    xs = [p[0] for p in outline]
    ys = [p[1] for p in outline]
    x0, x1, y0, y1 = min(xs), max(xs), min(ys), max(ys)
    return [(x0, y0), (x1, y0), (x1, y1), (x0, y1), (x0, y0)]


def _polygon_px(shape: Dict[str, Any], width_px: int, height_px: int) -> List[Tuple[float, float]]:
    pts = shape.get("points") or []
    if isinstance(pts, str):
        try:
            pts = json.loads(pts)
        except Exception:
            pts = []
    return [
        (float(p[0]) / 100.0 * float(width_px or 1), float(p[1]) / 100.0 * float(height_px or 1))
        for p in pts
    ]


def auto_detect_envelope_shapes(workspace_id: int, page: Dict[str, Any], px_per_m: float, level: str) -> List[Dict[str, Any]]:
    """Create drawable measurement shapes for the detected building envelope.

    For every detected outline a polygon is created for two take-off rows:

    - **External walls / cladding** (External, m²) - the outer envelope the user
      assigns to the real external substrate (render, brick, cladding, ...).
    - **Floor plan** (Internal, m²) for ``level`` - the internal floor area, the
      flat-rate basis for internal works including paint systems.

    Rows are created on the fly when missing so drawing can start immediately.
    """
    detections = auto_detect_building_envelope(str(page.get("image_path") or ""))
    if not detections:
        return []
    label_ext = " · ".join(x for x in [str(level or ""), "external envelope"] if str(x).strip())
    label_flr = " · ".join(x for x in [str(level or ""), "floor area"] if str(x).strip())
    ext_id = _ensure_mapper_row(
        workspace_id, "External", "External walls / cladding", label_ext,
        "Render", "Exterior acrylic", "m²", page.get("page_label", ""),
    )
    flr_id = _ensure_mapper_row(
        workspace_id, "Internal", "Floor plan", label_flr,
        "Concrete floor", "Interior coatings", "m²", page.get("page_label", ""),
    )
    pxpm = to_float(px_per_m)
    w_px = int(page.get("width_px") or 1000)
    h_px = int(page.get("height_px") or 1000)
    shapes: List[Dict[str, Any]] = []
    for i, det in enumerate(detections):
        pts = det["points"]
        px_pts = [(p[0] / 100.0 * w_px, p[1] / 100.0 * h_px) for p in pts]
        area_m2 = polygon_area(px_pts) / (pxpm * pxpm) if pxpm > 0 else 0.0
        perimeter_m = 0.0
        if pxpm > 0:
            perimeter_m = sum(
                math.hypot(px_pts[j][0] - px_pts[(j + 1) % len(px_pts)][0], px_pts[j][1] - px_pts[(j + 1) % len(px_pts)][1])
                for j in range(len(px_pts))
            ) / pxpm
        shapes.append({
            "id": f"auto_ext_{i}", "takeoff_row_id": int(ext_id),
            "label": f"External walls / cladding · {label_ext} (auto)", "unit": "m2",
            "colour": line_colour_for("External", "External walls / cladding"),
            "kind": "polygon", "points": pts,
            "length_m": 0.0, "area_m2": round(area_m2, 3), "perimeter_m": round(perimeter_m, 3),
            "quantity_status": "Detected", "moved": 1,
            "notes": "Auto-detected envelope - drag corners or move to match the real plan, then pick the external substrate row for each face.",
        })
        shapes.append({
            "id": f"auto_flr_{i}", "takeoff_row_id": int(flr_id),
            "label": f"Floor plan · {label_flr} (auto)", "unit": "m2",
            "colour": line_colour_for("Internal", "Floor plan"),
            "kind": "polygon", "points": pts,
            "length_m": 0.0, "area_m2": round(area_m2, 3), "perimeter_m": round(perimeter_m, 3),
            "quantity_status": "Detected", "moved": 1,
            "notes": "Auto-detected floor area - adjust to the real level footprint.",
        })
    return shapes


def _ensure_mapper_row(workspace_id: int, section: str, element: str, location: str,
                       substrate: str, finish_system: str, unit: str, source_page: str) -> int:
    """Return an existing drawable take-off row or create it (used by auto-detect)."""
    existing = ldf(
        "SELECT id FROM takeoff_rows WHERE workspace_id=? AND section=? AND element=? AND location=? AND unit=? ORDER BY id",
        (workspace_id, section, element, location, unit),
    )
    if not existing.empty:
        return int(existing.iloc[0]["id"])
    unit_norm = normalise_line_unit(unit)
    rate = default_rate_for(substrate, element, finish_system, unit)
    return lexecute(
        """INSERT INTO takeoff_rows(workspace_id,section,element,location,substrate,finish_system,quantity,unit,quantity_status,source_page,source_reference,inclusion_status,coats,coverage_m2_per_litre,productivity_m2_per_hour,rate_per_unit,confidence,notes,row_role,created_at,updated_at)
           VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
        (workspace_id, section, element, location, substrate, finish_system, 0, unit_norm, "To measure", source_page, "", "INCLUSION",
         3 if unit_norm == "m2" else 2, 12, 8, rate, "To review", f"Auto-detected from {source_page}.", "", now_stamp(), now_stamp()),
    )


def render_measurement_overlay(page: Dict[str, Any], lines: Sequence[Dict[str, Any]]) -> Optional[Image.Image]:
    """Static PNG snapshot of the measurement lines drawn on the page."""
    path = Path(str(page.get("image_path") or ""))
    if not path.exists():
        return None
    img = Image.open(path).convert("RGBA")
    overlay = Image.new("RGBA", img.size, (0, 0, 0, 0))
    draw = ImageDraw.Draw(overlay, "RGBA")
    w, h = img.size
    if hasattr(lines, "to_dict"):
        lines = lines.to_dict("records")
    for ln in lines or []:
        colour = str(ln.get("colour") or "#1f6fb2")
        hexc = colour.lstrip("#")
        try:
            rgb = tuple(int(hexc[i:i + 2], 16) for i in (0, 2, 4))
        except Exception:
            rgb = (31, 111, 178)
        label = str(ln.get("label") or "")
        if str(ln.get("kind") or "line") == "polygon":
            pts = ln.get("points") or []
            if isinstance(pts, str):
                import json as _json
                try:
                    pts = _json.loads(pts)
                except Exception:
                    pts = []
            if len(pts) >= 3:
                coords = [(to_float(p[0]) / 100 * w, to_float(p[1]) / 100 * h) for p in pts]
                draw.polygon(coords, outline=rgb + (255,), width=4)
                for cx, cy in coords:
                    draw.ellipse((cx - 5, cy - 5, cx + 5, cy + 5), fill=rgb + (255,))
                xs = [c[0] for c in coords]
                ys = [c[1] for c in coords]
                if label:
                    draw.text((min(xs) + 6, min(ys) + 6), label, fill=(20, 20, 20, 255))
            continue
        x1 = to_float(ln.get("x1")) / 100 * w
        y1 = to_float(ln.get("y1")) / 100 * h
        x2 = to_float(ln.get("x2")) / 100 * w
        y2 = to_float(ln.get("y2")) / 100 * h
        draw.line((x1, y1, x2, y2), fill=rgb + (255,), width=4)
        r = 6
        draw.ellipse((x1 - r, y1 - r, x1 + r, y1 + r), fill=rgb + (255,))
        draw.ellipse((x2 - r, y2 - r, x2 + r, y2 + r), fill=rgb + (255,))
        if label:
            draw.text((min(x1, x2) + 6, min(y1, y2) + 6), label, fill=(20, 20, 20, 255))
    return Image.alpha_composite(img, overlay)


_PDF_RENDER_LONG_EDGE_PX = 2400
_AI_IMAGE_LONG_EDGE_PX = 1600


def _pdf_render_zoom(page: Any) -> float:
    """Zoom capped so a single rendered page can never exceed the long-edge limit."""
    rect = page.rect
    long_edge = max(float(rect.width), float(rect.height)) or 1.0
    return min(1.7, _PDF_RENDER_LONG_EDGE_PX / long_edge)


_RENDER_WORKER_PATH = str(Path(__file__).resolve().parent / "pb_render_worker.py")
_RENDER_PAGE_TIMEOUT_SEC = 180
_RENDER_RETRY_ZOOM_RATIO = 0.5


class _RenderWorkerSession:
    """One long-lived child process that renders pages from a single PDF.

    The child keeps the document open and rasterises one page at a time on
    demand, so a whole batch costs a single Python/Fitz startup and the parent
    never holds the document open while a page is being rendered. A crash or
    memory spike in the MuPDF rasteriser kills only this worker, never the app.
    """

    def __init__(self, pdf_path: str) -> None:
        self.pdf_path = pdf_path
        self._queue: queue.Queue[str] = queue.Queue()
        self._proc = subprocess.Popen(
            [sys.executable, "-u", _RENDER_WORKER_PATH, pdf_path],
            stdin=subprocess.PIPE,
            stdout=subprocess.PIPE,
            stderr=subprocess.STDOUT,
            text=True,
        )
        self._alive = True
        self._thread = threading.Thread(target=self._read_lines, daemon=True)
        self._thread.start()

    def _read_lines(self) -> None:
        try:
            if self._proc.stdout is not None:
                for line in self._proc.stdout:
                    stripped = line.strip()
                    if stripped:
                        self._queue.put(stripped)
        except Exception:
            pass
        finally:
            self._queue.put("__EOF__")

    def render(self, page_no: int, zoom: float, out_path: str, timeout: float) -> Tuple[bool, int, int, str]:
        if not self._alive:
            return False, 0, 0, "render worker is not running"
        if self._proc.poll() is not None:
            self._alive = False
            return False, 0, 0, "render worker exited"
        try:
            if self._proc.stdin is not None:
                self._proc.stdin.write(f"RENDER {page_no} {zoom} {shlex.quote(out_path)}\n")
                self._proc.stdin.flush()
        except Exception as exc:
            self._alive = False
            return False, 0, 0, f"could not reach render worker: {exc}"
        try:
            line = self._queue.get(timeout=timeout)
        except queue.Empty:
            self.close()
            return False, 0, 0, f"page {page_no} render timed out after {int(timeout)}s"
        if line == "__EOF__":
            self._alive = False
            return False, 0, 0, "render worker exited unexpectedly"
        if line.startswith("OK "):
            parts = line.split()
            if len(parts) == 3:
                try:
                    return True, int(parts[1]), int(parts[2]), ""
                except ValueError:
                    pass
            return True, 0, 0, ""
        if line.startswith("ERR "):
            return False, 0, 0, line[4:].strip()[-300:]
        return False, 0, 0, line[-300:]

    def close(self) -> None:
        self._alive = False
        try:
            if self._proc.poll() is None:
                if self._proc.stdin is not None:
                    try:
                        self._proc.stdin.write("QUIT\n")
                        self._proc.stdin.flush()
                    except Exception:
                        pass
                try:
                    self._proc.wait(timeout=5)
                except subprocess.TimeoutExpired:
                    self._proc.kill()
                    self._proc.wait(timeout=5)
        except Exception:
            try:
                self._proc.kill()
            except Exception:
                pass
        # Let the reader thread finish consuming the pipe before the wrapper is
        # finalised, so the interpreter never reports "Exception ignored" on a
        # half-closed pipe during shutdown.
        try:
            self._thread.join(timeout=5)
        except Exception:
            pass
        for stream in (self._proc.stdout, self._proc.stdin):
            try:
                if stream is not None:
                    stream.close()
            except Exception:
                pass


def _render_pdf_pages_in_worker(
    pdf_path: str,
    jobs: Sequence[Tuple[int, float, str]],
    timeout: float = _RENDER_PAGE_TIMEOUT_SEC,
    progress_cb: Optional[Callable[[int, int, int], None]] = None,
) -> List[Tuple[int, bool, int, int, str]]:
    """Render a batch of PDF pages through one persistent child process.

    Returns one ``(page_no, ok, width, height, error)`` entry per job, in
    order. A failed or timed-out attempt is retried once at half zoom; if the
    worker dies it is respawned for the retry and all remaining pages.
    """
    results: List[Tuple[int, bool, int, int, str]] = []
    total = len(jobs)
    session: Optional[_RenderWorkerSession] = None
    try:
        for idx, (page_no, zoom, out_path) in enumerate(jobs):
            ok, width, height, error = False, 0, 0, ""
            for attempt_zoom in (zoom, zoom * _RENDER_RETRY_ZOOM_RATIO):
                if session is None or not session._alive:
                    try:
                        session = _RenderWorkerSession(pdf_path)
                    except Exception as exc:
                        error = f"could not start render worker: {exc}"
                        break
                ok, width, height, error = session.render(page_no, attempt_zoom, out_path, timeout)
                if ok:
                    break
            if ok:
                results.append((page_no, True, width, height, ""))
            else:
                results.append((page_no, False, 0, 0, error or f"page {page_no} failed to render"))
            if progress_cb is not None:
                progress_cb(idx + 1, total, page_no)
    finally:
        if session is not None:
            session.close()
    return results


def _render_pdf_page_safely(
    pdf_path: str, page_no: int, zoom: float, out_path: str
) -> Tuple[bool, int, int, str]:
    """Render one PDF page in a child process so MuPDF crashes cannot kill the app."""
    results = _render_pdf_pages_in_worker(pdf_path, [(page_no, zoom, out_path)])
    ok, width, height, error = results[0][1], results[0][2], results[0][3], results[0][4]
    return ok, width, height, error


def _ai_page_bytes(path: Path, max_long_edge: int = _AI_IMAGE_LONG_EDGE_PX) -> bytes:
    """PNG bytes for AI upload, downscaled so large pages stay small in memory."""
    with Image.open(path) as img:
        img = img.convert("RGB")
        long_edge = max(img.width, img.height)
        if long_edge > max_long_edge:
            ratio = max_long_edge / float(long_edge)
            img = img.resize((max(1, int(img.width * ratio)), max(1, int(img.height * ratio))))
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        return buf.getvalue()


def index_document_pages(document_id: int) -> Tuple[int, str]:
    """Register a document's pages cheaply: PDFs list pages with extracted
    text and classification but render nothing, so huge files never touch the
    pixmap pipeline. Single-page image/office files are processed immediately
    because they are small.
    """
    docs = lquery("SELECT * FROM documents WHERE id=?", (document_id,))
    if not docs:
        return 0, "Document not found"
    doc = docs[0]
    path = Path(str(doc["path"]))
    if not path.exists():
        return 0, "File is missing from PlanReader storage"
    workspace_id = int(doc["workspace_id"])
    if path.suffix.lower() != ".pdf":
        return process_document(document_id)

    if fitz is None:
        raise RuntimeError("PyMuPDF is not installed; PDFs cannot be processed.")
    pdf = fitz.open(path)
    extracted: List[str] = []
    count = 0
    for index, page in enumerate(pdf):
        page_no = index + 1
        text = page.get_text("text") or ""
        extracted.append(text)
        page_type, label = classify_page(text, path.name, page_no)
        lexecute(
            """INSERT INTO pages(document_id,workspace_id,page_no,page_label,page_type,scale_text,px_per_m,image_path,width_px,height_px,render_zoom,extracted_text,selected,created_at)
               VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
            (document_id, workspace_id, page_no, label, page_type, "", None, "", 0, 0, _pdf_render_zoom(page), text, 1, now_stamp()),
        )
        count += 1
    pdf.close()
    gc.collect()
    lexecute(
        "UPDATE documents SET page_count=?, extracted_text=? WHERE id=?",
        (count, "\n\n".join(extracted)[:2_000_000], document_id),
    )
    return count, "Indexed"


def process_document(
    document_id: int,
    force: bool = False,
    page_ids: Optional[Sequence[int]] = None,
    progress_cb: Optional[Callable[[int, int, int], None]] = None,
) -> Tuple[int, str]:
    """Render a document's page images. When page_ids is given, only those PDF
    pages are rendered; existing page rows are updated in place.
    """
    docs = lquery("SELECT * FROM documents WHERE id=?", (document_id,))
    if not docs:
        return 0, "Document not found"
    doc = docs[0]
    if not force:
        unrendered = lquery(
            "SELECT COUNT(*) AS n FROM pages WHERE document_id=? AND (image_path='' OR image_path IS NULL)",
            (document_id,),
        )
        if unrendered and int(unrendered[0]["n"] or 0) == 0:
            return int(doc.get("page_count") or 0), "Already processed"
    path = Path(str(doc["path"]))
    if not path.exists():
        return 0, "File is missing from PlanReader storage"
    workspace_id = int(doc["workspace_id"])
    pages_dir = workspace_path(workspace_id) / "pages"
    suffix = path.suffix.lower()
    extracted_all: List[str] = []
    created = 0
    if force and page_ids is None:
        for row in lquery("SELECT image_path FROM pages WHERE document_id=?", (document_id,)):
            try:
                Path(str(row.get("image_path") or "")).unlink(missing_ok=True)
            except Exception:
                pass
        lexecute("DELETE FROM pages WHERE document_id=?", (document_id,))

    if suffix == ".pdf":
        if fitz is None:
            raise RuntimeError("PyMuPDF is not installed; PDFs cannot be processed.")
        jobs: List[Tuple[int, float, str, str, str, str]] = []
        pdf = fitz.open(path)
        for index, page in enumerate(pdf):
            page_no = index + 1
            if page_ids is not None and page_no not in page_ids:
                continue
            text = page.get_text("text") or ""
            extracted_all.append(text)
            zoom = _pdf_render_zoom(page)
            page_type, label = classify_page(text, path.name, page_no)
            image_path = pages_dir / f"doc_{document_id}_page_{page_no}.png"
            jobs.append((page_no, zoom, str(image_path), text, page_type, label))
        pdf.close()
        gc.collect()
        if not jobs:
            return 0, "No pages selected"
        results = _render_pdf_pages_in_worker(
            str(path),
            [(page_no, zoom, image_path) for page_no, zoom, image_path, _t, _ty, _l in jobs],
            _RENDER_PAGE_TIMEOUT_SEC,
            progress_cb,
        )
        result_by_page = {r[0]: r for r in results}
        failures: List[str] = []
        for page_no, zoom, image_path, text, page_type, label in jobs:
            _pno, ok, width, height, render_error = result_by_page[page_no]
            if not ok:
                failures.append(f"page {page_no}: {render_error}")
                continue
            existing = lquery("SELECT id FROM pages WHERE document_id=? AND page_no=?", (document_id, page_no))
            if existing:
                lexecute(
                    "UPDATE pages SET image_path=?,width_px=?,height_px=?,render_zoom=?,page_label=?,page_type=?,extracted_text=?,selected=1 WHERE id=?",
                    (image_path, width, height, zoom, label, page_type, text, existing[0]["id"]),
                )
            else:
                lexecute(
                    """INSERT INTO pages(document_id,workspace_id,page_no,page_label,page_type,scale_text,px_per_m,image_path,width_px,height_px,render_zoom,extracted_text,selected,created_at)
                       VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                    (document_id, workspace_id, page_no, label, page_type, "", None, image_path, width, height, zoom, text, 1, now_stamp()),
                )
            created += 1
        if failures:
            return created, f"Processed with {len(failures)} page(s) that failed to render: {'; '.join(failures[:5])}"
    elif suffix in {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff"}:
        img = Image.open(path).convert("RGB")
        image_path = pages_dir / f"doc_{document_id}_page_1.png"
        img.save(image_path)
        page_type, label = classify_page("", path.name, 1)
        lexecute(
            """INSERT INTO pages(document_id,workspace_id,page_no,page_label,page_type,scale_text,px_per_m,image_path,width_px,height_px,extracted_text,selected,created_at)
               VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?)""",
            (document_id, workspace_id, 1, label, page_type, "", None, str(image_path), img.width, img.height, "", 1, now_stamp()),
        )
        created = 1
    elif suffix == ".docx":
        if DocxDocument is None:
            raise RuntimeError("python-docx is not installed.")
        document = DocxDocument(path)
        text = "\n".join(p.text for p in document.paragraphs)
        extracted_all.append(text)
        image_path = pages_dir / f"doc_{document_id}_page_1.png"
        width, height = placeholder_image(path.name, text, image_path)
        page_type, label = classify_page(text, path.name, 1)
        lexecute(
            """INSERT INTO pages(document_id,workspace_id,page_no,page_label,page_type,scale_text,px_per_m,image_path,width_px,height_px,extracted_text,selected,created_at)
               VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?)""",
            (document_id, workspace_id, 1, label, page_type, "", None, str(image_path), width, height, text, 1, now_stamp()),
        )
        created = 1
    elif suffix in {".xlsx", ".xls", ".csv"}:
        blocks: List[str] = []
        if suffix == ".csv":
            frame = pd.read_csv(path)
            blocks.append(frame.head(300).to_csv(index=False))
        else:
            sheets = pd.read_excel(path, sheet_name=None)
            for sheet_name, frame in sheets.items():
                blocks.append(f"SHEET: {sheet_name}\n{frame.head(200).to_csv(index=False)}")
        text = "\n\n".join(blocks)
        extracted_all.append(text)
        image_path = pages_dir / f"doc_{document_id}_page_1.png"
        width, height = placeholder_image(path.name, text, image_path)
        page_type, label = classify_page(text, path.name, 1)
        lexecute(
            """INSERT INTO pages(document_id,workspace_id,page_no,page_label,page_type,scale_text,px_per_m,image_path,width_px,height_px,extracted_text,selected,created_at)
               VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?)""",
            (document_id, workspace_id, 1, label, page_type, "", None, str(image_path), width, height, text, 1, now_stamp()),
        )
        created = 1
    else:
        try:
            text = path.read_text(errors="ignore")
        except Exception:
            text = "Unsupported file format. The file remains registered as a source document."
        extracted_all.append(text)
        image_path = pages_dir / f"doc_{document_id}_page_1.png"
        width, height = placeholder_image(path.name, text, image_path)
        page_type, label = classify_page(text, path.name, 1)
        lexecute(
            """INSERT INTO pages(document_id,workspace_id,page_no,page_label,page_type,scale_text,px_per_m,image_path,width_px,height_px,extracted_text,selected,created_at)
               VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?)""",
            (document_id, workspace_id, 1, label, page_type, "", None, str(image_path), width, height, text, 1, now_stamp()),
        )
        created = 1

    counted = lquery("SELECT COUNT(*) AS n FROM pages WHERE document_id=?", (document_id,))
    page_total = int(counted[0]["n"]) if counted else created
    lexecute(
        "UPDATE documents SET page_count=?, extracted_text=? WHERE id=?",
        (page_total, "\n\n".join(extracted_all)[:2_000_000], document_id),
    )
    seed_drawing_register(workspace_id)
    return created, "Processed"


def seed_drawing_register(workspace_id: int) -> None:
    pages = lquery(
        """
        SELECT p.id,p.page_label,p.page_type,p.scale_text,p.page_no,d.file_name
        FROM pages p JOIN documents d ON d.id=p.document_id
        WHERE p.workspace_id=? ORDER BY d.id,p.page_no
        """,
        (workspace_id,),
    )
    existing_keys = {
        (str(r.get("title") or ""), str(r.get("source_reference") or ""))
        for r in lquery("SELECT title,source_reference FROM register_items WHERE workspace_id=? AND register_name='drawing_register'", (workspace_id,))
    }
    for page in pages:
        key = (str(page.get("page_label") or ""), f"{page.get('file_name')} p{page.get('page_no')}")
        if key in existing_keys:
            continue
        lexecute(
            """INSERT INTO register_items(workspace_id,register_name,item_no,title,detail,priority,source_reference,status,created_at)
               VALUES(?,?,?,?,?,?,?,?,?)""",
            (
                workspace_id,
                "drawing_register",
                str(page.get("page_label") or ""),
                str(page.get("page_label") or ""),
                str(page.get("page_type") or ""),
                str(page.get("scale_text") or ""),
                f"{page.get('file_name')} p{page.get('page_no')}",
                "Reviewed" if page.get("page_type") != "Other" else "To classify",
                now_stamp(),
            ),
        )


# -----------------------------------------------------------------------------
# Take-off and model calculations
# -----------------------------------------------------------------------------


def paint_litres(quantity: float, unit: str, coats: float, coverage: float) -> float:
    if unit != "m²" or coverage <= 0:
        return 0.0
    return max(0.0, quantity * coats / coverage)


def labour_hours(quantity: float, unit: str, productivity: float) -> float:
    if productivity <= 0:
        return 0.0
    if unit in {"m²", "lm", "No.", "item"}:
        return max(0.0, quantity / productivity)
    return 0.0


def row_value(quantity: float, rate: float) -> float:
    return max(0.0, quantity * rate)


DEFAULT_RATES_M2 = {
    "Plasterboard": 28.0,
    "Wet-area plasterboard": 30.0,
    "Fibre cement": 42.0,
    "Precast concrete": 48.0,
    "Masonry / blockwork": 42.0,
    "Timber door": 220.0,
    "Timber trim / joinery": 15.0,
    "Structural steel": 95.0,
    "Metalwork": 40.0,
    "Concrete floor": 55.0,
    "Soffit": 45.0,
    "Previously painted substrate": 32.0,
    "Other": 25.0,
}


def default_rate_for(substrate: Any, element: Any, finish_system: Any, unit: Any) -> float:
    """Return an editable default $/unit rate for a take-off row.

    The AI is instructed not to invent rates. This library applies Premier
    Brushworks' own default estimating rates so a take-off is always priced,
    with every figure remaining editable in the schedule.
    """
    sub = str(substrate or "").strip().lower()
    el = str(element or "").strip().lower()
    fin = str(finish_system or "").strip().lower()
    unit_s = str(unit or "").strip()

    if "door" in el or "door" in sub:
        if unit_s in {"No.", "item", "each"}:
            if "entry" in el or "front" in el or "feature" in el:
                return 650.0
            if "double" in el or "pair" in el:
                return 900.0
            if "fire" in el:
                return 500.0
            return 220.0
        return 220.0
    if "architrave" in el or "skirting" in el or "dado" in el or "chair rail" in el or "trim" in el or "timber trim" in sub:
        return 15.0
    if "handrail" in el or "balustrade" in el or "balcony rail" in el:
        return 115.0
    if "fascia" in el:
        return 20.0
    if "gutter" in el:
        return 18.0
    if "downpipe" in el:
        return 32.0
    if "window" in el:
        if unit_s in {"No.", "item", "each"}:
            return 280.0
        return 280.0
    if "stair" in el:
        return 2000.0
    if "interior coatings" in fin or "paint" in fin:
        return 60.0
    if "floor" in el or "balcony" in el or "deck" in el:
        if "epoxy" in fin:
            return 95.0
        if "seal" in fin or "clear" in fin:
            return 42.0
        return 60.0
    if "ceiling" in el or "soffit" in el or "ceiling flat" in fin:
        if "fibre" in sub or "soffit" in sub:
            return 38.0
        return 25.0
    if "steel" in sub or "metal" in sub:
        if unit_s in {"lm", "m"}:
            return 38.0
        return 95.0
    if "render" in sub or "brick" in sub or "block" in sub or "masonry" in sub:
        return 42.0
    if "previously painted" in sub:
        return 32.0
    if "fibre cement" in sub:
        return 42.0
    if "plasterboard" in sub:
        return 28.0
    if "concrete" in sub:
        return 48.0
    if "timber" in sub:
        if unit_s in {"lm", "m"}:
            return 15.0
        return 55.0
    for key, rate in DEFAULT_RATES_M2.items():
        if key.lower() in sub:
            return rate
    return 0.0


def polygon_area(points: Sequence[Sequence[float]]) -> float:
    if len(points) < 3:
        return 0.0
    area = 0.0
    for i in range(len(points)):
        x1, y1 = points[i]
        x2, y2 = points[(i + 1) % len(points)]
        area += x1 * y2 - x2 * y1
    return abs(area) / 2.0


def finish_to_colour(finish: Any) -> str:
    """Map a free-text finish/colour description to a plotly-friendly hex colour."""
    text = str(finish or "").lower()
    pairs = [
        ("charcoal", "#3A3A3C"), ("graphite", "#4A4A4C"), ("black", "#222222"),
        ("white", "#F2F0EB"), ("off-white", "#EDE9E0"), ("cream", "#E8E1CC"),
        ("ivory", "#F4F0E3"), ("grey", "#9A9A96"), ("gray", "#9A9A96"),
        ("silver", "#C6C6C2"), ("zinc", "#8E8E8A"), ("aluminium", "#B8B8B4"),
        ("colorbond", "#6E7B6E"), ("weathered", "#8A8578"), ("zincalume", "#A8ACAE"),
        ("timber", "#9A6A3A"), ("wood", "#9A6A3A"), ("cedar", "#A3703C"),
        ("oak", "#B58348"), ("stained", "#6E4B2C"), ("birch", "#C9A97A"),
        ("brick", "#9A4B32"), ("red brick", "#A24A33"), ("render", "#D8CFBE"),
        ("dune", "#CFC3AC"), ("beige", "#CBB79A"), ("sand", "#C9B79C"),
        ("buff", "#C9B48F"), ("clay", "#B5653C"), ("terracotta", "#B9653A"),
        ("brown", "#7A5230"), ("cocoa", "#6E4B33"), ("espresso", "#4A3424"),
        ("blue", "#4A74A4"), ("navy", "#2C4058"), ("teal", "#3E7A70"),
        ("green", "#4F7A4F"), ("sage", "#8A9378"), ("olive", "#6E7238"),
        ("red", "#A04444"), ("crimson", "#8E2E2E"), ("burgundy", "#6E2E3C"),
        ("rust", "#9A4B2E"), ("orange", "#C07830"), ("ochre", "#C09A3C"),
        ("yellow", "#C9B93C"), ("gold", "#B89A2E"), ("bronze", "#8E6E2E"),
        ("purple", "#6E4A8E"), ("pink", "#C48A96"), ("tan", "#B49A74"),
        ("dark", "#4A4A48"), ("light", "#E6E2D8"), ("prefinished", "#B8B4A8"),
        ("powdercoat", "#B0ACA2"),
    ]
    for key, color in pairs:
        if key in text:
            return color
    return "#C9BFA6"


def add_cuboid(fig: go.Figure, x: float, y: float, z: float, w: float, d: float, h: float, name: str, opacity: float = 0.65, hover: str = "", color: str = "") -> None:
    xs = [x, x+w, x+w, x, x, x+w, x+w, x]
    ys = [y, y, y+d, y+d, y, y, y+d, y+d]
    zs = [z, z, z, z, z+h, z+h, z+h, z+h]
    i = [0, 0, 0, 1, 1, 2, 4, 4, 5, 6, 3, 3]
    j = [1, 2, 3, 2, 5, 3, 5, 6, 6, 7, 7, 4]
    k = [2, 3, 4, 5, 4, 7, 6, 7, 1, 2, 4, 0]
    fig.add_trace(
        go.Mesh3d(
            x=xs,
            y=ys,
            z=zs,
            i=i,
            j=j,
            k=k,
            name=name,
            opacity=opacity,
            hovertext=hover or name,
            hoverinfo="text",
            flatshading=True,
            showscale=False,
            **({"color": color} if color else {}),
        )
    )


def build_3d_figure(workspace_id: int) -> go.Figure:
    masses = lquery("SELECT * FROM model_masses WHERE workspace_id=? ORDER BY z,id", (workspace_id,))
    zones = lquery(
        """
        SELECT z.*,p.width_px,p.height_px,p.page_label FROM mapped_zones z
        JOIN pages p ON p.id=z.page_id WHERE z.workspace_id=? ORDER BY z.id
        """,
        (workspace_id,),
    )
    fig = go.Figure()
    for mass in masses:
        add_cuboid(
            fig,
            to_float(mass.get("x")),
            to_float(mass.get("y")),
            to_float(mass.get("z")),
            max(0.05, to_float(mass.get("width"), 1)),
            max(0.05, to_float(mass.get("depth"), 1)),
            max(0.05, to_float(mass.get("height"), 2.7)),
            str(mass.get("label") or "Building mass"),
            0.72 if str(mass.get("confidence") or "").lower() in {"measured", "verified"} else 0.42,
            f"{mass.get('label')}<br>{mass.get('width')} × {mass.get('depth')} × {mass.get('height')} m<br>{mass.get('confidence')}<br>{mass.get('source_reference')}",
            finish_to_colour(mass.get("finish")),
        )
    if not masses:
        # Plan zones become conceptual room/building masses.
        for idx, zone in enumerate(zones):
            pxpm = to_float(zone.get("px_per_m"))
            if pxpm <= 0 or str(zone.get("view_type") or "").lower() not in {"floor plan", "plan", "room footprint", "building footprint"}:
                continue
            x = to_float(zone.get("x_px")) / pxpm
            y = to_float(zone.get("y_px")) / pxpm
            w = max(0.05, to_float(zone.get("w_px")) / pxpm)
            d = max(0.05, to_float(zone.get("h_px")) / pxpm)
            h = max(0.05, to_float(zone.get("wall_height_m"), 2.7))
            add_cuboid(
                fig,
                x,
                y,
                0,
                w,
                d,
                h,
                str(zone.get("name") or f"Mapped zone {idx+1}"),
                0.48,
                f"Concept mass from {zone.get('page_label')}<br>{w:.2f} × {d:.2f} × {h:.2f} m",
            )
    openings = lquery("SELECT * FROM model_openings WHERE workspace_id=? ORDER BY id", (workspace_id,))
    mass_by_id = {int(m["id"]): m for m in masses}
    for opening in openings:
        mass = mass_by_id.get(to_int(opening.get("mass_id")))
        if not mass:
            continue
        face = str(opening.get("face") or "Front").lower()
        ox = to_float(opening.get("offset_x"))
        oz = to_float(opening.get("offset_z"))
        ow = max(0.05, to_float(opening.get("width"), .9))
        oh = max(0.05, to_float(opening.get("height"), 2.1))
        x, y, z = to_float(mass.get("x")), to_float(mass.get("y")), to_float(mass.get("z"))
        w, d, h = to_float(mass.get("width")), to_float(mass.get("depth")), to_float(mass.get("height"))
        if face in {"front", "south"}:
            xs, ys, zs = [x+ox, x+ox+ow, x+ox+ow, x+ox], [y-0.01]*4, [z+oz, z+oz, z+oz+oh, z+oz+oh]
        elif face in {"back", "north"}:
            xs, ys, zs = [x+ox, x+ox+ow, x+ox+ow, x+ox], [y+d+0.01]*4, [z+oz, z+oz, z+oz+oh, z+oz+oh]
        elif face in {"left", "west"}:
            xs, ys, zs = [x-0.01]*4, [y+ox, y+ox+ow, y+ox+ow, y+ox], [z+oz, z+oz, z+oz+oh, z+oz+oh]
        else:
            xs, ys, zs = [x+w+0.01]*4, [y+ox, y+ox+ow, y+ox+ow, y+ox], [z+oz, z+oz, z+oz+oh, z+oz+oh]
        fig.add_trace(
            go.Mesh3d(
                x=xs,
                y=ys,
                z=zs,
                i=[0, 0], j=[1, 2], k=[2, 3],
                name=str(opening.get("label") or opening.get("opening_type") or "Opening"),
                opacity=.88,
                hovertext=f"{opening.get('opening_type')} · {opening.get('width')} × {opening.get('height')} m",
                hoverinfo="text",
                showscale=False,
            )
        )
    fig.update_layout(
        height=720,
        margin=dict(l=0, r=0, t=45, b=0),
        title="Interactive building model",
        scene=dict(
            xaxis_title="X (m)",
            yaxis_title="Y (m)",
            zaxis_title="Height (m)",
            aspectmode="data",
            camera=dict(eye=dict(x=1.5, y=-1.7, z=1.25)),
        ),
        legend=dict(orientation="h"),
    )
    return fig


def generate_obj(workspace_id: int) -> str:
    masses = lquery("SELECT * FROM model_masses WHERE workspace_id=? ORDER BY id", (workspace_id,))
    lines = ["# Premier Brushworks PlanReader OBJ", f"# Generated {now_stamp()}"]
    vertex_offset = 1
    for mass in masses:
        x, y, z = to_float(mass.get("x")), to_float(mass.get("y")), to_float(mass.get("z"))
        w, d, h = to_float(mass.get("width")), to_float(mass.get("depth")), to_float(mass.get("height"))
        lines.append(f"o {safe_name(mass.get('label'),'mass')}")
        vertices = [
            (x,y,z),(x+w,y,z),(x+w,y+d,z),(x,y+d,z),
            (x,y,z+h),(x+w,y,z+h),(x+w,y+d,z+h),(x,y+d,z+h),
        ]
        for vx,vy,vz in vertices:
            lines.append(f"v {vx:.6f} {vy:.6f} {vz:.6f}")
        faces = [(1,2,3,4),(5,8,7,6),(1,5,6,2),(2,6,7,3),(3,7,8,4),(5,1,4,8)]
        for face in faces:
            lines.append("f " + " ".join(str(vertex_offset + n - 1) for n in face))
        vertex_offset += 8
    return "\n".join(lines) + "\n"


# -----------------------------------------------------------------------------
# OpenAI plan reading
# -----------------------------------------------------------------------------


def resolve_openai_key(session_key: str = "") -> str:
    if session_key.strip():
        return session_key.strip()
    try:
        secret = st.secrets.get("OPENAI_API_KEY", "")
        if secret:
            return str(secret)
    except Exception:
        pass
    return os.environ.get("OPENAI_API_KEY", "")


def resolve_gemini_key(session_key: str = "") -> str:
    if session_key.strip():
        return session_key.strip()
    try:
        secret = st.secrets.get("GEMINI_API_KEY", "")
        if secret:
            return str(secret)
    except Exception:
        pass
    return os.environ.get("GEMINI_API_KEY", "")


def resolve_ai_provider(session_provider: str = "") -> str:
    if session_provider and session_provider in AI_PROVIDERS:
        return session_provider
    try:
        secret = st.secrets.get("AI_PROVIDER", "")
        if secret and secret in AI_PROVIDERS:
            return str(secret)
    except Exception:
        pass
    configured = os.environ.get("AI_PROVIDER", "")
    return configured if configured in AI_PROVIDERS else DEFAULT_AI_PROVIDER


def resolve_ai_key(provider: str, session_key: str = "") -> str:
    if provider == "Google Gemini":
        return resolve_gemini_key(session_key)
    return resolve_openai_key(session_key)


def default_ai_model(provider: str) -> str:
    if provider == "Google Gemini":
        return DEFAULT_GEMINI_MODEL
    return DEFAULT_OPENAI_MODEL


def image_data_url(path: Path) -> str:
    data = base64.b64encode(_ai_page_bytes(path)).decode("ascii")
    return f"data:image/png;base64,{data}"


def ai_schema() -> Dict[str, Any]:
    takeoff_props = {
        "section": {"type": "string"},
        "element": {"type": "string"},
        "location": {"type": "string"},
        "substrate": {"type": "string"},
        "finish_system": {"type": "string"},
        "quantity": {"type": "number"},
        "unit": {"type": "string"},
        "quantity_status": {"type": "string"},
        "source_page": {"type": "string"},
        "source_reference": {"type": "string"},
        "inclusion_status": {"type": "string"},
        "coats": {"type": "number"},
        "coverage_m2_per_litre": {"type": "number"},
        "productivity_m2_per_hour": {"type": "number"},
        "rate_per_unit": {"type": "number"},
        "confidence": {"type": "string"},
        "notes": {"type": "string"},
        "row_role": {"type": "string", "enum": ["", "floor_area"]},
    }
    register_props = {
        "register_name": {"type": "string"},
        "item_no": {"type": "string"},
        "title": {"type": "string"},
        "detail": {"type": "string"},
        "priority": {"type": "string"},
        "source_reference": {"type": "string"},
        "status": {"type": "string"},
    }
    mass_props = {
        "label": {"type": "string"},
        "level_name": {"type": "string"},
        "x": {"type": "number"},
        "y": {"type": "number"},
        "z": {"type": "number"},
        "width": {"type": "number"},
        "depth": {"type": "number"},
        "height": {"type": "number"},
        "finish": {"type": "string"},
        "source_reference": {"type": "string"},
        "confidence": {"type": "string"},
        "notes": {"type": "string"},
    }
    opening_props = {
        "mass_label": {"type": "string"},
        "label": {"type": "string"},
        "opening_type": {"type": "string"},
        "face": {"type": "string"},
        "offset_x": {"type": "number"},
        "offset_z": {"type": "number"},
        "width": {"type": "number"},
        "height": {"type": "number"},
        "count": {"type": "integer"},
        "notes": {"type": "string"},
        "source_reference": {"type": "string"},
    }
    return {
        "type": "object",
        "properties": {
            "executive_summary": {"type": "string"},
            "drawing_issue": {"type": "string"},
            "takeoff_rows": {
                "type": "array",
                "items": {"type": "object", "properties": takeoff_props, "required": list(takeoff_props), "additionalProperties": False},
            },
            "register_items": {
                "type": "array",
                "items": {"type": "object", "properties": register_props, "required": list(register_props), "additionalProperties": False},
            },
            "model_masses": {
                "type": "array",
                "items": {"type": "object", "properties": mass_props, "required": list(mass_props), "additionalProperties": False},
            },
            "model_openings": {
                "type": "array",
                "items": {"type": "object", "properties": opening_props, "required": list(opening_props), "additionalProperties": False},
            },
            "unknowns": {"type": "array", "items": {"type": "string"}},
        },
        "required": ["executive_summary", "drawing_issue", "takeoff_rows", "register_items", "model_masses", "model_openings", "unknowns"],
        "additionalProperties": False,
    }


def _openai_generate(api_key: str, model: str, prompt: str, blocks: List[Tuple[str, str]], schema: Dict[str, Any], schema_name: str) -> Dict[str, Any]:
    if OpenAI is None:
        raise RuntimeError("The openai Python package is not installed.")
    if not api_key:
        raise RuntimeError("OpenAI API key is not configured.")
    client = OpenAI(api_key=api_key)
    content: List[Dict[str, Any]] = [{"type": "input_text", "text": prompt}]
    for kind, value in blocks:
        if kind == "text":
            content.append({"type": "input_text", "text": value})
        else:
            content.append({"type": "input_image", "image_url": image_data_url(Path(value)), "detail": "high"})
    try:
        response = client.responses.create(
            model=model,
            input=[{"role": "user", "content": content}],
            text={"format": {"type": "json_schema", "name": schema_name, "strict": True, "schema": schema}},
        )
        raw = response.output_text
        data = json.loads(raw)
    except Exception as structured_error:
        fallback_prompt = prompt + "\nReturn valid JSON matching this schema exactly:\n" + json.dumps(schema)
        fallback_content = [{"type": "input_text", "text": fallback_prompt}] + content[1:]
        response = client.responses.create(model=model, input=[{"role": "user", "content": fallback_content}])
        raw = response.output_text
        match = re.search(r"\{.*\}", raw, re.S)
        if not match:
            raise RuntimeError(f"Structured output failed: {structured_error}; fallback did not return JSON.")
        data = json.loads(match.group(0))
    return data


def _extract_json_from_text(text: str) -> Dict[str, Any]:
    """Return the first complete JSON object in ``text``.

    Tolerates markdown code fences and any trailing prose the model may add,
    and scans string-aware for the balanced braces of a complete object when a
    direct parse fails. Raises ValueError if no complete object can be parsed.
    """
    cleaned = str(text or "").strip()
    cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned)
    cleaned = re.sub(r"\s*```$", "", cleaned)
    if cleaned:
        try:
            return json.loads(cleaned)
        except json.JSONDecodeError:
            pass
    start = cleaned.find("{")
    while start != -1:
        depth = 0
        in_string = False
        escaped = False
        for i in range(start, len(cleaned)):
            char = cleaned[i]
            if in_string:
                if escaped:
                    escaped = False
                elif char == "\\":
                    escaped = True
                elif char == '"':
                    in_string = False
                continue
            if char == '"':
                in_string = True
            elif char == "{":
                depth += 1
            elif char == "}":
                depth -= 1
                if depth == 0:
                    try:
                        return json.loads(cleaned[start:i + 1])
                    except json.JSONDecodeError:
                        break
        start = cleaned.find("{", start + 1)
    raise ValueError("Gemini returned JSON that could not be parsed as a complete object.")


def _retry_after_seconds(text: str, default: float = 30.0) -> float:
    """Read the provider's 'Please retry in Xs' hint from a rate-limit body."""
    match = re.search(r"retry in\s+([\d.]+)\s*s", str(text or ""), re.I)
    if match:
        try:
            return min(float(match.group(1)), 120.0)
        except ValueError:
            pass
    return default


def _gemini_generate(api_key: str, model: str, prompt: str, blocks: List[Tuple[str, str]], schema: Dict[str, Any], schema_name: str) -> Dict[str, Any]:
    if not api_key:
        raise RuntimeError("Google Gemini API key is not configured.")
    parts: List[Dict[str, Any]] = []
    if prompt:
        parts.append({"text": prompt})
    for kind, value in blocks:
        if kind == "text":
            parts.append({"text": value})
        else:
            path = Path(value)
            parts.append({"inline_data": {"mime_type": "image/png", "data": base64.b64encode(_ai_page_bytes(path)).decode("ascii")}})
    parts.append({"text": "Return a single valid JSON object matching this schema exactly, with no prose:\n" + json.dumps(schema)})
    body = {
        "contents": [{"role": "user", "parts": parts}],
        "generationConfig": {"responseMimeType": "application/json", "temperature": 0.1, "maxOutputTokens": 32768},
    }
    url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent"
    last_error = ""
    transient_codes = {500, 502, 503, 504}
    for attempt in range(4):
        try:
            resp = requests.post(url, params={"key": api_key}, json=body, timeout=300)
        except requests.exceptions.RequestException as exc:
            last_error = f"Gemini request failed: {exc}"
            time.sleep(min(2 ** attempt, 10))
            continue
        if resp.status_code == 429:
            wait = _retry_after_seconds(resp.text)
            if attempt >= 3:
                raise RuntimeError(
                    f"Gemini is rate-limited (free-tier quota for {model}). Wait about {wait:.0f}s and try again, "
                    "use a paid key, or switch providers. On the free tier gemini-3.5-flash-lite allows far more "
                    "requests per minute for draft take-offs."
                )
            last_error = f"Gemini rate-limited; retrying after {wait:.0f}s."
            time.sleep(wait)
            continue
        if resp.status_code in transient_codes:
            last_error = f"Gemini returned HTTP {resp.status_code}; the provider was temporarily unavailable."
            time.sleep(min(2 ** attempt, 10))
            continue
        resp.raise_for_status()
        payload = resp.json()
        try:
            text = payload["candidates"][0]["content"]["parts"][0]["text"]
        except (KeyError, IndexError, TypeError):
            raise RuntimeError(f"Gemini returned no usable content: {str(payload)[:1000]}")
        try:
            return _extract_json_from_text(text)
        except ValueError as exc:
            last_error = str(exc)
            continue
    raise RuntimeError(f"Gemini could not return valid JSON after retries: {last_error}")


def run_ai_structured(provider: str, api_key: str, model: str, prompt: str, blocks: List[Tuple[str, str]], schema: Dict[str, Any], schema_name: str) -> Dict[str, Any]:
    if provider == "Google Gemini":
        return _gemini_generate(api_key, model, prompt, blocks, schema, schema_name)
    return _openai_generate(api_key, model, prompt, blocks, schema, schema_name)


def _ai_error_hint(exc: Exception) -> str:
    msg = str(exc)
    low = msg.lower()
    if "api key" in low or "invalid_api_key" in low or "401" in msg or "403" in msg or "permission" in low:
        return ("The AI provider rejected the API key. In Render go to your service and set OPENAI_API_KEY "
                "(or set AI_PROVIDER to 'Google Gemini' with a GEMINI_API_KEY from aistudio.google.com for the free tier), "
                "then redeploy. You can also paste a working key in the sidebar session key box.")
    if "rate-limited" in low or "retry after" in low or "retrying after" in low:
        return msg
    if "quota" in low or "rate limit" in low or "429" in msg or "insufficient" in low:
        return "The AI provider is out of quota or rate-limited. Wait a moment, or switch providers in the sidebar."
    if "connection" in low or "timeout" in low or "max retries" in low or "dns" in low or "resolve" in low:
        return "Could not reach the AI provider from this server. Check network/outbound access and try again."
    return msg


def run_ai_plan_read(workspace_id: int, page_ids: Sequence[int], api_key: str, model: str, provider: str = "OpenAI") -> Dict[str, Any]:
    if not page_ids:
        raise RuntimeError("Select at least one page.")
    pages = lquery(
        f"SELECT p.*,d.file_name FROM pages p JOIN documents d ON d.id=p.document_id WHERE p.id IN ({','.join('?' for _ in page_ids)}) ORDER BY p.id",
        tuple(page_ids),
    )
    prompt = """
You are a senior Australian commercial painting estimator and plan take-off reviewer.
Read the supplied architectural drawing pages and extracted text using the Premier Brushworks subscription take-off method.

Required method:
1. Establish a source and drawing register.
2. Separate INCLUSIONS, EXCLUSIONS, SEPARATE ITEMS and PROVISIONAL items.
3. Create clarification/RFI items wherever scope, finish, scale, issue, access, opening deductions or dimensions are uncertain.
4. Create measurable painting take-off rows for internal walls, ceilings, doors, frames, joinery, external walls/cladding, soffits, steel, concrete coatings and specialist coatings only when supported.
5. For every floor plan level, read and report the internal floor area (m²) as its own measurement row: section='Internal', element='Floor area', unit='m²', row_role='floor_area', quantity=the internal floor area read from that level. Read internal wall areas as normal wall rows with row_role=''. If a floor area is not dimensioned on the plan, set quantity=0 and quantity_status='To measure'. Do not add a rate to floor-area rows.
6. Never invent dimensions. If a quantity cannot be measured from the supplied evidence, use quantity=0 and quantity_status='To measure'.
7. For model geometry, only create rectangular building masses where width/depth/height are supported. Use confidence='Measured' only for clear dimensions, 'Derived' for calculated dimensions, and 'Assumed' for placeholders.
8. Leave rate_per_unit at zero; the PlanReader default estimating rate library is applied automatically for every row on import. Use practical default coats, coverage and productivity only as editable estimating defaults, and explain them in notes.
9. Surface areas must be net or clearly marked gross/provisional. Identify exclusions such as glazing, tiles, prefinished metal, signage, roofing and specialist systems.

Return structured data only. References must name the drawing/page or visible note that supports each item.
"""
    blocks: List[Tuple[str, str]] = []
    for page in pages:
        text_excerpt = str(page.get("extracted_text") or "")[:12000]
        blocks.append(("text", f"SOURCE PAGE: {page.get('file_name')} · {page.get('page_label')} · page {page.get('page_no')} · classified {page.get('page_type')}\nEXTRACTED TEXT:\n{text_excerpt}"))
        image_path = Path(str(page.get("image_path") or ""))
        if image_path.exists():
            blocks.append(("image", str(image_path)))
    schema = ai_schema()
    data = run_ai_structured(provider, api_key, model, prompt, blocks, schema, "paint_takeoff_analysis")
    lexecute(
        "INSERT INTO ai_runs(workspace_id,run_type,model,source_pages,status,response_json,error_message,created_at) VALUES(?,?,?,?,?,?,?,?)",
        (workspace_id, "Plan take-off and 3D", f"{provider} · {model}", json.dumps(list(page_ids)), "Completed", json.dumps(data), "", now_stamp()),
    )
    return data


def render_ai_schema() -> Dict[str, Any]:
    facade_props = {
        "face": {"type": "string"},
        "main_material": {"type": "string"},
        "colour_description": {"type": "string"},
        "colour_hex": {"type": "string"},
        "window_count": {"type": "integer"},
        "door_count": {"type": "integer"},
        "notes": {"type": "string"},
    }
    mass_props = {
        "label": {"type": "string"},
        "width": {"type": "number"},
        "depth": {"type": "number"},
        "height": {"type": "number"},
        "finish": {"type": "string"},
        "confidence": {"type": "string"},
        "source_reference": {"type": "string"},
        "notes": {"type": "string"},
    }
    opening_props = {
        "mass_label": {"type": "string"},
        "label": {"type": "string"},
        "opening_type": {"type": "string"},
        "face": {"type": "string"},
        "offset_x": {"type": "number"},
        "offset_z": {"type": "number"},
        "width": {"type": "number"},
        "height": {"type": "number"},
        "count": {"type": "integer"},
        "notes": {"type": "string"},
    }
    return {
        "type": "object",
        "properties": {
            "executive_summary": {"type": "string"},
            "building_form": {
                "type": "object",
                "properties": {
                    "storeys": {"type": "integer"},
                    "roof_style": {"type": "string"},
                    "roof_material": {"type": "string"},
                    "roof_colour": {"type": "string"},
                    "overall_height_m": {"type": "number"},
                    "notes": {"type": "string"},
                },
                "required": ["storeys", "roof_style", "roof_material", "roof_colour", "overall_height_m", "notes"],
                "additionalProperties": False,
            },
            "facades": {
                "type": "array",
                "items": {"type": "object", "properties": facade_props, "required": list(facade_props), "additionalProperties": False},
            },
            "model_masses": {
                "type": "array",
                "items": {"type": "object", "properties": mass_props, "required": list(mass_props), "additionalProperties": False},
            },
            "model_openings": {
                "type": "array",
                "items": {"type": "object", "properties": opening_props, "required": list(opening_props), "additionalProperties": False},
            },
            "unknowns": {"type": "array", "items": {"type": "string"}},
        },
        "required": ["executive_summary", "building_form", "facades", "model_masses", "model_openings", "unknowns"],
        "additionalProperties": False,
    }


def run_ai_render_read(workspace_id: int, page_ids: Sequence[int], api_key: str, model: str, provider: str = "OpenAI") -> Dict[str, Any]:
    """Read render / artist's impression images and extract building form for the 3D model.

    Uses the render as a secondary evidence source: form, storeys, roof, per-facade
    material/colour and window/door placement. Dimensions from a render are treated as
    'Assumed' and always remain editable — the render never overrides measured plan data.
    """
    if not page_ids:
        raise RuntimeError("Select at least one render / artist's impression page.")
    pages = lquery(
        f"SELECT p.*,d.file_name FROM pages p JOIN documents d ON d.id=p.document_id WHERE p.id IN ({','.join('?' for _ in page_ids)}) ORDER BY p.id",
        tuple(page_ids),
    )
    prompt = """
You are a senior Australian architectural estimator reviewing render / artist's impression images
to inform a conceptual 3D building model for a painting take-off.

The supplied images are architectural renders or artist's impressions, NOT dimensioned drawings.
Use them only as secondary evidence for building FORM and APPEARANCE.

Required method:
1. Describe the building form: number of storeys, roof style, roof material and roof colour.
2. For each visible facade (front/south, back/north, left/west, right/east), describe the main
   material (e.g. render, weatherboard, brick, fibre cement, metal cladding, timber), the
   colour/colour description, and estimate the number of windows and doors on that face.
3. Create a single rectangular building mass (or clearly separate masses for distinct wings)
   using ASSUMED dimensions only. Prefer a height derived from the storey count using an
   approximate 3.0 m per storey (2.7 m floor-to-ceiling plus floor/roof allowance) unless the
   render shows a clear scale reference. Width/depth should reflect the render's visible
   proportions. Use confidence='Assumed' and explain the basis in notes.
4. Create model openings (windows, doors, garage doors) on the appropriate faces to match the
   render. Use ASSUMED standard sizes unless a scale reference is visible: windows 1.8 m wide
   x 1.2 m high at 0.9 m sill, doors 0.9 m x 2.1 m, garage door 4.8 m x 2.2 m. Offset each
   opening along the face so the pattern resembles the render (centres/columns evenly spaced).
5. Set model_masses finish to the dominant external finish + colour (e.g. 'White render',
   'Charcoal Colorbond cladding') so the 3D view is colour-coded.
6. Never invent measured dimensions. Everything from a render is 'Assumed' until verified against
   a dimensioned plan. List any uncertainties.

Return structured data only. References must name the render page that supports each item.
"""
    blocks: List[Tuple[str, str]] = []
    for page in pages:
        blocks.append(("text", f"RENDER PAGE: {page.get('file_name')} · {page.get('page_label')} · classified {page.get('page_type')}"))
        image_path = Path(str(page.get("image_path") or ""))
        if image_path.exists():
            blocks.append(("image", str(image_path)))
    schema = render_ai_schema()
    data = run_ai_structured(provider, api_key, model, prompt, blocks, schema, "render_analysis")
    lexecute(
        "INSERT INTO ai_runs(workspace_id,run_type,model,source_pages,status,response_json,error_message,created_at) VALUES(?,?,?,?,?,?,?,?)",
        (workspace_id, "Render / artist's impression", f"{provider} · {model}", json.dumps(list(page_ids)), "Completed", json.dumps(data), "", now_stamp()),
    )
    return data


def apply_render_to_model(workspace_id: int, data: Dict[str, Any], mode: str = "merge") -> Dict[str, int]:
    """Apply a render analysis to the 3D model.

    mode='merge' keeps any existing measured masses and only adds missing pieces.
    mode='replace' replaces assumed masses (never measured ones) with the render-derived form.
    Never overrides a mass whose confidence is 'Measured' or 'Verified'.
    """
    counts = {"masses": 0, "openings": 0}
    building_form = data.get("building_form") or {}
    exec_summary = str(data.get("executive_summary") or "")
    if exec_summary:
        lexecute("UPDATE workspaces SET executive_summary=?, updated_at=? WHERE id=?", (exec_summary, now_stamp(), workspace_id))
    existing = lquery("SELECT * FROM model_masses WHERE workspace_id=? ORDER BY id", (workspace_id,))
    measured_existing = {int(m["id"]): m for m in existing if str(m.get("confidence") or "").lower() in {"measured", "verified"}}

    if mode == "replace":
        for m in existing:
            if str(m.get("confidence") or "").lower() not in {"measured", "verified"}:
                lexecute("DELETE FROM model_openings WHERE mass_id=?", (m["id"],))
                lexecute("DELETE FROM model_masses WHERE id=?", (m["id"],))
        existing = lquery("SELECT * FROM model_masses WHERE workspace_id=? ORDER BY id", (workspace_id,))
        measured_existing = {int(m["id"]): m for m in existing if str(m.get("confidence") or "").lower() in {"measured", "verified"}}

    mass_id_by_label: Dict[str, int] = {}
    for row in data.get("model_masses", []):
        label = str(row.get("label") or "Building mass").strip() or "Building mass"
        finish = str(row.get("finish") or "").strip()
        confidence = str(row.get("confidence") or "Assumed").capitalize()
        if confidence.lower() not in {"assumed", "derived"}:
            confidence = "Assumed"
        match = next((m for m in existing if str(m.get("label") or "").strip() == label), None)
        if match and str(match.get("confidence") or "").lower() in {"measured", "verified"}:
            mass_id_by_label[label] = int(match["id"])
            if finish:
                lexecute("UPDATE model_masses SET finish=? WHERE id=?", (finish, match["id"]))
            continue
        if match:
            lexecute(
                "UPDATE model_masses SET width=?,depth=?,height=?,finish=?,confidence=?,source_reference=?,notes=? WHERE id=?",
                (to_float(row.get("width"), 1), to_float(row.get("depth"), 1), to_float(row.get("height"), 2.7),
                 finish, confidence, str(row.get("source_reference") or "Render / artist's impression"),
                 str(row.get("notes") or ""), match["id"]),
            )
            mass_id_by_label[label] = int(match["id"])
        else:
            new_id = lexecute(
                """INSERT INTO model_masses(workspace_id,label,level_name,x,y,z,width,depth,height,finish,source_reference,confidence,notes,created_at)
                   VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                (workspace_id, label, "Ground", 0, 0, 0, to_float(row.get("width"), 1), to_float(row.get("depth"), 1),
                 to_float(row.get("height"), 2.7), finish, str(row.get("source_reference") or "Render / artist's impression"),
                 confidence, str(row.get("notes") or ""), now_stamp()),
            )
            mass_id_by_label[label] = int(new_id)
            counts["masses"] += 1

    for opening in data.get("model_openings", []):
        label = str(opening.get("mass_label") or "").strip()
        mass_id = mass_id_by_label.get(label) or (next(iter(measured_existing.keys())) if measured_existing else None)
        if mass_id is None:
            continue
        lexecute(
            """INSERT INTO model_openings(workspace_id,mass_id,label,opening_type,face,offset_x,offset_z,width,height,count,notes,source_reference,created_at)
               VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?)""",
            (workspace_id, mass_id, str(opening.get("label") or opening.get("opening_type") or "Opening"),
             str(opening.get("opening_type") or "Window"), str(opening.get("face") or "Front"),
             to_float(opening.get("offset_x")), to_float(opening.get("offset_z")),
             max(0.05, to_float(opening.get("width"), 0.9)), max(0.05, to_float(opening.get("height"), 2.1)),
             max(1, to_int(opening.get("count"), 1)), str(opening.get("notes") or ""),
             f"Render / artist's impression · {str(opening.get('notes') or '').strip() or 'assumed from render'}", now_stamp()),
        )
        counts["openings"] += 1
    return counts


def import_ai_result(workspace_id: int, data: Dict[str, Any]) -> Dict[str, int]:
    counts = {"takeoff": 0, "registers": 0, "masses": 0, "openings": 0}
    if data.get("executive_summary"):
        lexecute("UPDATE workspaces SET executive_summary=?,drawing_issue=?,updated_at=? WHERE id=?", (str(data.get("executive_summary")), str(data.get("drawing_issue") or ""), now_stamp(), workspace_id))
    for row in data.get("takeoff_rows", []):
        row = dict(row)
        if not to_float(row.get("rate_per_unit")):
            row["rate_per_unit"] = default_rate_for(
                row.get("substrate"), row.get("element"), row.get("finish_system"), row.get("unit")
            )
        row_role = str(row.get("row_role") or "").strip()
        if row_role not in {"", "floor_area"}:
            row_role = ""
        row["row_role"] = row_role
        values = [row.get(col, "") for col in TAKEOFF_COLUMNS]
        lexecute(
            """
            INSERT INTO takeoff_rows(workspace_id,section,element,location,substrate,finish_system,quantity,unit,quantity_status,source_page,source_reference,inclusion_status,coats,coverage_m2_per_litre,productivity_m2_per_hour,rate_per_unit,confidence,notes,row_role,created_at,updated_at)
            VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
            """,
            (workspace_id, *values, row_role, now_stamp(), now_stamp()),
        )
        counts["takeoff"] += 1
    for row in data.get("register_items", []):
        register_name = str(row.get("register_name") or "clarifications").lower().replace(" ", "_")
        if register_name not in REGISTER_NAMES:
            register_name = "clarifications"
        lexecute(
            """INSERT INTO register_items(workspace_id,register_name,item_no,title,detail,priority,source_reference,status,created_at)
               VALUES(?,?,?,?,?,?,?,?,?)""",
            (
                workspace_id,
                register_name,
                row.get("item_no", ""),
                row.get("title", ""),
                row.get("detail", ""),
                row.get("priority", ""),
                row.get("source_reference", ""),
                row.get("status", "Open"),
                now_stamp(),
            ),
        )
        counts["registers"] += 1
    mass_label_to_id: Dict[str, int] = {}
    for row in data.get("model_masses", []):
        mass_id = lexecute(
            """INSERT INTO model_masses(workspace_id,label,level_name,x,y,z,width,depth,height,finish,source_reference,confidence,notes,created_at)
               VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
            (
                workspace_id,
                row.get("label", "Building mass"),
                row.get("level_name", "Ground"),
                row.get("x", 0), row.get("y", 0), row.get("z", 0),
                row.get("width", 1), row.get("depth", 1), row.get("height", 2.7),
                row.get("finish", ""), row.get("source_reference", ""), row.get("confidence", "Assumed"), row.get("notes", ""), now_stamp(),
            ),
        )
        mass_label_to_id[str(row.get("label", "")).lower()] = mass_id
        counts["masses"] += 1
    for row in data.get("model_openings", []):
        mass_id = mass_label_to_id.get(str(row.get("mass_label") or "").lower())
        lexecute(
            """INSERT INTO model_openings(workspace_id,mass_id,label,opening_type,face,offset_x,offset_z,width,height,count,notes,source_reference,created_at)
               VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?)""",
            (
                workspace_id, mass_id, row.get("label", "Opening"), row.get("opening_type", "Door"), row.get("face", "Front"),
                row.get("offset_x", 0), row.get("offset_z", 0), row.get("width", .9), row.get("height", 2.1), row.get("count", 1),
                row.get("notes", ""), row.get("source_reference", ""), now_stamp(),
            ),
        )
        counts["openings"] += 1
    for unknown in data.get("unknowns", []):
        lexecute(
            """INSERT INTO register_items(workspace_id,register_name,item_no,title,detail,priority,source_reference,status,created_at)
               VALUES(?,?,?,?,?,?,?,?,?)""",
            (workspace_id, "clarifications", "", "AI review unknown", str(unknown), "High", "AI plan review", "Open", now_stamp()),
        )
        counts["registers"] += 1
    return counts


# -----------------------------------------------------------------------------
# Export and JobHub push
# -----------------------------------------------------------------------------


# -----------------------------------------------------------------------------
# Workspace settings, levels and per-level pricing
# -----------------------------------------------------------------------------

_LEVEL_ORDER = ["Ground", "Level 1", "Level 2", "Level 3", "Level 4", "Level 5", "Roof"]


def workspace_setting(workspace_id: int, key: str, default: Any = None) -> Any:
    rows = lquery("SELECT value FROM workspace_settings WHERE workspace_id=? AND key=?", (workspace_id, key))
    return rows[0]["value"] if rows else default


def set_workspace_setting(workspace_id: int, key: str, value: Any) -> None:
    lexecute(
        """INSERT INTO workspace_settings(workspace_id,key,value,updated_at) VALUES(?,?,?,?)
           ON CONFLICT(workspace_id,key) DO UPDATE SET value=excluded.value, updated_at=excluded.updated_at""",
        (workspace_id, key, str(value if value is not None else ""), now_stamp()),
    )


def level_of(location: Any) -> str:
    text = str(location or "")
    lower = text.lower()
    if "ground" in lower:
        return "Ground"
    for level in ["Level 5", "Level 4", "Level 3", "Level 2", "Level 1"]:
        if level.lower() in lower:
            return level
    if "roof" in lower:
        return "Roof"
    return "Unassigned"


def level_sort_key(level: str) -> int:
    return _LEVEL_ORDER.index(level) if level in _LEVEL_ORDER else len(_LEVEL_ORDER)


def _level_rewrite(location: str, target_level: str) -> str:
    """Rewrite a row's location to a new level, keeping the rest of the text."""
    text = str(location or "").strip()
    lower = text.lower()
    for existing in ["Ground", "Level 5", "Level 4", "Level 3", "Level 2", "Level 1", "Roof"]:
        if existing.lower() in lower:
            return text.replace(existing, target_level, 1)
    return f"{target_level} · {text}" if text else str(target_level)


def copy_takeoff_rows_to_level(workspace_id: int, row_ids: Sequence[int], target_level: str) -> int:
    """Duplicate the selected take-off rows for another level (locations rewritten)."""
    created = 0
    target_level = str(target_level or "").strip()
    if not target_level:
        return 0
    for rid in row_ids:
        rows = lquery("SELECT * FROM takeoff_rows WHERE id=?", (int(rid),))
        if not rows:
            continue
        r = rows[0]
        if int(r.get("workspace_id")) != int(workspace_id):
            continue
        new_loc = _level_rewrite(str(r.get("location") or ""), target_level)
        dup = lquery(
            "SELECT id FROM takeoff_rows WHERE workspace_id=? AND section=? AND element=? AND location=? AND substrate=? AND unit=?",
            (workspace_id, str(r.get("section") or ""), str(r.get("element") or ""), new_loc,
             str(r.get("substrate") or ""), str(r.get("unit") or "")),
        )
        if dup:
            continue
        lexecute(
            """INSERT INTO takeoff_rows(workspace_id,section,element,location,substrate,finish_system,quantity,unit,quantity_status,source_page,source_reference,inclusion_status,coats,coverage_m2_per_litre,productivity_m2_per_hour,rate_per_unit,confidence,notes,row_role,created_at,updated_at)
               VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
            (workspace_id, r.get("section", ""), r.get("element", ""), new_loc, r.get("substrate", ""),
             r.get("finish_system", ""), 0, r.get("unit", ""), "To measure", r.get("source_page", ""),
             r.get("source_reference", ""), r.get("inclusion_status", ""), r.get("coats", 2),
             r.get("coverage_m2_per_litre", 0), r.get("productivity_m2_per_hour", 0), r.get("rate_per_unit", 0),
             "To review", f"Copied to {target_level}.", r.get("row_role", ""), now_stamp(), now_stamp()),
        )
        created += 1
    return created


def _quote_settings(workspace_id: int) -> Dict[str, Any]:
    def f(key: str, default: float) -> float:
        try:
            return float(workspace_setting(workspace_id, key, default))
        except (TypeError, ValueError):
            return float(default)

    return {
        "gst_rate_pct": f("gst_rate_pct", 10.0),
        "pricing_margin_pct": f("pricing_margin_pct", 0.0),
        "default_coverage_m2_per_litre": f("default_coverage_m2_per_litre", 12.0),
        "default_productivity_m2_per_hour": f("default_productivity_m2_per_hour", 8.0),
        "default_wall_height_m": f("default_wall_height_m", 2.7),
        "internal_pricing_basis": str(workspace_setting(workspace_id, "internal_pricing_basis", "wall_m2") or "wall_m2").strip().lower(),
        "quote_header": str(workspace_setting(workspace_id, "quote_header", "Premier Brushworks Pty Ltd")),
        "quote_footer": str(workspace_setting(workspace_id, "quote_footer", "Valid for 30 days. All quantities require estimator review against the current issued drawings and specification before pricing or construction use.")),
    }


def pricing_basis_label(basis: Any) -> str:
    """Human label for the project-level internal pricing basis setting."""
    return "Floor m²" if str(basis or "").strip().lower() == "floor_m2" else "Wall m²"


def takeoff_work_rows(takeoff: pd.DataFrame) -> pd.DataFrame:
    """Take-off rows that represent priced work (excludes floor-area measurement rows).

    Floor-area rows are per-level measurement rows used to drive floor-m² pricing;
    they carry real m² but are never painted quantities, so they are filtered out
    of painted-area totals and JobHub line pushes.
    """
    if takeoff.empty or "row_role" not in takeoff.columns:
        return takeoff
    return takeoff.loc[takeoff["row_role"].fillna("").ne("floor_area")]


def is_internal_wall_row(section: Any, element: Any) -> bool:
    """True when a row is an internal wall m² row that can be priced off floor area."""
    sec = str(section or "").strip().lower()
    elm = str(element or "").strip().lower()
    if any(token in sec for token in ["external", "exterior", "facade", "elevation", "soffit", "cladding", "roof", "canopy", "balcony", "deck"]):
        return False
    return "wall" in elm


def floor_area_by_level(takeoff: pd.DataFrame) -> Dict[str, float]:
    """Per-level internal floor area (m²) from rows tagged ``floor_area``."""
    out: Dict[str, float] = {}
    if takeoff.empty or "row_role" not in takeoff.columns:
        return out
    floor = takeoff.loc[
        takeoff["row_role"].fillna("").eq("floor_area") & takeoff["unit"].astype(str).eq("m²")
    ]
    for r in floor.itertuples(index=False):
        lvl = level_of(r.location)
        out[lvl] = out.get(lvl, 0.0) + to_float(r.quantity)
    return out


def per_level_summary(workspace_id: int) -> pd.DataFrame:
    """Priced quantities grouped by building level for the quote summary.

    Floor-area measurement rows (``row_role='floor_area'``) are reported in a
    dedicated ``floor_m2`` column and excluded from the painted m²/value totals.
    """
    takeoff = dataframe_for_takeoff(workspace_id)
    if takeoff.empty:
        return pd.DataFrame(columns=["level", "rows", "m2", "floor_m2", "lm", "count", "paint_litres", "labour_hours", "value_ex_gst"])
    takeoff = takeoff.copy()
    takeoff["level"] = [level_of(loc) for loc in takeoff["location"]]
    work = takeoff_work_rows(takeoff)
    out = []
    for level, group in takeoff.groupby("level", dropna=False):
        work_group = work.loc[work["level"].eq(level)]
        floor_group = takeoff.loc[
            takeoff["level"].eq(level)
            & takeoff["row_role"].fillna("").eq("floor_area")
            & takeoff["unit"].eq("m²")
        ]
        out.append({
            "level": level,
            "rows": int(len(work_group)),
            "m2": float(work_group.loc[work_group["unit"].eq("m²"), "quantity"].sum()),
            "floor_m2": float(floor_group["quantity"].sum()),
            "lm": float(work_group.loc[work_group["unit"].eq("lm"), "quantity"].sum()),
            "count": float(work_group.loc[work_group["unit"].isin({"No.", "item"}), "quantity"].sum()),
            "paint_litres": float(work_group["paint_litres"].sum()),
            "labour_hours": float(work_group["labour_hours"].sum()),
            "value_ex_gst": float(work_group["value_ex_gst"].sum()),
        })
    df = pd.DataFrame(out)
    df["sort"] = df["level"].map(level_sort_key)
    df = df.sort_values("sort", ignore_index=True).drop(columns=["sort"])
    return df


def quote_summary_frame(workspace_id: int) -> pd.DataFrame:
    """Per-level quote table including mark-up, GST and total."""
    levels = per_level_summary(workspace_id)
    if levels.empty:
        return levels
    settings = _quote_settings(workspace_id)
    margin = settings["pricing_margin_pct"] / 100.0
    gst = settings["gst_rate_pct"] / 100.0
    levels = levels.copy()
    levels["markup_ex_gst"] = levels["value_ex_gst"] * (1 + margin)
    levels["gst"] = levels["markup_ex_gst"] * gst
    levels["total_inc_gst"] = levels["markup_ex_gst"] + levels["gst"]
    return levels


def per_level_summary_csv(workspace_id: int) -> str:
    frame = quote_summary_frame(workspace_id)
    return frame.to_csv(index=False) if not frame.empty else ""


def quote_workbook_bytes(workspace_id: int) -> bytes:
    """Excel quotation with per-level pricing, mark-up, GST and full detail."""
    workspace = lquery("SELECT * FROM workspaces WHERE id=?", (workspace_id,))[0]
    settings = _quote_settings(workspace_id)
    levels = quote_summary_frame(workspace_id)
    takeoff = dataframe_for_takeoff(workspace_id)
    margin = settings["pricing_margin_pct"] / 100.0
    gst = settings["gst_rate_pct"] / 100.0
    header = pd.DataFrame([
        ["Quotation", str(settings["quote_header"])],
        ["Job number", workspace.get("job_no", "")],
        ["Project", workspace.get("job_name", "")],
        ["Builder / client", workspace.get("builder_client", "")],
        ["Site address", workspace.get("site_address", "")],
        ["Drawing issue", workspace.get("drawing_issue", "")],
        ["Estimator", workspace.get("estimator", "")],
        ["Prepared", now_stamp()],
        ["Mark-up (ex GST)", f"{settings['pricing_margin_pct']:.1f}%"],
        ["GST", f"{settings['gst_rate_pct']:.1f}%"],
    ], columns=["Field", "Value"])
    totals = pd.DataFrame({
        "Metric": ["Value ex GST", "Mark-up", "Subtotal ex GST", "GST", "Total inc GST"],
        "Amount": [
            float(takeoff["value_ex_gst"].sum()) if not takeoff.empty else 0.0,
            float(takeoff["value_ex_gst"].sum()) * margin if not takeoff.empty else 0.0,
            float(takeoff["value_ex_gst"].sum()) * (1 + margin) if not takeoff.empty else 0.0,
            float(takeoff["value_ex_gst"].sum()) * (1 + margin) * gst if not takeoff.empty else 0.0,
            float(takeoff["value_ex_gst"].sum()) * (1 + margin) * (1 + gst) if not takeoff.empty else 0.0,
        ],
    })
    sheets: List[Tuple[str, pd.DataFrame]] = [
        ("Quote Header", header),
        ("Per-Level Summary", levels),
        ("Totals", totals),
        ("Take-off Detail", takeoff),
        ("Door Schedule", register_df(workspace_id, "door_schedule")),
        ("Inclusions", register_df(workspace_id, "inclusions")),
        ("Exclusions", register_df(workspace_id, "exclusions")),
        ("Separate Clarifications", register_df(workspace_id, "clarifications")),
        ("Assumptions", register_df(workspace_id, "assumptions")),
        ("RFIs", register_df(workspace_id, "rfis")),
    ]
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        for sheet_name, frame in sheets:
            frame.to_excel(writer, index=False, sheet_name=sheet_name[:31])
            ws = writer.book[sheet_name[:31]]
            ws.freeze_panes = "A2"
            ws.auto_filter.ref = ws.dimensions
            for cell in ws[1]:
                from openpyxl.styles import Alignment, Font, PatternFill
                cell.fill = PatternFill("solid", fgColor="171717")
                cell.font = Font(color="FFFFFF", bold=True)
                cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            for col in ws.columns:
                letter = col[0].column_letter
                max_len = min(65, max(10, max(len(str(c.value or "")) for c in col) + 2))
                ws.column_dimensions[letter].width = max_len
                for cell in col:
                    cell.alignment = Alignment(vertical="top", wrap_text=True)
    return output.getvalue()


def quote_pdf_bytes(workspace_id: int) -> bytes:
    """Portable PDF quotation: job header, per-level pricing and totals."""
    workspace = lquery("SELECT * FROM workspaces WHERE id=?", (workspace_id,))[0]
    settings = _quote_settings(workspace_id)
    levels = quote_summary_frame(workspace_id)
    takeoff = dataframe_for_takeoff(workspace_id)
    margin = settings["pricing_margin_pct"] / 100.0
    gst = settings["gst_rate_pct"] / 100.0
    sub_total = float(takeoff["value_ex_gst"].sum()) if not takeoff.empty else 0.0
    markup = sub_total * margin
    subtotal_ex = sub_total + markup
    gst_amount = subtotal_ex * gst
    grand = subtotal_ex + gst_amount

    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import HRFlowable, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=18 * mm, rightMargin=18 * mm, topMargin=16 * mm, bottomMargin=16 * mm)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("PTitle", parent=styles["Title"], fontSize=18, spaceAfter=2)
    small = ParagraphStyle("PSmall", parent=styles["Normal"], fontSize=9, textColor=colors.HexColor("#5f574f"))
    h2 = ParagraphStyle("PH2", parent=styles["Heading2"], fontSize=13, spaceBefore=10, spaceAfter=4)

    flow: List[Any] = []
    flow.append(Paragraph(str(settings["quote_header"]), title_style))
    flow.append(Paragraph(f"Quotation for {workspace.get('job_name', '')}", styles["Normal"]))
    flow.append(Spacer(1, 4))
    info = Table([
        ["Job number", str(workspace.get("job_no") or ""), "Builder / client", str(workspace.get("builder_client") or "")],
        ["Site address", str(workspace.get("site_address") or ""), "Drawing issue", str(workspace.get("drawing_issue") or "")],
        ["Estimator", str(workspace.get("estimator") or ""), "Prepared", now_stamp()],
    ], colWidths=[28 * mm, 70 * mm, 32 * mm, 70 * mm])
    info.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8.5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
        ("TOPPADDING", (0, 0), (-1, -1), 2),
        ("GRID", (0, 0), (-1, -1), 0.2, colors.HexColor("#d8d2c8")),
    ]))
    flow.append(info)
    flow.append(Spacer(1, 8))

    flow.append(Paragraph("Per-level pricing (ex GST)", h2))
    flow.append(Paragraph(
        f"Internal pricing basis: <b>{pricing_basis_label(settings['internal_pricing_basis'])}</b>"
        + (" — internal wall rates apply to each level's floor area (floor m²), with measured wall m² shown for reference." if settings["internal_pricing_basis"] == "floor_m2" else " — internal rates apply to each measured wall area."),
        small,
    ))
    if levels.empty:
        flow.append(Paragraph("No take-off rows yet.", small))
    else:
        head = [["Level", "Rows", "m²", "Floor m²", "lm", "Count", "Paint (L)", "Labour (hrs)", "Value ex GST", "Mark-up", "GST", "Total inc GST"]]
        body = head + [
            [str(r["level"]), int(r["rows"]), f"{r['m2']:,.1f}", f"{r['floor_m2']:,.1f}", f"{r['lm']:,.1f}", f"{r['count']:,.1f}",
             f"{r['paint_litres']:,.1f}", f"{r['labour_hours']:,.1f}", f"${r['value_ex_gst']:,.2f}",
             f"${r['markup_ex_gst'] - r['value_ex_gst']:,.2f}", f"${r['gst']:,.2f}", f"${r['total_inc_gst']:,.2f}"]
            for r in levels.to_dict("records")
        ]
        body.append(["TOTAL", int(levels["rows"].sum()), f"{levels['m2'].sum():,.1f}", f"{levels['floor_m2'].sum():,.1f}", f"{levels['lm'].sum():,.1f}",
                     f"{levels['count'].sum():,.1f}", f"{levels['paint_litres'].sum():,.1f}",
                     f"{levels['labour_hours'].sum():,.1f}", f"${sub_total:,.2f}", f"${markup:,.2f}",
                     f"${gst_amount:,.2f}", f"${grand:,.2f}"])
        table = Table(body, repeatRows=1, colWidths=[20 * mm, 12 * mm, 15 * mm, 15 * mm, 12 * mm, 13 * mm, 16 * mm, 19 * mm, 23 * mm, 17 * mm, 14 * mm, 23 * mm])
        table.setStyle(TableStyle([
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 7.2),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#171717")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#d8d2c8")),
            ("FONTNAME", (0, -1), (-1, -1), "Helvetica-Bold"),
            ("BACKGROUND", (0, -1), (-1, -1), colors.HexColor("#f3efe7")),
            ("TOPPADDING", (0, 0), (-1, -1), 2),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
        ]))
        flow.append(table)
    flow.append(Spacer(1, 6))
    flow.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#8B6F47")))
    flow.append(Paragraph(f"<b>Mark-up (ex GST):</b> {settings['pricing_margin_pct']:.1f}% · <b>GST:</b> {settings['gst_rate_pct']:.1f}% · <b>Total inc GST:</b> ${grand:,.2f}", h2))
    flow.append(Paragraph(str(settings["quote_footer"]), small))
    doc.build(flow)
    return buf.getvalue()


def reconcile_ai_vs_drawn(workspace_id: int) -> pd.DataFrame:
    """Compare AI-drafted quantities against drawn/mapped quantities per line item.

    Every take-off row is classified by its basis (AI draft, Drawn on plan,
    JobHub import, or Manual). Rows are grouped by section/element/location/unit
    and the AI quantity is compared with the drawn quantity so differences stand
    out before the take-off is priced.
    """
    takeoff = dataframe_for_takeoff(workspace_id)
    if takeoff.empty:
        return pd.DataFrame(columns=["section", "element", "location", "unit", "ai_qty", "drawn_qty", "variance", "status"])
    records = []
    for r in takeoff.itertuples(index=False):
        rid = int(r.id)
        drawn = lquery("SELECT COUNT(*) AS c FROM measurement_lines WHERE takeoff_row_id=?", (rid,))[0]["c"]
        source_ref = str(r.source_reference or "")
        notes = str(r.notes or "")
        is_ai = "AI" in source_ref or "AI plan review" in notes or "AI-generated" in notes or "AI draft" in notes
        if drawn:
            basis = "Drawn"
        elif is_ai:
            basis = "AI"
        elif str(r.source_page or "") == "JobHub import":
            basis = "JobHub"
        else:
            basis = "Manual"
        records.append({
            "id": rid, "section": str(r.section or ""), "element": str(r.element or ""),
            "location": str(r.location or ""), "unit": str(r.unit or ""),
            "quantity": to_float(r.quantity), "basis": basis,
        })
    df = pd.DataFrame(records)
    key = ["section", "element", "location", "unit"]
    out = []
    for _, g in df.groupby(key, dropna=False):
        ai_qty = float(g.loc[g["basis"].eq("AI"), "quantity"].sum())
        dr_qty = float(g.loc[g["basis"].eq("Drawn"), "quantity"].sum())
        bases = sorted(set(g["basis"].tolist()))
        variance = 0.0
        if ai_qty == 0 and dr_qty == 0:
            status = "Manual / not yet measured"
        elif ai_qty > 0 and dr_qty > 0:
            variance = dr_qty - ai_qty
            tol = max(0.01, 0.02 * max(ai_qty, dr_qty))
            status = "Matched" if abs(variance) <= tol else "Difference"
        elif ai_qty > 0:
            variance = -ai_qty
            status = "AI only (not drawn)"
        else:
            variance = dr_qty
            status = "Drawn only (no AI draft)"
        out.append({
            "section": g["section"].iloc[0], "element": g["element"].iloc[0],
            "location": g["location"].iloc[0], "unit": g["unit"].iloc[0],
            "ai_qty": ai_qty, "drawn_qty": dr_qty, "variance": variance,
            "basis": ", ".join(bases), "status": status,
        })
    return pd.DataFrame(out)


def dataframe_for_takeoff(workspace_id: int) -> pd.DataFrame:
    df = ldf("SELECT * FROM takeoff_rows WHERE workspace_id=? ORDER BY id", (workspace_id,))
    if df.empty:
        return df
    basis = str(workspace_setting(workspace_id, "internal_pricing_basis", "wall_m2") or "wall_m2").strip().lower()
    floor_by_level = floor_area_by_level(df)
    paint_rows: List[float] = []
    labour_rows: List[float] = []
    value_rows: List[float] = []
    priced_qty_rows: List[float] = []
    basis_rows: List[str] = []
    for r in df.itertuples():
        qty = to_float(r.quantity)
        unit = str(r.unit or "")
        role = str(r.row_role or "").strip()
        coats = to_float(r.coats, 2)
        coverage = to_float(r.coverage_m2_per_litre, 12)
        productivity = to_float(r.productivity_m2_per_hour, 8)
        if role == "floor_area":
            paint_rows.append(0.0)
            labour_rows.append(0.0)
            value_rows.append(0.0)
            priced_qty_rows.append(qty if unit == "m²" else 0.0)
            basis_rows.append("Floor area (reference)")
            continue
        litres = paint_litres(qty, unit, coats, coverage)
        hours = labour_hours(qty, unit, productivity)
        paint_rows.append(litres)
        labour_rows.append(hours)
        if basis == "floor_m2" and unit == "m²" and is_internal_wall_row(r.section, r.element):
            floor_m2 = floor_by_level.get(level_of(r.location), 0.0)
            priced_qty_rows.append(floor_m2)
            basis_rows.append("Floor m²")
            value_rows.append(row_value(floor_m2, to_float(r.rate_per_unit)))
        else:
            priced_qty_rows.append(qty)
            basis_rows.append("Wall m²" if basis == "floor_m2" and unit == "m²" else "Quantity")
            value_rows.append(row_value(qty, to_float(r.rate_per_unit)))
    df["paint_litres"] = paint_rows
    df["labour_hours"] = labour_rows
    df["priced_quantity"] = priced_qty_rows
    df["pricing_basis"] = basis_rows
    df["value_ex_gst"] = value_rows
    return df


def _norm_key(text: Any) -> str:
    return re.sub(r"[^a-z0-9]+", "", str(text).lower().replace("²", "2").replace("³", "3"))


_TAKEOFF_HEADER_SYNONYMS = {
    "section": ["section", "division", "category", "part", "area type", "trade"],
    "element": ["element", "item", "scope", "work", "work item", "description", "task", "component"],
    "location": ["location", "room", "zone", "floor", "level", "site", "location/area", "area/location"],
    "substrate": ["substrate", "surface", "material", "background", "base material", "substrate/surface", "wall type", "surface type"],
    "finish_system": ["finish system", "finish", "system", "paint system", "spec", "coating", "finish spec", "paint type"],
    "quantity": ["quantity", "qty", "quantity (m2)", "qty (m2)", "quantity m2", "qty m2", "area m2", "area (m2)", "amount", "count", "qty (no.)", "qty (no)", "quantity (no.)"],
    "floor_area": ["floor area", "floor area m2", "floor m2", "internal floor area", "internal floor area m2", "gross floor area", "nett floor area", "net floor area", "floor area (m2)"],
    "unit": ["unit", "uom", "units", "unit of measure", "measure"],
    "quantity_status": ["quantity status", "status", "measurement status", "qty status", "takeoff status"],
    "source_page": ["source page", "page", "drawing", "sheet", "plan no", "drawing no", "ref page"],
    "source_reference": ["source reference", "reference", "ref", "source", "doc ref", "item ref"],
    "inclusion_status": ["inclusion status", "inclusion", "inc/exc", "inc exc", "include/exclude"],
    "coats": ["coats", "coat", "no coats", "number of coats", "coats no"],
    "coverage_m2_per_litre": ["coverage m2 per litre", "coverage", "m2 per litre", "m2/l", "m2/litre", "sqm/l", "spread rate"],
    "productivity_m2_per_hour": ["productivity m2 per hour", "productivity", "m2 per hour", "m2/hr", "m2/hour", "labour rate", "labour", "labor", "production rate"],
    "rate_per_unit": ["rate per unit", "rate", "unit rate", "rate per m2", "price", "rate ($)", "value"],
    "confidence": ["confidence", "basis", "measurement basis"],
    "notes": ["notes", "note", "comments", "remarks", "comment"],
}

_NORMALISED_HEADERS: Dict[str, str] = {}
for _target, _words in _TAKEOFF_HEADER_SYNONYMS.items():
    for _word in _words:
        _NORMALISED_HEADERS[_norm_key(_word)] = _target


def _match_takeoff_header(header: Any) -> Optional[str]:
    raw_lower = str(header or "").strip().lower()
    if not raw_lower:
        return None
    if "$" in raw_lower and _norm_key(raw_lower) in {"m2", "m", "sqm"}:
        return "rate_per_unit"
    key = _norm_key(raw_lower)
    if key in _NORMALISED_HEADERS:
        return _NORMALISED_HEADERS[key]
    if key in {"m2", "sqm", "sqmt", "area"}:
        return "quantity"
    if key in {"m2l", "sqml"}:
        return "coverage_m2_per_litre"
    if key in {"m2h", "m2hr", "sqmh"}:
        return "productivity_m2_per_hour"
    best_target: Optional[str] = None
    best_score = 0
    for target, words in _TAKEOFF_HEADER_SYNONYMS.items():
        score = 0
        for word in words:
            norm_word = _norm_key(word)
            if norm_word and norm_word in key:
                score += len(norm_word)
        if score > best_score:
            best_score = score
            best_target = target
    return best_target if best_score >= 4 else None


def _normalise_unit(raw: Any) -> str:
    text = str(raw or "").strip().lower()
    if not text:
        return ""
    if text in {"m2", "sqm", "sq m", "m²", "square metres", "square meter", "square metres"} or "m2" in text or "m²" in text:
        return "m²"
    if text in {"lm", "m", "lin m", "lineal", "lin m", "linear metre", "metre"} or "lm" == text:
        return "lm"
    if text in {"no", "no.", "nos", "ea", "each", "count", "each (no.)"}:
        return "No."
    if text in {"l", "litre", "litres", "ltr", "lt"}:
        return "L"
    if "item" in text:
        return "item"
    if "allow" in text:
        return "allowance"
    return text.title()


def _parse_qty(raw: Any) -> float:
    text = str(raw or "").strip()
    if not text:
        return 0.0
    cleaned = re.sub(r"[^0-9.\-]", "", text)
    return to_float(cleaned, 0.0)


def detect_takeoff_columns(upload: Any, header_row: Optional[int] = None) -> Tuple[List[str], List[List[Any]], int, int, int]:
    """Read an uploaded .xlsx/.xls/.csv take-off file.

    Returns (raw column headers, body rows, detected header row index, match score,
    total data rows). The header row is auto-detected from the row containing the
    most recognised take-off column names; when nothing is recognised it falls back
    to the first row so the caller can map columns manually.
    """
    name = str(getattr(upload, "name", "") or "takeoff")
    data = upload.getvalue()
    suffix = Path(name).suffix.lower()
    try:
        if suffix == ".csv":
            df = pd.read_csv(io.BytesIO(data), header=None, encoding_errors="ignore")
        else:
            df = pd.read_excel(io.BytesIO(data), sheet_name=0, header=None)
    except Exception as exc:
        raise RuntimeError(f"Could not read take-off file: {exc}")

    if df.empty or df.shape[1] == 0:
        raise RuntimeError("The take-off file is empty.")

    best_score = 0
    detected: Optional[int] = None
    if header_row is None:
        for i in range(min(30, len(df))):
            row = [str(v) for v in df.iloc[i].tolist()]
            score = sum(1 for cell in row if _match_takeoff_header(cell))
            if score > best_score:
                best_score = score
                detected = i
        if detected is None or best_score < 2:
            detected = 0
    else:
        detected = header_row

    detected = max(0, min(int(detected), len(df) - 1))
    raw_headers = [str(v).strip() for v in df.iloc[detected].tolist()]
    body = [list(line) for line in df.iloc[detected + 1:].astype(object).values.tolist()]
    return raw_headers, body, detected, best_score, len(df)


def parse_takeoff_file(upload: Any, mapping: Optional[Dict[int, str]] = None,
                       raw_headers: Optional[List[str]] = None,
                       body: Optional[List[List[Any]]] = None) -> Tuple[pd.DataFrame, List[str]]:
    """Build take-off rows from an uploaded file.

    mapping maps column index -> take-off field name. When None, columns are
    matched automatically from the header text. Returns (rows DataFrame, warnings).
    """
    name = str(getattr(upload, "name", "") or "takeoff")
    warnings: List[str] = []
    if raw_headers is None or body is None:
        raw_headers, body, _used, _score, _total = detect_takeoff_columns(upload)
    if mapping is None:
        mapping = {}
        used: List[str] = []
        for idx, h in enumerate(raw_headers):
            target = _match_takeoff_header(h)
            if target and target not in used:
                mapping[idx] = target
                used.append(target)

    rows: List[Dict[str, Any]] = []
    for line in body:
        row: Dict[str, Any] = {c: "" for c in TAKEOFF_COLUMNS}
        row["row_role"] = ""
        for idx, target in mapping.items():
            if idx >= len(line):
                continue
            value = line[idx]
            if target == "quantity":
                row["quantity"] = _parse_qty(value)
            elif target == "floor_area":
                row["quantity"] = _parse_qty(value)
                row["unit"] = "m²"
                row["row_role"] = "floor_area"
            elif target == "unit":
                row["unit"] = _normalise_unit(value)
            elif target in {"coats", "coverage_m2_per_litre", "productivity_m2_per_hour", "rate_per_unit"}:
                row[target] = to_float(value)
            else:
                row[target] = str(value or "").strip()
        text_signature = " ".join(str(row.get(c) or "") for c in ["element", "location", "section", "source_reference"]).strip()
        if not text_signature:
            continue
        if not row["row_role"]:
            element_text = str(row.get("element") or "").lower()
            location_text = str(row.get("location") or "").lower()
            if "floor area" in element_text or "internal floor" in element_text or "floor m2" in element_text or (location_text.startswith("floor area") or "floor area" in location_text):
                row["row_role"] = "floor_area"
        if not row.get("unit"):
            qty_header = next((raw_headers[i] for i, t in mapping.items() if t in {"quantity", "floor_area"}), "")
            qty_lower = qty_header.lower()
            if qty_header and "no" in qty_lower:
                row["unit"] = "No."
            elif qty_header and ("m2" in qty_lower or "m²" in qty_lower or "sq" in qty_lower or "area" in qty_lower):
                row["unit"] = "m²"
            elif qty_header and "lm" in qty_lower:
                row["unit"] = "lm"
            else:
                row["unit"] = "m²"
        row["quantity_status"] = row["quantity_status"] or ("Measured" if to_float(row["quantity"]) > 0 else "To measure")
        row["inclusion_status"] = row["inclusion_status"] or "INCLUSION"
        row["source_reference"] = row["source_reference"] or name
        rows.append(row)

    if not rows:
        raise RuntimeError("No take-off lines were found under the detected header row.")

    if len(mapping) < 4:
        warnings.append("Only a few columns were recognised — map them manually below or review the imported rows.")

    return pd.DataFrame(rows, columns=TAKEOFF_COLUMNS + ["row_role"]), warnings


def takeoff_import_panel(workspace_id: int, widget_key: str = "takeoff_import") -> None:
    """Render the take-off file import UI (header row picker, column mapping, preview, import)."""
    with st.expander("Import a take-off from an Excel or CSV file"):
        up = st.file_uploader("Take-off file (.xlsx, .xls, .csv)", type=["xlsx", "xls", "csv"], key=f"{widget_key}_file")
        if up is None:
            return
        try:
            raw_headers, body, used_row, best_score, total_rows = detect_takeoff_columns(up)
        except Exception as exc:
            st.error(f"Could not read that file as a take-off: {exc}")
            return
        hrow = st.number_input("Header row (1-based)", min_value=1, max_value=max(1, total_rows), value=int(used_row) + 1, key=f"{widget_key}_header_row")
        if int(hrow) != used_row + 1:
            try:
                raw_headers, body, used_row, best_score, total_rows = detect_takeoff_columns(up, header_row=int(hrow) - 1)
            except Exception as exc:
                st.error(f"Could not read that file as a take-off: {exc}")
                return
        if best_score < 2:
            st.info("No standard take-off column names were recognised in this file — set the header row above, then map the columns below.")
        cols = [h if str(h).strip() else f"Column {i+1}" for i, h in enumerate(raw_headers)]
        sig = f"{up.name}:{getattr(up, 'size', '')}:{used_row}"
        if st.session_state.get(f"{widget_key}_sig") != sig:
            st.session_state[f"{widget_key}_sig"] = sig
            auto = {i: _match_takeoff_header(h) for i, h in enumerate(raw_headers)}
            st.session_state[f"{widget_key}_map"] = pd.DataFrame({"Column": cols, "Maps to": [auto.get(i) or "" for i in range(len(cols))]})
            st.session_state.pop(f"{widget_key}_editor", None)
        map_df = st.data_editor(st.session_state[f"{widget_key}_map"], key=f"{widget_key}_editor", hide_index=True, use_container_width=True,
            column_config={"Column": st.column_config.TextColumn(disabled=True), "Maps to": st.column_config.SelectboxColumn(options=[""] + TAKEOFF_COLUMNS + ["floor_area"])})
        st.caption("Map a quantity column headed 'Floor area (m²)' / 'Floor m²' to **floor_area** so it becomes a per-level floor-area measurement row that can drive floor-m² pricing.")
        st.session_state[f"{widget_key}_map"] = map_df
        mapping = {}
        for i, m in enumerate(map_df["Maps to"]):
            if str(m).strip():
                mapping[i] = str(m).strip()
        if len(mapping) < 3:
            st.info("Use the 'Maps to' dropdowns above to map each column (at least an element/description, a quantity and a unit).")
        try:
            parsed_takeoff, import_warnings = parse_takeoff_file(up, mapping=mapping, raw_headers=raw_headers, body=body)
        except Exception as exc:
            st.error(f"Could not build take-off rows: {exc}")
            return
        for warning in import_warnings:
            st.warning(warning)
        st.dataframe(parsed_takeoff, use_container_width=True, hide_index=True)
        if st.button(f"Import {len(parsed_takeoff)} rows into the take-off schedule", type="primary", key=f"{widget_key}_button"):
            imported = 0
            for row in parsed_takeoff.to_dict("records"):
                if not any(str(row.get(c) or "").strip() for c in ["section", "element", "location", "source_reference"]):
                    continue
                if not to_float(row.get("rate_per_unit")):
                    row["rate_per_unit"] = default_rate_for(row.get("substrate"), row.get("element"), row.get("finish_system"), row.get("unit"))
                row_role = str(row.get("row_role") or "").strip()
                if row_role not in {"", "floor_area"}:
                    row_role = ""
                values = [row.get(col, "") for col in TAKEOFF_COLUMNS]
                lexecute("""INSERT INTO takeoff_rows(workspace_id,section,element,location,substrate,finish_system,quantity,unit,quantity_status,source_page,source_reference,inclusion_status,coats,coverage_m2_per_litre,productivity_m2_per_hour,rate_per_unit,confidence,notes,row_role,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""", (workspace_id, *values, row_role, now_stamp(), now_stamp()))
                imported += 1
            st.success(f"Imported {imported} take-off rows. Open the Take-off schedule tab to review them.")
            for k in (f"{widget_key}_file", f"{widget_key}_sig", f"{widget_key}_map", f"{widget_key}_editor", f"{widget_key}_header_row"):
                st.session_state.pop(k, None)
            st.rerun()


def register_df(workspace_id: int, name: str) -> pd.DataFrame:
    return ldf(
        "SELECT item_no,title,detail,priority,source_reference,status FROM register_items WHERE workspace_id=? AND register_name=? ORDER BY id",
        (workspace_id, name),
    )


def excel_export_bytes(workspace_id: int) -> bytes:
    workspace = lquery("SELECT * FROM workspaces WHERE id=?", (workspace_id,))[0]
    takeoff = dataframe_for_takeoff(workspace_id)
    pages = ldf(
        """SELECT p.page_label,p.page_type,p.scale_text,p.page_no,d.file_name,p.selected FROM pages p JOIN documents d ON d.id=p.document_id WHERE p.workspace_id=? ORDER BY d.id,p.page_no""",
        (workspace_id,),
    )
    docs = ldf("SELECT file_name,source_type,category,page_count,uploaded_at FROM documents WHERE workspace_id=? ORDER BY id", (workspace_id,))
    masses = ldf("SELECT label,level_name,x,y,z,width,depth,height,finish,source_reference,confidence,notes FROM model_masses WHERE workspace_id=? ORDER BY id", (workspace_id,))
    openings = ldf("SELECT label,opening_type,face,offset_x,offset_z,width,height,count,notes,source_reference FROM model_openings WHERE workspace_id=? ORDER BY id", (workspace_id,))
    zones = ldf("SELECT name,view_type,area_m2,wall_height_m,substrate,finish_system,quantity_status,source_reference FROM mapped_zones WHERE workspace_id=? ORDER BY id", (workspace_id,))
    project = pd.DataFrame(
        [
            ["Job number", workspace.get("job_no", "")],
            ["Project", workspace.get("job_name", "")],
            ["Builder / client", workspace.get("builder_client", "")],
            ["Site address", workspace.get("site_address", "")],
            ["Drawing issue", workspace.get("drawing_issue", "")],
            ["Estimator", workspace.get("estimator", "")],
            ["Status", workspace.get("status", "")],
            ["Generated", now_stamp()],
            ["Take-off method", "Premier Brushworks subscription paint take-off workflow"],
            ["3D model status", "Conceptual unless masses are marked Measured or Verified"],
        ],
        columns=["Field", "Value"],
    )
    summary = pd.DataFrame({"Executive Summary": [workspace.get("executive_summary", "")]})
    sheets: List[Tuple[str, pd.DataFrame]] = [
        ("Project Information", project),
        ("Executive Summary", summary),
        ("Source Documents", docs),
        ("Drawing Register", pages),
        ("Take-off Schedule", takeoff),
        ("Mapped Zones", zones),
        ("Door Schedule", register_df(workspace_id, "door_schedule")),
        ("Inclusions", register_df(workspace_id, "inclusions")),
        ("Exclusions", register_df(workspace_id, "exclusions")),
        ("Separate Clarifications", register_df(workspace_id, "clarifications")),
        ("Assumptions", register_df(workspace_id, "assumptions")),
        ("RFIs", register_df(workspace_id, "rfis")),
        ("Source & Basis", register_df(workspace_id, "source_basis")),
        ("Colours & Finishes", register_df(workspace_id, "colour_finish_schedule")),
        ("Access Constraints", register_df(workspace_id, "access_constraints")),
        ("Risks", register_df(workspace_id, "risks")),
        ("3D Masses", masses),
        ("3D Openings", openings),
    ]
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        for sheet_name, frame in sheets:
            frame.to_excel(writer, index=False, sheet_name=sheet_name[:31])
            ws = writer.book[sheet_name[:31]]
            ws.freeze_panes = "A2"
            ws.auto_filter.ref = ws.dimensions
            for cell in ws[1]:
                from openpyxl.styles import Alignment, Font, PatternFill
                cell.fill = PatternFill("solid", fgColor="171717")
                cell.font = Font(color="FFFFFF", bold=True)
                cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            for col in ws.columns:
                letter = col[0].column_letter
                max_len = min(65, max(10, max(len(str(c.value or "")) for c in col) + 2))
                ws.column_dimensions[letter].width = max_len
                for cell in col:
                    cell.alignment = Alignment(vertical="top", wrap_text=True)
    return output.getvalue()


def zip_export_bytes(workspace_id: int) -> bytes:
    workspace = lquery("SELECT * FROM workspaces WHERE id=?", (workspace_id,))[0]
    output = io.BytesIO()
    with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr(f"{safe_name(workspace.get('job_no'))}_takeoff.xlsx", excel_export_bytes(workspace_id))
        zf.writestr("model/building_geometry.json", json.dumps({
            "workspace": workspace,
            "masses": lquery("SELECT * FROM model_masses WHERE workspace_id=?", (workspace_id,)),
            "openings": lquery("SELECT * FROM model_openings WHERE workspace_id=?", (workspace_id,)),
            "mapped_zones": lquery("SELECT * FROM mapped_zones WHERE workspace_id=?", (workspace_id,)),
        }, indent=2, default=str))
        zf.writestr("model/building.obj", generate_obj(workspace_id))
        takeoff = dataframe_for_takeoff(workspace_id)
        zf.writestr("data/takeoff_schedule.csv", takeoff.to_csv(index=False))
        for doc in lquery("SELECT * FROM documents WHERE workspace_id=?", (workspace_id,)):
            path = Path(str(doc.get("path") or ""))
            if path.exists():
                zf.write(path, f"source_documents/{safe_name(doc.get('file_name'))}")
        for page in lquery("SELECT * FROM pages WHERE workspace_id=?", (workspace_id,)):
            path = Path(str(page.get("image_path") or ""))
            if path.exists():
                zf.write(path, f"rendered_pages/{safe_name(page.get('page_label'))}_{page.get('id')}.png")
        zf.writestr(
            "README.txt",
            "Premier Brushworks PlanReader export. The 3D model is conceptual unless geometry is marked Measured or Verified. All quantities must be reviewed against the current issued drawings and specifications before pricing or construction use.\n",
        )
    return output.getvalue()


def ensure_jobhub_takeoff_tables(bridge: JobHubBridge) -> None:
    if bridge.kind == "postgres":
        package_sql = """
        CREATE TABLE IF NOT EXISTS painting_takeoff_packages (
            id SERIAL PRIMARY KEY, job_id INTEGER NOT NULL, takeoff_no TEXT, takeoff_date TEXT,
            status TEXT, source_documents TEXT, interior_total_m2 REAL DEFAULT 0,
            exterior_total_m2 REAL DEFAULT 0, total_labour_hours REAL DEFAULT 0,
            total_paint_litres REAL DEFAULT 0, generated_method TEXT, assumptions TEXT,
            ai_notes TEXT, created_by TEXT, created_at TEXT, updated_at TEXT, notes TEXT)
        """
        line_sql = """
        CREATE TABLE IF NOT EXISTS painting_takeoff_lines (
            id SERIAL PRIMARY KEY, package_id INTEGER NOT NULL, area_type TEXT, location_area TEXT,
            substrate TEXT, labour_category TEXT, m2 REAL DEFAULT 0, unit TEXT, quantity REAL DEFAULT 0,
            coats REAL DEFAULT 0, productivity_m2_per_hour REAL DEFAULT 0, labour_hours REAL DEFAULT 0,
            finish_type TEXT, element_count REAL DEFAULT 0, lineal_metres REAL DEFAULT 0,
            paint_litres REAL DEFAULT 0, flags TEXT, notes TEXT, created_at TEXT)
        """
    else:
        package_sql = """
        CREATE TABLE IF NOT EXISTS painting_takeoff_packages (
            id INTEGER PRIMARY KEY AUTOINCREMENT, job_id INTEGER NOT NULL, takeoff_no TEXT, takeoff_date TEXT,
            status TEXT, source_documents TEXT, interior_total_m2 REAL DEFAULT 0,
            exterior_total_m2 REAL DEFAULT 0, total_labour_hours REAL DEFAULT 0,
            total_paint_litres REAL DEFAULT 0, generated_method TEXT, assumptions TEXT,
            ai_notes TEXT, created_by TEXT, created_at TEXT, updated_at TEXT, notes TEXT)
        """
        line_sql = """
        CREATE TABLE IF NOT EXISTS painting_takeoff_lines (
            id INTEGER PRIMARY KEY AUTOINCREMENT, package_id INTEGER NOT NULL, area_type TEXT, location_area TEXT,
            substrate TEXT, labour_category TEXT, m2 REAL DEFAULT 0, unit TEXT, quantity REAL DEFAULT 0,
            coats REAL DEFAULT 0, productivity_m2_per_hour REAL DEFAULT 0, labour_hours REAL DEFAULT 0,
            finish_type TEXT, element_count REAL DEFAULT 0, lineal_metres REAL DEFAULT 0,
            paint_litres REAL DEFAULT 0, flags TEXT, notes TEXT, created_at TEXT)
        """
    bridge.execute(package_sql)
    bridge.execute(line_sql)


def push_takeoff_to_jobhub(workspace_id: int, bridge: JobHubBridge, created_by: str) -> Tuple[int, int]:
    workspace = lquery("SELECT * FROM workspaces WHERE id=?", (workspace_id,))[0]
    job_id = workspace.get("jobhub_job_id")
    if not job_id:
        raise RuntimeError("This workspace is not linked to a JobHub job.")
    takeoff = dataframe_for_takeoff(workspace_id)
    if takeoff.empty:
        raise RuntimeError("There are no take-off rows to send.")
    takeoff = takeoff_work_rows(takeoff)
    if takeoff.empty:
        raise RuntimeError("There are no priced work rows to send (only floor-area measurement rows).")
    ensure_jobhub_takeoff_tables(bridge)
    docs = ", ".join(d["file_name"] for d in lquery("SELECT file_name FROM documents WHERE workspace_id=? ORDER BY id", (workspace_id,)))
    internal_mask = takeoff["section"].astype(str).str.lower().str.contains("internal|ceiling|door|joinery")
    external_mask = takeoff["section"].astype(str).str.lower().str.contains("external|facade|elevation|soffit|canopy")
    m2_mask = takeoff["unit"].astype(str).eq("m²")
    interior = float(takeoff.loc[internal_mask & m2_mask, "quantity"].fillna(0).sum())
    exterior = float(takeoff.loc[external_mask & m2_mask, "quantity"].fillna(0).sum())
    total_hours = float(takeoff["labour_hours"].fillna(0).sum())
    total_litres = float(takeoff["paint_litres"].fillna(0).sum())
    takeoff_no = f"PR-{workspace.get('job_no')}-{datetime.now().strftime('%Y%m%d-%H%M')}"
    if bridge.kind == "postgres":
        package_id = bridge.execute(
            """INSERT INTO painting_takeoff_packages(job_id,takeoff_no,takeoff_date,status,source_documents,interior_total_m2,exterior_total_m2,total_labour_hours,total_paint_litres,generated_method,assumptions,ai_notes,created_by,created_at,updated_at,notes)
               VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?) RETURNING id""",
            (job_id,takeoff_no,datetime.now().date().isoformat(),"Draft from PlanReader",docs,interior,exterior,total_hours,total_litres,"PB PlanReader subscription method","Quantities require estimator review.",workspace.get("executive_summary",""),created_by,now_stamp(),now_stamp(),"3D geometry is conceptual unless marked measured."),
            returning=True,
        )
    else:
        bridge.execute(
            """INSERT INTO painting_takeoff_packages(job_id,takeoff_no,takeoff_date,status,source_documents,interior_total_m2,exterior_total_m2,total_labour_hours,total_paint_litres,generated_method,assumptions,ai_notes,created_by,created_at,updated_at,notes)
               VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
            (job_id,takeoff_no,datetime.now().date().isoformat(),"Draft from PlanReader",docs,interior,exterior,total_hours,total_litres,"PB PlanReader subscription method","Quantities require estimator review.",workspace.get("executive_summary",""),created_by,now_stamp(),now_stamp(),"3D geometry is conceptual unless marked measured."),
        )
        package_id = bridge.query("SELECT id FROM painting_takeoff_packages WHERE takeoff_no=?", (takeoff_no,))[0]["id"]
    line_count = 0
    for _, row in takeoff.iterrows():
        unit = str(row.get("unit") or "")
        qty = to_float(row.get("quantity"))
        m2 = qty if unit == "m²" else 0
        lm = qty if unit == "lm" else 0
        count = qty if unit in {"No.", "item"} else 0
        bridge.execute(
            """INSERT INTO painting_takeoff_lines(package_id,area_type,location_area,substrate,labour_category,m2,unit,quantity,coats,productivity_m2_per_hour,labour_hours,finish_type,element_count,lineal_metres,paint_litres,flags,notes,created_at)
               VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
            (
                package_id,row.get("section",""),row.get("location",""),row.get("substrate",""),row.get("element",""),m2,unit,qty,row.get("coats",0),row.get("productivity_m2_per_hour",0),row.get("labour_hours",0),row.get("finish_system",""),count,lm,row.get("paint_litres",0),row.get("confidence",""),f"{row.get('notes','')} | Source: {row.get('source_reference','')}",now_stamp(),
            ),
        )
        line_count += 1
    return int(package_id), line_count


def pull_takeoff_from_jobhub(workspace_id: int, bridge: JobHubBridge) -> int:
    """Import take-off rows from JobHub's ``job_takeoff_rows`` into this workspace."""
    workspace = lquery("SELECT * FROM workspaces WHERE id=?", (workspace_id,))[0]
    job_id = workspace.get("jobhub_job_id")
    if not job_id:
        raise RuntimeError("This workspace is not linked to a JobHub job.")
    tables = set(bridge.table_names())
    if "job_takeoff_rows" not in tables:
        raise RuntimeError("No take-off rows exist in JobHub for linked jobs.")
    cols = set(bridge.columns("job_takeoff_rows"))
    selectable = [c for c in
                  ["id", "internal_external", "area_location", "substrate", "labour_category",
                   "qty_m2", "lineal_m", "count", "coats", "rate_ex_gst", "labour_hours",
                   "paint_litres", "value_ex_gst", "source_note", "confidence", "updated_at"]
                  if c in cols]
    rows = bridge.query(
        f"SELECT {', '.join(selectable)} FROM job_takeoff_rows WHERE job_id=? ORDER BY id",
        (int(job_id),),
    )
    if not rows:
        return 0
    # Replace any rows previously pulled from JobHub so re-imports stay in sync.
    lexecute(
        "DELETE FROM takeoff_rows WHERE workspace_id=? AND source_page='JobHub import'",
        (workspace_id,),
    )
    created = 0
    for row in rows:
        location = str(row.get("area_location") or "").strip()
        substrate = str(row.get("substrate") or "").strip()
        if not location or not substrate:
            continue
        qty_m2 = to_float(row.get("qty_m2"))
        lineal_m = to_float(row.get("lineal_m"))
        count = to_float(row.get("count"))
        if qty_m2 > 0:
            quantity, unit = qty_m2, "m²"
        elif lineal_m > 0:
            quantity, unit = lineal_m, "lm"
        elif count > 0:
            quantity, unit = count, "No."
        else:
            quantity, unit = 0.0, "m²"
        internal_external = str(row.get("internal_external") or "Internal")
        section = {
            "internal": "Internal walls and ceilings",
            "exterior": "Exterior walls and cladding",
            "external": "Exterior walls and cladding",
        }.get(internal_external.strip().lower(), "Internal walls and ceilings")
        rate = to_float(row.get("rate_ex_gst"))
        if rate <= 0:
            rate = default_rate_for(str(row.get("substrate") or ""), str(row.get("labour_category") or ""), "", unit)
        lexecute(
            """
            INSERT INTO takeoff_rows(workspace_id,section,element,location,substrate,finish_system,
                quantity,unit,quantity_status,source_page,source_reference,inclusion_status,
                coats,coverage_m2_per_litre,productivity_m2_per_hour,rate_per_unit,confidence,notes,row_role,created_at,updated_at)
            VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
            """,
            (
                workspace_id,
                section,
                str(row.get("labour_category") or ""),
                location,
                substrate,
                "",
                quantity,
                unit,
                "Measured" if quantity > 0 else "To measure",
                "JobHub import",
                str(row.get("source_note") or ""),
                "",
                to_float(row.get("coats"), 1),
                12.0,
                8.0,
                rate,
                str(row.get("confidence") or ""),
                (f"Imported from JobHub · labour {to_float(row.get('labour_hours')):.1f} hrs · "
                 f"paint {to_float(row.get('paint_litres')):.1f} L · value {to_float(row.get('value_ex_gst')):.2f}")
                if any(v is not None for v in [row.get("labour_hours"), row.get("paint_litres"), row.get("value_ex_gst")])
                else "Imported from JobHub",
                "",
                now_stamp(),
                now_stamp(),
            ),
        )
        created += 1
    lexecute(
        "UPDATE workspaces SET status='Draft', updated_at=? WHERE id=?",
        (now_stamp(), workspace_id),
    )
    return created


# -----------------------------------------------------------------------------
# Shared JobHub bridge: job numbering, one-file progress marker, live take-off
# -----------------------------------------------------------------------------


def ensure_shared_jobhub_schema(bridge: JobHubBridge) -> List[str]:
    """Create the shared PlanReader <-> JobHub tables used by live sync.

    Only creates tables that do not already exist so an existing JobHub
    database is never altered in place. Returns the names actually created.
    """
    created = []
    pk = "INTEGER PRIMARY KEY AUTOINCREMENT" if bridge.kind == "sqlite" else "SERIAL PRIMARY KEY"
    ddl = {
        "jobs": f"""
            CREATE TABLE IF NOT EXISTS jobs (
                id {pk}, job_no TEXT UNIQUE, job_name TEXT, builder_client_id INTEGER,
                site_address TEXT, status TEXT, leading_hand TEXT, start_date TEXT,
                end_date TEXT, contract_value REAL, notes TEXT)
        """,
        "job_takeoff_rows": f"""
            CREATE TABLE IF NOT EXISTS job_takeoff_rows (
                id {pk}, job_id INTEGER NOT NULL, internal_external TEXT DEFAULT 'Internal',
                area_location TEXT NOT NULL, substrate TEXT NOT NULL, labour_category TEXT,
                qty_m2 REAL DEFAULT 0, lineal_m REAL DEFAULT 0, count REAL DEFAULT 0,
                coats REAL DEFAULT 1, rate_ex_gst REAL DEFAULT 0, labour_hours REAL DEFAULT 0,
                paint_litres REAL DEFAULT 0, value_ex_gst REAL DEFAULT 0, source_note TEXT,
                confidence TEXT, updated_at TEXT NOT NULL,
                UNIQUE(job_id, area_location, substrate))
        """,
        "job_colour_schedules": f"""
            CREATE TABLE IF NOT EXISTS job_colour_schedules (
                id {pk}, job_id INTEGER NOT NULL, area_location TEXT NOT NULL, surface TEXT NOT NULL,
                colour TEXT, finish TEXT, product TEXT, notes TEXT, hex TEXT, updated_at TEXT NOT NULL,
                UNIQUE(job_id, area_location, surface))
        """,
        "job_document_blobs": f"""
            CREATE TABLE IF NOT EXISTS job_document_blobs (
                id {pk}, job_id INTEGER NOT NULL, file_name TEXT NOT NULL, mime_type TEXT,
                doc_type TEXT, notes TEXT, blob_data TEXT NOT NULL, created_at TEXT NOT NULL,
                UNIQUE(job_id, file_name))
        """,
    }
    for name, sql in ddl.items():
        try:
            bridge.execute(sql)
            created.append(name)
        except Exception:
            pass
    return created


def next_jobhub_job_no(bridge: Optional[JobHubBridge]) -> str:
    """Next sequential PB number from the shared JobHub jobs table.

    Mirrors JobHub's ``next_job_no()``: no PB rows -> PB25001, otherwise the
    last PB number + 1, formatted as ``{prefix}{int(digits)+1:05d}``. Falls
    back to a dated local number when JobHub is not reachable.
    """
    if not bridge:
        return f"PB-{datetime.now().strftime('%y%m%d')}"
    try:
        tables = set(bridge.table_names())
    except Exception:
        return f"PB-{datetime.now().strftime('%y%m%d')}"
    if "jobs" not in tables:
        return "PB25001"
    try:
        cols = set(bridge.columns("jobs"))
    except Exception:
        cols = set()
    if "job_no" not in cols:
        return "PB25001"
    try:
        rows = bridge.query("SELECT job_no FROM jobs WHERE job_no LIKE 'PB%' ORDER BY job_no DESC LIMIT 1")
    except Exception:
        return f"PB-{datetime.now().strftime('%y%m%d')}"
    if not rows:
        return "PB25001"
    last = str(rows[0].get("job_no") or "")
    digits = "".join(c for c in last if c.isdigit())
    prefix = "".join(c for c in last if not c.isdigit())
    if not digits:
        return "PB25001"
    return f"{prefix}{int(digits) + 1:05d}"


def create_linked_jobhub_job(bridge: Optional[JobHubBridge], job_no: str, job_name: str = "", site_address: str = "", status: str = "Active") -> Optional[int]:
    """Insert a shared JobHub job row if that job number does not exist yet.

    Returns the JobHub ``jobs.id`` for the job, or None when the bridge is
    unavailable. Used when creating a standalone workspace with the
    ``link to JobHub`` option so both apps share the same job number.
    """
    if not bridge:
        return None
    job_no = str(job_no or "").strip()
    if not job_no:
        return None
    try:
        tables = set(bridge.table_names())
    except Exception:
        return None
    if "jobs" not in tables:
        try:
            ensure_shared_jobhub_schema(bridge)
        except Exception:
            return None
    try:
        existing = bridge.query("SELECT id FROM jobs WHERE job_no=?", (job_no,))
        if existing:
            return int(existing[0]["id"])
    except Exception:
        return None
    try:
        bridge.execute(
            "INSERT INTO jobs(job_no,job_name,site_address,status,notes,start_date,end_date) VALUES(?,?,?,?,?,?,?)",
            (job_no, str(job_name or ""), str(site_address or ""), str(status or "Active"),
             "Created by PB PlanReader sync.", datetime.now().date().isoformat(), ""),
        )
    except Exception:
        return None
    try:
        existing = bridge.query("SELECT id FROM jobs WHERE job_no=?", (job_no,))
        if existing:
            return int(existing[0]["id"])
    except Exception:
        pass
    return None


def _jobhub_takeoff_lines_csv(workspace: Dict[str, Any], takeoff: pd.DataFrame) -> str:
    """Take-off rows in JobHub's shared ``job_takeoff_rows`` column order."""
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow(["internal_external", "area_location", "substrate", "labour_category",
                     "qty_m2", "lineal_m", "count", "coats", "rate_ex_gst",
                     "labour_hours", "paint_litres", "value_ex_gst", "source_note", "confidence"])
    for _, row in takeoff_work_rows(takeoff).iterrows():
        unit = str(row.get("unit") or "")
        qty = to_float(row.get("quantity"))
        m2 = qty if unit == "m²" else 0
        lm = qty if unit == "lm" else 0
        count = qty if unit in {"No.", "item"} else 0
        section = str(row.get("section") or "Internal walls and ceilings")
        internal_external = "Internal" if "internal" in section.lower() else ("External" if "external" in section.lower() else "Internal")
        writer.writerow([
            internal_external,
            str(row.get("location") or row.get("element") or "General"),
            str(row.get("substrate") or "Other"),
            str(row.get("element") or ""),
            m2, lm, count,
            to_float(row.get("coats"), 1),
            to_float(row.get("rate_per_unit")),
            to_float(row.get("labour_hours")),
            to_float(row.get("paint_litres")),
            to_float(row.get("value_ex_gst")),
            str(row.get("source_reference") or ""),
            str(row.get("confidence") or ""),
        ])
    return buf.getvalue()


def _takeoff_summary_csv(takeoff: pd.DataFrame) -> str:
    if takeoff.empty:
        return "section,m2,lineal_m,count,paint_litres,labour_hours,value_ex_gst\n"
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow(["section", "m2", "lineal_m", "count", "paint_litres", "labour_hours", "value_ex_gst"])
    for section, group in takeoff_work_rows(takeoff).groupby("section", dropna=False):
        writer.writerow([
            str(section or "Unassigned"),
            float(group.loc[group["unit"].eq("m²"), "quantity"].sum()),
            float(group.loc[group["unit"].eq("lm"), "quantity"].sum()),
            float(group.loc[group["unit"].isin({"No.", "item"}), "quantity"].sum()),
            float(group["paint_litres"].sum()),
            float(group["labour_hours"].sum()),
            float(group["value_ex_gst"].sum()),
        ])
    return buf.getvalue()


def _progress_package_readme(workspace: Dict[str, Any], takeoff: pd.DataFrame, pages: pd.DataFrame) -> str:
    lines = [
        "PB PlanReader — progress marker package",
        "=======================================",
        f"Job number:        {workspace.get('job_no', '')}",
        f"Project:           {workspace.get('job_name', '')}",
        f"Builder / client:  {workspace.get('builder_client', '')}",
        f"Site address:      {workspace.get('site_address', '')}",
        f"Drawing issue:     {workspace.get('drawing_issue', '')}",
        f"Estimator:         {workspace.get('estimator', '')}",
        f"Status:            {workspace.get('status', '')}",
        f"Generated:         {now_stamp()}",
        f"3D model status:   Conceptual unless geometry is marked Measured or Verified.",
        "",
        "This package is the single progress marker for the job:",
        "  3d/3d_progress_marker.html      interactive 3D render (opens in a browser)",
        "  3d/building.obj                 OBJ geometry",
        "  3d/building_geometry.json       masses, openings, mapped zones, measurement lines",
        "  takeoff/paint_takeoff.xlsx      subscription-style Excel take-off pack",
        "  takeoff/takeoff_schedule.csv    full take-off schedule",
        "  takeoff/quantity_summary.csv    quantities by section",
        "  takeoff/takeoff_lines_jobhub.csv  take-off rows in JobHub shared format",
        "  registers/*.csv                 door schedule, inclusions, exclusions, RFIs, etc.",
        "  source_documents/               original plan / spec documents",
        "  rendered_pages/                 processed drawing page images",
        "  package_manifest.json           machine-readable manifest with totals",
        "",
    ]
    if takeoff.empty:
        lines.append("No measured take-off rows are in this package yet.")
    else:
        work = takeoff_work_rows(takeoff)
        lines.append(f"Take-off lines:   {len(work)}")
        lines.append(f"Measured m2:      {float(work.loc[work['unit'].eq('m2' if 'm2' in set(work['unit'].astype(str)) else 'm²'), 'quantity'].sum()):,.2f}")
        lines.append(f"Paint litres:     {float(work['paint_litres'].sum()):,.2f}")
        lines.append(f"Labour hours:     {float(work['labour_hours'].sum()):,.2f}")
        lines.append(f"Value ex GST:     ${float(work['value_ex_gst'].sum()):,.2f}")
        if "row_role" in takeoff.columns:
            floor_m2 = float(takeoff.loc[takeoff["row_role"].fillna("").eq("floor_area") & takeoff["unit"].eq("m²"), "quantity"].sum())
            if floor_m2:
                lines.append(f"Floor m2 (ref):   {floor_m2:,.2f}")
    lines += [
        "",
        "All quantities must be reviewed against the current issued drawings and",
        "specifications before pricing or construction use. The 3D model is not",
        "construction-grade BIM.",
    ]
    return "\n".join(lines)


def progress_package_bytes(workspace_id: int) -> bytes:
    """Build the single 'progress marker' ZIP: 3D render + take-off + job documents."""
    workspace = lquery("SELECT * FROM workspaces WHERE id=?", (workspace_id,))[0]
    output = io.BytesIO()
    takeoff = dataframe_for_takeoff(workspace_id)
    pages = ldf(
        """SELECT p.id,p.page_label,p.page_type,p.scale_text,p.page_no,p.selected,p.image_path,d.file_name
           FROM pages p JOIN documents d ON d.id=p.document_id WHERE p.workspace_id=? ORDER BY d.id,p.page_no""",
        (workspace_id,),
    )
    with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("README.txt", _progress_package_readme(workspace, takeoff, pages))
        work_takeoff = takeoff_work_rows(takeoff)
        manifest = {
            "app": APP_NAME,
            "app_version": APP_VERSION,
            "job_no": workspace.get("job_no", ""),
            "job_name": workspace.get("job_name", ""),
            "builder_client": workspace.get("builder_client", ""),
            "site_address": workspace.get("site_address", ""),
            "drawing_issue": workspace.get("drawing_issue", ""),
            "estimator": workspace.get("estimator", ""),
            "status": workspace.get("status", ""),
            "generated": now_stamp(),
            "executive_summary": workspace.get("executive_summary", ""),
            "totals": {
                "m2": float(work_takeoff.loc[work_takeoff["unit"].eq("m²"), "quantity"].sum()) if not work_takeoff.empty else 0.0,
                "lm": float(work_takeoff.loc[work_takeoff["unit"].eq("lm"), "quantity"].sum()) if not work_takeoff.empty else 0.0,
                "count": float(work_takeoff.loc[work_takeoff["unit"].isin({"No.", "item"}), "quantity"].sum()) if not work_takeoff.empty else 0.0,
                "paint_litres": float(work_takeoff["paint_litres"].sum()) if not work_takeoff.empty else 0.0,
                "labour_hours": float(work_takeoff["labour_hours"].sum()) if not work_takeoff.empty else 0.0,
                "value_ex_gst": float(work_takeoff["value_ex_gst"].sum()) if not work_takeoff.empty else 0.0,
            },
            "pages": pages.to_dict("records"),
            "measured_rows": int((work_takeoff["quantity_status"].astype(str).str.lower().str.contains("measur") if not work_takeoff.empty else pd.Series(dtype=bool)).sum()),
        }
        zf.writestr("package_manifest.json", json.dumps(manifest, indent=2, default=str))
        zf.writestr("3d/3d_progress_marker.html", build_3d_figure(workspace_id).to_html(full_html=True, include_plotlyjs=True))
        zf.writestr("3d/building.obj", generate_obj(workspace_id))
        zf.writestr("3d/building_geometry.json", json.dumps({
            "workspace": workspace,
            "masses": lquery("SELECT * FROM model_masses WHERE workspace_id=?", (workspace_id,)),
            "openings": lquery("SELECT * FROM model_openings WHERE workspace_id=?", (workspace_id,)),
            "mapped_zones": lquery("SELECT * FROM mapped_zones WHERE workspace_id=?", (workspace_id,)),
            "measurement_lines": lquery("SELECT * FROM measurement_lines WHERE workspace_id=?", (workspace_id,)),
        }, indent=2, default=str))
        zf.writestr("takeoff/paint_takeoff.xlsx", excel_export_bytes(workspace_id))
        if not takeoff.empty:
            zf.writestr("takeoff/takeoff_schedule.csv", takeoff.to_csv(index=False))
            zf.writestr("takeoff/quantity_summary.csv", _takeoff_summary_csv(takeoff))
            zf.writestr("takeoff/takeoff_lines_jobhub.csv", _jobhub_takeoff_lines_csv(workspace, takeoff))
        for register_name, label in [
            ("door_schedule", "door_schedule"),
            ("inclusions", "inclusions"),
            ("exclusions", "exclusions"),
            ("clarifications", "separate_clarifications"),
            ("assumptions", "assumptions"),
            ("rfis", "rfis"),
            ("source_basis", "source_and_basis"),
            ("colour_finish_schedule", "colours_and_finishes"),
            ("access_constraints", "access_constraints"),
            ("risks", "risks"),
        ]:
            frame = register_df(workspace_id, register_name)
            if not frame.empty:
                zf.writestr(f"registers/{label}.csv", frame.to_csv(index=False))
        if not pages.empty:
            zf.writestr("registers/drawing_register.csv", pages.drop(columns=["image_path"], errors="ignore").to_csv(index=False))
        else:
            zf.writestr("registers/drawing_register.csv", "page_label,page_type,scale_text,page_no,file_name,selected\n")
        for doc in lquery("SELECT * FROM documents WHERE workspace_id=?", (workspace_id,)):
            path = Path(str(doc.get("path") or ""))
            if path.exists():
                zf.write(path, f"source_documents/{safe_name(doc.get('file_name'))}")
        for page in pages.itertuples(index=False):
            path = Path(str(getattr(page, "image_path") or ""))
            if path.exists():
                zf.write(path, f"rendered_pages/{safe_name(page.file_name)}_{int(page.id)}.png")
    return output.getvalue()


def _sync_jobhub_takeoff_rows(bridge: JobHubBridge, job_id: int, takeoff: pd.DataFrame) -> int:
    """Upsert the workspace take-off into the shared ``job_takeoff_rows`` table."""
    takeoff = takeoff_work_rows(takeoff)
    if takeoff.empty:
        return 0
    stamp = now_stamp()
    payloads = []
    for _, row in takeoff.iterrows():
        unit = str(row.get("unit") or "")
        qty = to_float(row.get("quantity"))
        m2 = qty if unit == "m²" else 0
        lm = qty if unit == "lm" else 0
        count = qty if unit in {"No.", "item"} else 0
        section = str(row.get("section") or "Internal walls and ceilings")
        internal_external = "Internal" if "internal" in section.lower() else ("External" if "external" in section.lower() else "Internal")
        payloads.append((
            int(job_id), internal_external,
            str(row.get("location") or row.get("element") or "General"),
            str(row.get("substrate") or "Other"),
            str(row.get("element") or ""),
            m2, lm, count,
            to_float(row.get("coats"), 1),
            to_float(row.get("rate_per_unit")),
            to_float(row.get("labour_hours")),
            to_float(row.get("paint_litres")),
            to_float(row.get("value_ex_gst")),
            str(row.get("source_reference") or ""),
            str(row.get("confidence") or ""),
            stamp,
        ))
    for payload in payloads:
        bridge.execute(
            """
            INSERT INTO job_takeoff_rows(job_id,internal_external,area_location,substrate,labour_category,
                qty_m2,lineal_m,count,coats,rate_ex_gst,labour_hours,paint_litres,value_ex_gst,source_note,confidence,updated_at)
            VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
            ON CONFLICT(job_id, area_location, substrate) DO UPDATE SET
                internal_external=excluded.internal_external, labour_category=excluded.labour_category,
                qty_m2=excluded.qty_m2, lineal_m=excluded.lineal_m, count=excluded.count,
                coats=excluded.coats, rate_ex_gst=excluded.rate_ex_gst,
                labour_hours=excluded.labour_hours, paint_litres=excluded.paint_litres,
                value_ex_gst=excluded.value_ex_gst, source_note=excluded.source_note,
                confidence=excluded.confidence, updated_at=excluded.updated_at
            """,
            payload,
        )
    return len(payloads)


def push_progress_marker_to_jobhub(workspace_id: int, bridge: JobHubBridge, created_by: str) -> Dict[str, Any]:
    """Send the one-file progress package to the shared JobHub database.

    Stores the ZIP (3D render + take-off + job documents) as a progress marker
    blob on the linked job, upserts live take-off rows into the shared
    ``job_takeoff_rows`` table, records the marker in ``planreader_documents``
    when present, and keeps the existing draft take-off package history.
    """
    workspace = lquery("SELECT * FROM workspaces WHERE id=?", (workspace_id,))[0]
    job_id = workspace.get("jobhub_job_id")
    if not job_id:
        raise RuntimeError("Link this workspace to a JobHub job first (or create it as a linked job).")
    takeoff = dataframe_for_takeoff(workspace_id)
    if takeoff.empty:
        raise RuntimeError("There are no take-off rows to send in the progress marker.")
    ensure_shared_jobhub_schema(bridge)
    file_name = f"progress_marker_{safe_name(workspace.get('job_no'))}_{datetime.now().strftime('%Y%m%d_%H%M')}.zip"
    blob_bytes = progress_package_bytes(workspace_id)
    stamp = now_stamp()
    try:
        bridge.execute(
            """
            INSERT INTO job_document_blobs(job_id,file_name,mime_type,doc_type,notes,blob_data,created_at)
            VALUES(?,?,?,?,?,?,?)
            ON CONFLICT(job_id, file_name) DO UPDATE SET
                mime_type=excluded.mime_type, doc_type=excluded.doc_type, notes=excluded.notes,
                blob_data=excluded.blob_data, created_at=excluded.created_at
            """,
            (int(job_id), file_name, "application/zip", "Progress Marker",
             f"PB PlanReader progress marker: 3D render + take-off + job documents. Generated by {created_by}.",
             base64.b64encode(blob_bytes).decode("ascii"), stamp),
        )
    except Exception:
        pass
    synced = _sync_jobhub_takeoff_rows(bridge, int(job_id), takeoff)
    try:
        ensure_planreader_document_table(bridge)
        bridge.execute(
            "INSERT INTO planreader_documents(job_id,file_name,mime_type,storage_path,source_app,uploaded_by,uploaded_at,notes) VALUES(?,?,?,?,?,?,?,?)",
            (int(job_id), file_name, "application/zip", "", "PlanReader", created_by, stamp,
             "One-file progress marker package (3D render + take-off + documents)."),
        )
    except Exception:
        pass
    package_id, line_count = push_takeoff_to_jobhub(workspace_id, bridge, created_by)
    return {
        "file_name": file_name,
        "size_bytes": len(blob_bytes),
        "blob_recorded": True,
        "takeoff_rows_synced": synced,
        "package_id": int(package_id),
        "package_lines": int(line_count),
        "job_id": int(job_id),
    }


def publish_job_to_jobhub(workspace_id: int, bridge: JobHubBridge, created_by: str) -> Dict[str, Any]:
    """Publish the final take-off and quotation to the shared JobHub job.

    Creates a final take-off package marked ``Published``, writes the priced
    Excel quotation as a ``Final quotation`` document blob, stores the one-file
    progress package, and updates the shared ``jobs.status`` to ``Published``.
    """
    workspace = lquery("SELECT * FROM workspaces WHERE id=?", (workspace_id,))[0]
    job_id = workspace.get("jobhub_job_id")
    if not job_id:
        raise RuntimeError("Link this workspace to a JobHub job before publishing.")
    takeoff = dataframe_for_takeoff(workspace_id)
    if takeoff.empty:
        raise RuntimeError("There are no take-off rows to publish.")
    takeoff = takeoff_work_rows(takeoff)
    ensure_shared_jobhub_schema(bridge)
    ensure_jobhub_takeoff_tables(bridge)
    stamp = now_stamp()
    job_no = safe_name(workspace.get("job_no"))
    quote_name = f"quotation_{job_no}_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx"
    quote_bytes = quote_workbook_bytes(workspace_id)
    try:
        bridge.execute(
            """INSERT INTO job_document_blobs(job_id,file_name,mime_type,doc_type,notes,blob_data,created_at)
               VALUES(?,?,?,?,?,?,?)
               ON CONFLICT(job_id, file_name) DO UPDATE SET
                   mime_type=excluded.mime_type, doc_type=excluded.doc_type, notes=excluded.notes,
                   blob_data=excluded.blob_data, created_at=excluded.created_at""",
            (int(job_id), quote_name, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
             "Final quotation",
             f"Priced per-level quotation generated by PB PlanReader ({created_by}).",
             base64.b64encode(quote_bytes).decode("ascii"), stamp),
        )
    except Exception:
        pass
    progress_name = f"progress_marker_{job_no}_{datetime.now().strftime('%Y%m%d_%H%M')}.zip"
    progress_bytes = progress_package_bytes(workspace_id)
    try:
        bridge.execute(
            """INSERT INTO job_document_blobs(job_id,file_name,mime_type,doc_type,notes,blob_data,created_at)
               VALUES(?,?,?,?,?,?,?)
               ON CONFLICT(job_id, file_name) DO UPDATE SET
                   mime_type=excluded.mime_type, doc_type=excluded.doc_type, notes=excluded.notes,
                   blob_data=excluded.blob_data, created_at=excluded.created_at""",
            (int(job_id), progress_name, "application/zip", "Progress Marker",
             f"Published take-off + 3D render + documents ({created_by}).",
             base64.b64encode(progress_bytes).decode("ascii"), stamp),
        )
    except Exception:
        pass
    synced = _sync_jobhub_takeoff_rows(bridge, int(job_id), takeoff)
    docs = ", ".join(d["file_name"] for d in lquery("SELECT file_name FROM documents WHERE workspace_id=? ORDER BY id", (workspace_id,)))
    internal_mask = takeoff["section"].astype(str).str.lower().str.contains("internal|ceiling|door|joinery")
    external_mask = takeoff["section"].astype(str).str.lower().str.contains("external|facade|elevation|soffit|canopy")
    m2_mask = takeoff["unit"].astype(str).eq("m²")
    interior = float(takeoff.loc[internal_mask & m2_mask, "quantity"].fillna(0).sum())
    exterior = float(takeoff.loc[external_mask & m2_mask, "quantity"].fillna(0).sum())
    total_hours = float(takeoff["labour_hours"].fillna(0).sum())
    total_litres = float(takeoff["paint_litres"].fillna(0).sum())
    takeoff_no = f"PR-{workspace.get('job_no')}-PUB-{datetime.now().strftime('%Y%m%d-%H%M')}"
    if bridge.kind == "postgres":
        package_id = bridge.execute(
            """INSERT INTO painting_takeoff_packages(job_id,takeoff_no,takeoff_date,status,source_documents,interior_total_m2,exterior_total_m2,total_labour_hours,total_paint_litres,generated_method,assumptions,ai_notes,created_by,created_at,updated_at,notes)
               VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?) RETURNING id""",
            (job_id, takeoff_no, datetime.now().date().isoformat(), "Published", docs, interior, exterior,
             total_hours, total_litres, "PB PlanReader subscription method",
             "Final quantities approved for pricing.", workspace.get("executive_summary", ""),
             created_by, stamp, stamp, "Published by PB PlanReader."),
            returning=True,
        )
    else:
        bridge.execute(
            """INSERT INTO painting_takeoff_packages(job_id,takeoff_no,takeoff_date,status,source_documents,interior_total_m2,exterior_total_m2,total_labour_hours,total_paint_litres,generated_method,assumptions,ai_notes,created_by,created_at,updated_at,notes)
               VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
            (job_id, takeoff_no, datetime.now().date().isoformat(), "Published", docs, interior, exterior,
             total_hours, total_litres, "PB PlanReader subscription method",
             "Final quantities approved for pricing.", workspace.get("executive_summary", ""),
             created_by, stamp, stamp, "Published by PB PlanReader."),
        )
        package_id = bridge.query("SELECT id FROM painting_takeoff_packages WHERE takeoff_no=?", (takeoff_no,))[0]["id"]
    line_count = 0
    for _, row in takeoff.iterrows():
        unit = str(row.get("unit") or "")
        qty = to_float(row.get("quantity"))
        m2 = qty if unit == "m²" else 0
        lm = qty if unit == "lm" else 0
        count = qty if unit in {"No.", "item"} else 0
        bridge.execute(
            """INSERT INTO painting_takeoff_lines(package_id,area_type,location_area,substrate,labour_category,m2,unit,quantity,coats,productivity_m2_per_hour,labour_hours,finish_type,element_count,lineal_metres,paint_litres,flags,notes,created_at)
               VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
            (package_id, row.get("section", ""), row.get("location", ""), row.get("substrate", ""), row.get("element", ""),
             m2, unit, qty, row.get("coats", 0), row.get("productivity_m2_per_hour", 0), row.get("labour_hours", 0),
             row.get("finish_system", ""), count, lm, row.get("paint_litres", 0), row.get("confidence", ""),
             f"{row.get('notes', '')} | Source: {row.get('source_reference', '')}", stamp),
        )
        line_count += 1
    try:
        bridge.execute(
            "UPDATE jobs SET status=?, notes=COALESCE(notes, '') || ' | Published by PB PlanReader.' WHERE id=?",
            ("Published", int(job_id)),
        )
    except Exception:
        pass
    try:
        ensure_planreader_document_table(bridge)
        bridge.execute(
            "INSERT INTO planreader_documents(job_id,file_name,mime_type,storage_path,source_app,uploaded_by,uploaded_at,notes) VALUES(?,?,?,?,?,?,?,?)",
            (int(job_id), quote_name, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "",
             "PlanReader", created_by, stamp, "Final priced quotation."),
        )
    except Exception:
        pass
    return {
        "package_id": int(package_id),
        "package_lines": int(line_count),
        "job_id": int(job_id),
        "quotation": quote_name,
        "progress_marker": progress_name,
        "takeoff_rows_synced": synced,
        "job_status": "Published",
    }


def list_jobhub_progress_markers(bridge: Optional[JobHubBridge], job_id: int) -> List[Dict[str, Any]]:
    """Progress marker records already stored on a JobHub job."""
    if not bridge or not job_id:
        return []
    try:
        tables = set(bridge.table_names())
    except Exception:
        return []
    if "job_document_blobs" not in tables:
        return []
    try:
        cols = set(bridge.columns("job_document_blobs"))
    except Exception:
        cols = set()
    selectable = [c for c in ["id", "file_name", "mime_type", "doc_type", "notes", "created_at"] if c in cols]
    if "doc_type" not in cols:
        rows = bridge.query(
            f"SELECT {', '.join(selectable)} FROM job_document_blobs WHERE job_id=? ORDER BY id DESC",
            (int(job_id),),
        )
    else:
        rows = bridge.query(
            f"SELECT {', '.join(selectable)} FROM job_document_blobs WHERE job_id=? AND doc_type LIKE ? ORDER BY id DESC",
            (int(job_id), "%Progress Marker%"),
        )
    return rows if rows else []


# -----------------------------------------------------------------------------
# UI helpers and pages
# -----------------------------------------------------------------------------


def hero(workspace: Optional[Dict[str, Any]] = None) -> None:
    if workspace:
        title = f"{workspace.get('job_no','')} — {workspace.get('job_name','')}"
        sub = f"{workspace.get('builder_client','')} · {workspace.get('site_address','')}"
    else:
        title = APP_NAME
        sub = "Standalone commercial painting plan reader, subscription take-off workflow and interactive 3D building model."
    st.markdown(f"<div class='pb-hero'><h1>{title}</h1><p>{sub}</p></div>", unsafe_allow_html=True)


def current_workspace() -> Optional[Dict[str, Any]]:
    workspace_id = st.session_state.get("workspace_id")
    if not workspace_id:
        return None
    rows = lquery("SELECT * FROM workspaces WHERE id=?", (int(workspace_id),))
    return rows[0] if rows else None


def login_screen(bridge: Optional[JobHubBridge]) -> None:
    app_css()
    hero()
    st.markdown("<div class='pb-card'>", unsafe_allow_html=True)
    st.subheader("Sign in")
    try:
        has_user_table = bool(bridge) and "app_users" in set(bridge.table_names())
    except Exception:
        has_user_table = False
    if has_user_table:
        st.caption("Use the same username and password as JobHub.")
        with st.form("login_form"):
            username = st.text_input("Username")
            password = st.text_input("Password", type="password")
            submitted = st.form_submit_button("Sign in", type="primary", use_container_width=True)
        if submitted:
            user = authenticate_jobhub_user(bridge, username, password)
            if user:
                st.session_state["planreader_user"] = user
                st.rerun()
            st.error("Invalid username or password.")
    else:
        local_password = os.environ.get("PLANREADER_PASSWORD", "")
        if local_password:
            with st.form("local_login"):
                password = st.text_input("PlanReader password", type="password")
                submitted = st.form_submit_button("Sign in", type="primary", use_container_width=True)
            if submitted:
                if password == local_password:
                    st.session_state["planreader_user"] = {"username": "planreader", "role": "admin"}
                    st.rerun()
                st.error("Incorrect password.")
        else:
            st.info("No JobHub user table or PLANREADER_PASSWORD is configured. This local instance is open.")
            if st.button("Continue", type="primary"):
                st.session_state["planreader_user"] = {"username": "local", "role": "admin"}
                st.rerun()
    st.markdown("</div>", unsafe_allow_html=True)
    st.stop()


def sidebar_workspace_selector(bridge: Optional[JobHubBridge]) -> Optional[int]:
    st.sidebar.markdown("## PB PlanReader")
    st.sidebar.caption(f"Standalone app · v{APP_VERSION}")
    if st.sidebar.button("Sign out"):
        for key in ["planreader_user", "workspace_id"]:
            st.session_state.pop(key, None)
        st.rerun()
    st.sidebar.markdown("### Job workspace")
    if bridge:
        try:
            jobs = fetch_jobhub_jobs(bridge)
        except Exception as exc:
            jobs = []
            st.sidebar.error(f"JobHub connection error: {exc}")
        if jobs:
            labels = [f"{j.get('job_no')} — {j.get('job_name')} | {j.get('builder_client')}" for j in jobs]
            picked = st.sidebar.selectbox("Open a JobHub job", labels)
            selected_job = jobs[labels.index(picked)]
            if st.sidebar.button("Open linked job", type="primary", use_container_width=True):
                st.session_state["workspace_id"] = open_jobhub_workspace(selected_job)
                st.rerun()
        else:
            st.sidebar.caption("No JobHub jobs were found.")
    workspaces = lquery("SELECT * FROM workspaces ORDER BY updated_at DESC,id DESC")
    if workspaces:
        labels = [f"#{w['id']} · {w.get('job_no','')} — {w.get('job_name','')}" for w in workspaces]
        current_id = st.session_state.get("workspace_id")
        current_index = 0
        for i, w in enumerate(workspaces):
            if int(w["id"]) == int(current_id or -1):
                current_index = i
                break
        picked_local = st.sidebar.selectbox("Saved PlanReader workspaces", labels, index=current_index)
        selected_ws = workspaces[labels.index(picked_local)]
        if st.sidebar.button("Open saved workspace", use_container_width=True):
            st.session_state["workspace_id"] = int(selected_ws["id"])
            st.rerun()
    with st.sidebar.expander("Create standalone workspace"):
        default_job_no = next_jobhub_job_no(bridge)
        with st.form("create_workspace"):
            job_no = st.text_input("Job number (auto: next shared JobHub number)", value=default_job_no)
            job_name = st.text_input("Project name")
            builder = st.text_input("Builder / client")
            address = st.text_input("Site address")
            link_to_jobhub = st.checkbox("Create this job in the shared JobHub database", value=bool(bridge))
            create = st.form_submit_button("Create workspace", use_container_width=True)
        if create:
            if not job_name.strip():
                st.error("Project name is required.")
            else:
                workspace_id = create_standalone_workspace(job_no.strip(), job_name.strip(), builder.strip(), address.strip())
                if bridge and link_to_jobhub:
                    jobhub_id = create_linked_jobhub_job(bridge, job_no.strip(), job_name.strip(), address.strip())
                    if jobhub_id:
                        lexecute("UPDATE workspaces SET jobhub_job_id=? WHERE id=?", (int(jobhub_id), workspace_id))
                        st.sidebar.success(f"Linked to JobHub job #{jobhub_id}.")
                    else:
                        st.sidebar.error("Could not link to JobHub. The job number may already be taken.")
                st.session_state["workspace_id"] = workspace_id
                st.rerun()
    return st.session_state.get("workspace_id")


def dashboard_page(workspace: Dict[str, Any]) -> None:
    hero(workspace)
    docs = ldf("SELECT * FROM documents WHERE workspace_id=?", (workspace["id"],))
    pages = ldf("SELECT * FROM pages WHERE workspace_id=?", (workspace["id"],))
    takeoff = dataframe_for_takeoff(int(workspace["id"]))
    masses = ldf("SELECT * FROM model_masses WHERE workspace_id=?", (workspace["id"],))
    work = takeoff_work_rows(takeoff)
    cols = st.columns(6)
    cols[0].metric("Documents", len(docs))
    cols[1].metric("Drawing pages", len(pages))
    cols[2].metric("Take-off lines", len(work))
    cols[3].metric("Measured m²", f"{work.loc[work['unit'].eq('m²'),'quantity'].sum():,.1f}" if not work.empty else "0")
    cols[4].metric("Paint litres", f"{work['paint_litres'].sum():,.1f}" if not work.empty else "0")
    cols[5].metric("3D masses", len(masses))
    st.markdown("<div class='pb-card'>", unsafe_allow_html=True)
    st.subheader("Recommended workflow")
    st.write("1. Link or create a job → 2. Sync/upload every plan and specification → 3. Process drawings → 4. review the drawing register → 5. run AI and/or map measured zones → 6. review the take-off → 7. build the 3D model → 8. export or send the approved draft to JobHub.")
    st.markdown("<div class='pb-warning'><b>Accuracy rule:</b> the app labels geometry as Measured, Derived or Assumed. The model is not construction-grade BIM and must not be treated as exact until the estimator verifies dimensions against the current drawing issue.</div>", unsafe_allow_html=True)
    st.markdown("</div>", unsafe_allow_html=True)
    if not work.empty:
        section_summary = work.groupby("section", dropna=False).agg(quantity_rows=("id","count"), paint_litres=("paint_litres","sum"), labour_hours=("labour_hours","sum"), value_ex_gst=("value_ex_gst","sum")).reset_index()
        st.subheader("Take-off summary")
        st.dataframe(section_summary, use_container_width=True, hide_index=True)
    if not masses.empty or not ldf("SELECT * FROM mapped_zones WHERE workspace_id=?", (workspace["id"],)).empty:
        st.plotly_chart(build_3d_figure(int(workspace["id"])), use_container_width=True)


def page_thumbnail_bytes(path: str, max_w: int = 320) -> Optional[bytes]:
    try:
        p = Path(path)
        if not p.exists():
            return None
        img = Image.open(p)
        img = img.convert("RGB")
        ratio = max_w / float(img.width)
        if ratio < 1:
            img = img.resize((max_w, max(1, int(img.height * ratio))))
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        return buf.getvalue()
    except Exception:
        return None


def page_thumbnail(path: str, max_w: int = 320) -> Optional[bytes]:
    try:
        mtime = Path(path).stat().st_mtime
    except Exception:
        return None
    return _page_thumb(path, mtime, max_w)


@lru_cache(maxsize=512)
def _page_thumb(path: str, mtime: float, max_w: int) -> Optional[bytes]:
    try:
        p = Path(path)
        if not p.exists():
            return None
        img = Image.open(p)
        img = img.convert("RGB")
        ratio = max_w / float(img.width)
        if ratio < 1:
            img = img.resize((max_w, max(1, int(img.height * ratio))))
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        return buf.getvalue()
    except Exception:
        return None


def page_preview_picker(workspace_id: int, doc_ids: Sequence[int]) -> None:
    pages = lquery(
        "SELECT p.id,p.document_id,p.page_no,p.page_label,p.page_type,p.image_path,p.extracted_text,p.selected,d.file_name "
        "FROM pages p JOIN documents d ON d.id=p.document_id "
        f"WHERE p.workspace_id=? AND p.document_id IN ({','.join('?' for _ in doc_ids)}) ORDER BY d.id,p.page_no",
        (workspace_id, *doc_ids),
    )
    if not pages:
        st.info("No indexed pages to choose from yet.")
        return

    st.markdown("#### Choose pages to keep")
    st.caption("Tick only the pages you want rendered and shown. Unticked pages are ignored and never displayed.")
    keep_default = st.session_state.get("_pb_picker_keep_default")
    rows = []
    for pg in pages:
        has_img = bool(str(pg.get("image_path") or "").strip())
        excerpt = " ".join(str(pg.get("extracted_text") or "").split())[:70]
        if keep_default is None:
            keep = has_img
        elif keep_default:
            keep = True
        else:
            keep = False
        rows.append({
            "Keep": keep,
            "Page": int(pg["page_no"]),
            "File": str(pg.get("file_name") or ""),
            "Label": str(pg.get("page_label") or ""),
            "Type": str(pg.get("page_type") or ""),
            "Text": excerpt,
            "Rendered": "yes" if has_img else "no",
            "_pid": int(pg["id"]),
            "_doc": int(pg["document_id"]),
            "_no": int(pg["page_no"]),
            "_img": has_img,
        })
    edited = st.data_editor(
        pd.DataFrame(rows),
        hide_index=True,
        disabled=["Page", "File", "Label", "Type", "Text", "Rendered", "_pid", "_doc", "_no", "_img"],
        column_config={
            "Keep": st.column_config.CheckboxColumn("Keep"),
            "Rendered": st.column_config.TextColumn("Rendered", width="small"),
        },
        use_container_width=True,
    )
    c1, c2, c3, c4 = st.columns(4)
    if c1.button("Render & show kept pages", type="primary"):
        records = edited.to_dict("records")
        for row in records:
            lexecute("UPDATE pages SET selected=? WHERE id=?", (1 if row.get("Keep") else 0, row["_pid"]))
        to_render: Dict[int, List[int]] = {}
        for row in records:
            if row.get("Keep") and not row.get("_img"):
                to_render.setdefault(int(row["_doc"]), []).append(int(row["_no"]))
        rendered_msgs = []
        for doc_id, page_nos in to_render.items():
            try:
                count, msg = process_document(doc_id, page_ids=page_nos)
                rendered_msgs.append(f"doc {doc_id}: {count} page(s) — {msg}")
            except Exception as exc:
                rendered_msgs.append(f"doc {doc_id}: ERROR — {exc}")
        if rendered_msgs:
            st.code("\n".join(rendered_msgs))
        st.session_state.pop("pb_page_keep_editor", None)
        st.success("Selection saved; kept pages rendered.")
        st.rerun()
    if c2.button("Keep all"):
        st.session_state["_pb_picker_keep_default"] = True
        st.rerun()
    if c3.button("Keep none"):
        st.session_state["_pb_picker_keep_default"] = False
        st.rerun()
    if c4.button("Done"):
        st.session_state.pop("preview_doc_ids", None)
        st.session_state.pop("_pb_picker_keep_default", None)
        st.session_state.pop("pb_page_keep_editor", None)
        st.rerun()

    kept = [pg for pg in pages if bool(pg.get("selected")) and bool(str(pg.get("image_path") or "").strip())]
    if kept:
        st.markdown("#### Kept page previews")
        ncols = 4
        cols = st.columns(ncols)
        for i, pg in enumerate(kept):
            with cols[i % ncols]:
                thumb = page_thumbnail(str(pg.get("image_path") or ""))
                if thumb:
                    st.image(thumb, use_container_width=True)
                st.caption(f"p{int(pg['page_no'])} · {pg.get('page_type')}")


def project_documents_page(workspace: Dict[str, Any], bridge: Optional[JobHubBridge], user: Dict[str, Any]) -> None:
    hero(workspace)
    tabs = st.tabs(["Project details", "Linked documents", "Upload", "Process files", "File manager"])
    with tabs[0]:
        with st.form("project_details"):
            c1,c2 = st.columns(2)
            job_no = c1.text_input("Job number", value=str(workspace.get("job_no") or ""))
            job_name = c2.text_input("Project name", value=str(workspace.get("job_name") or ""))
            builder = c1.text_input("Builder / client", value=str(workspace.get("builder_client") or ""))
            address = c2.text_input("Site address", value=str(workspace.get("site_address") or ""))
            issue = c1.text_input("Drawing issue / revision", value=str(workspace.get("drawing_issue") or ""))
            estimator = c2.text_input("Estimator", value=str(workspace.get("estimator") or user.get("username") or ""))
            status = c1.selectbox("Take-off status", ["Draft","In review","Clarifications required","Measured","Approved"], index=max(0,["Draft","In review","Clarifications required","Measured","Approved"].index(str(workspace.get("status") or "Draft")) if str(workspace.get("status") or "Draft") in ["Draft","In review","Clarifications required","Measured","Approved"] else 0))
            summary = st.text_area("Executive summary", value=str(workspace.get("executive_summary") or ""), height=160)
            save = st.form_submit_button("Save project details", type="primary")
        if save:
            lexecute("UPDATE workspaces SET job_no=?,job_name=?,builder_client=?,site_address=?,drawing_issue=?,estimator=?,status=?,executive_summary=?,updated_at=? WHERE id=?", (job_no,job_name,builder,address,issue,estimator,status,summary,now_stamp(),workspace["id"]))
            st.success("Project details saved.")
            st.rerun()
    with tabs[1]:
        if not bridge or not workspace.get("jobhub_job_id"):
            st.info("This workspace is not linked to JobHub. Use Upload to add plans directly.")
        else:
            st.markdown("<div class='pb-note'>This searches common JobHub document/attachment tables. Database blobs, public URLs and locally reachable paths can be imported. A local path stored on a different Render service is not reachable; use shared object storage or upload it here.</div>", unsafe_allow_html=True)
            if st.button("Find and import every linked JobHub document", type="primary"):
                try:
                    records = discover_jobhub_document_records(bridge, int(workspace["jobhub_job_id"]))
                except Exception as exc:
                    records = None
                    st.warning(f"JobHub is unreachable right now ({exc}). No linked documents were changed; it will retry automatically.")
                if records:
                    messages=[]
                    success=0
                    for record in records:
                        if record.get("has_blob") and record.get("source_table") and record.get("record_id"):
                            try:
                                record["file_blob"] = bridge.fetch_document_blob(record["source_table"], int(record["record_id"]))
                            except Exception:
                                record["file_blob"] = None
                        ok,msg = copy_jobhub_document_to_workspace(record, int(workspace["id"]))
                        messages.append(msg)
                        success += int(ok)
                    st.success(f"Processed {len(records)} linked records; {success} are available in PlanReader.")
                    st.code("\n".join(messages))
                elif records is not None:
                    st.warning("No compatible linked-document records were found in the JobHub database.")
            linked = ldf("SELECT file_name,jobhub_table,jobhub_record_id,mime_type,uploaded_at FROM documents WHERE workspace_id=? AND source_type='JobHub linked document' ORDER BY id DESC", (workspace["id"],))
            st.dataframe(linked, use_container_width=True, hide_index=True)
    with tabs[2]:
        takeoff_import_panel(int(workspace["id"]), "doc_takeoff_import")
        uploads = st.file_uploader("Upload plans, specifications, schedules and images", type=["pdf","png","jpg","jpeg","webp","docx","xlsx","xls","csv","txt"], accept_multiple_files=True)
        category = st.selectbox("Document category", ["Plans","Specifications","Schedules","Addenda","Scope / tender documents","Site photos","Other"])
        mirror = st.checkbox("Record these uploads as linked PlanReader documents in the shared JobHub database", value=bool(bridge and workspace.get("jobhub_job_id")))
        total_bytes=sum((int(getattr(u,"size",0) or 0)) for u in uploads or [])
        if uploads and total_bytes>200*1024*1024:
            st.warning(f"Selected files total {total_bytes/1024/1024:.0f} MB. Uploads this large can exhaust this service's memory; split them into smaller batches.")
        if uploads and st.button("Save uploaded documents", type="primary"):
            added=0
            if bridge and mirror and workspace.get("jobhub_job_id"):
                ensure_planreader_document_table(bridge)
            new_ids=[]
            for upload in uploads:
                data = upload.getvalue()
                digest=sha256_bytes(data)
                if lquery("SELECT id FROM documents WHERE workspace_id=? AND sha256=?", (workspace["id"],digest)):
                    continue
                target=workspace_path(int(workspace["id"])) / "documents" / f"{digest[:12]}_{safe_name(upload.name)}"
                target.write_bytes(data)
                doc_id=lexecute("""INSERT INTO documents(workspace_id,source_type,jobhub_table,jobhub_record_id,file_name,mime_type,path,sha256,category,page_count,extracted_text,uploaded_at) VALUES(?,?,?,?,?,?,?,?,?,?,?,?)""", (workspace["id"],"PlanReader upload","","",upload.name,upload.type or mimetypes.guess_type(upload.name)[0] or "application/octet-stream",str(target),digest,category,0,"",now_stamp()))
                if bridge and mirror and workspace.get("jobhub_job_id"):
                    bridge.execute("INSERT INTO planreader_documents(job_id,file_name,mime_type,storage_path,source_app,uploaded_by,uploaded_at,notes) VALUES(?,?,?,?,?,?,?,?)", (workspace["jobhub_job_id"],upload.name,upload.type or "",str(target),"PlanReader",user.get("username",""),now_stamp(),"Stored on PlanReader service; metadata linked in JobHub database."))
                new_ids.append(doc_id)
                added+=1
            messages=[]
            if new_ids:
                progress=st.progress(0)
                status=st.empty()
            for i,doc_id in enumerate(new_ids):
                status.caption(f"Indexing document {i+1} of {len(new_ids)}…")
                try:
                    count,msg=index_document_pages(doc_id)
                    messages.append(f"#{doc_id}: {count} page(s) — {msg}")
                except Exception as exc:
                    messages.append(f"#{doc_id}: ERROR — {exc}")
                progress.progress(min((i+1)/len(new_ids),1.0))
            if new_ids:
                status.empty()
                progress.progress(1.0)
            if messages:
                st.code("\n".join(messages))
            st.success(f"Saved {added} new document(s).")
            if new_ids:
                st.session_state["preview_doc_ids"]=new_ids
                st.session_state["_pb_picker_keep_default"]=None
            st.rerun()
        if st.session_state.get("preview_doc_ids"):
            page_preview_picker(int(workspace["id"]), st.session_state["preview_doc_ids"])
    with tabs[3]:
        docs=ldf("SELECT id,file_name,category,page_count,source_type FROM documents WHERE workspace_id=? ORDER BY id", (workspace["id"],))
        if docs.empty:
            st.info("Upload or sync documents first.")
        else:
            st.dataframe(docs,use_container_width=True,hide_index=True)
            options={f"#{int(r.id)} — {r.file_name} ({int(r.page_count or 0)} pages)":int(r.id) for r in docs.itertuples()}
            selected=st.multiselect("Select documents",list(options.keys()),default=list(options.keys()))
            force=st.checkbox("Reprocess and replace existing rendered pages")
            if selected and st.button("Process selected documents", type="primary"):
                progress = st.progress(0)
                status = st.empty()
                messages = []
                selected_ids = [options[label] for label in selected]
                page_rows = ldf(
                    "SELECT id, page_count FROM documents WHERE id IN (%s)" % ",".join(["?"] * len(selected_ids)),
                    selected_ids,
                )
                page_by_doc = {int(r.id): int(r.page_count or 0) for r in page_rows.itertuples()}
                grand_total = max(1, sum(page_by_doc.values()))
                done = [0]

                def _render_progress(_completed, _total, page_no, doc_label):
                    done[0] += 1
                    progress.progress(min(done[0] / grand_total, 1.0))
                    status.caption(f"{doc_label}: rendering page {page_no}")

                for i, label in enumerate(selected):
                    estimated = page_by_doc.get(options[label]) or 1
                    try:
                        count, msg = process_document(
                            options[label],
                            force=force,
                            progress_cb=lambda c, t, p, l=label: _render_progress(c, t, p, l),
                        )
                        messages.append(f"{label}: {count} page(s) — {msg}")
                    except Exception as exc:
                        messages.append(f"{label}: ERROR — {exc}")
                        done[0] += estimated
                        progress.progress(min(done[0] / grand_total, 1.0))
                progress.progress(1.0)
                status.empty()
                st.session_state["pb_process_messages"] = messages
                st.rerun()
            if st.session_state.get("pb_process_messages"):
                st.code("\n".join(st.session_state["pb_process_messages"]))
    with tabs[4]:
        docs=ldf("SELECT id,file_name,source_type,category,page_count,uploaded_at,path FROM documents WHERE workspace_id=? ORDER BY id DESC", (workspace["id"],))
        st.dataframe(docs,use_container_width=True,hide_index=True)
        if not docs.empty:
            labels={f"#{int(r.id)} — {r.file_name}":int(r.id) for r in docs.itertuples()}
            selected=st.multiselect("Delete documents",list(labels.keys()))
            delete_files=st.checkbox("Delete stored physical files too",value=True)
            if selected and st.button("Delete selected",type="secondary"):
                for label in selected:
                    doc_id=labels[label]
                    row=lquery("SELECT path FROM documents WHERE id=?",(doc_id,))[0]
                    for page in lquery("SELECT image_path FROM pages WHERE document_id=?",(doc_id,)):
                        try: Path(str(page.get("image_path") or "")).unlink(missing_ok=True)
                        except Exception: pass
                    if delete_files:
                        try: Path(str(row.get("path") or "")).unlink(missing_ok=True)
                        except Exception: pass
                    lexecute("DELETE FROM documents WHERE id=?",(doc_id,))
                st.success("Selected documents removed from this workspace.")
                st.rerun()


def drawing_register_page(workspace: Dict[str, Any]) -> None:
    hero(workspace)
    target_register_id = st.session_state.pop("active_register_item_id", None)
    if target_register_id:
        st.info(f"🎯 Target clarification register item #{target_register_id} selected from Review & QA workspace.")
    pages=ldf("""SELECT p.id,p.page_label,p.page_type,p.scale_text,p.px_per_m,p.page_no,p.selected,d.file_name,p.image_path FROM pages p JOIN documents d ON d.id=p.document_id WHERE p.workspace_id=? ORDER BY d.id,p.page_no""",(workspace["id"],))
    if pages.empty:
        st.info("Process documents first.")
        return
    editable=pages[["id","page_label","page_type","scale_text","selected"]].copy()
    edited=st.data_editor(editable,use_container_width=True,hide_index=True,column_config={"id":st.column_config.NumberColumn(disabled=True),"page_type":st.column_config.SelectboxColumn(options=PAGE_TYPES),"selected":st.column_config.CheckboxColumn()},num_rows="fixed")
    if st.button("Save drawing register changes",type="primary"):
        for row in edited.to_dict("records"):
            lexecute("UPDATE pages SET page_label=?,page_type=?,scale_text=?,selected=? WHERE id=?",(row["page_label"],row["page_type"],row["scale_text"],1 if row["selected"] else 0,row["id"]))
        seed_drawing_register(int(workspace["id"]))
        st.success("Drawing register saved.")
        st.rerun()
    st.subheader("Drawing previews")
    labels=[f"#{int(r.id)} · {r.page_label} · {r.page_type}" for r in pages.itertuples()]
    chosen=st.selectbox("Preview page",labels)
    row=pages.iloc[labels.index(chosen)]
    if Path(str(row["image_path"])).exists():
        st.image(str(row["image_path"]),caption=f"{row['file_name']} · page {row['page_no']}",use_container_width=True)


def add_register_item_form(workspace_id:int, register_name:str, key_suffix:str="") -> None:
    with st.form(f"reg_form_{register_name}_{key_suffix}"):
        c1,c2=st.columns(2)
        item_no=c1.text_input("No. / code")
        title=c2.text_input("Title")
        detail=st.text_area("Detail")
        priority=c1.selectbox("Priority",["","Low","Medium","High","Critical"])
        source=c2.text_input("Source reference")
        status=c1.selectbox("Status",["Open","To review","Answered","Accepted","Excluded","Closed"])
        save=st.form_submit_button("Add item")
    if save:
        lexecute("INSERT INTO register_items(workspace_id,register_name,item_no,title,detail,priority,source_reference,status,created_at) VALUES(?,?,?,?,?,?,?,?,?)",(workspace_id,register_name,item_no,title,detail,priority,source,status,now_stamp()))
        st.success("Item added.")
        st.rerun()


def _run_with_progress(worker: Callable[[], Any], caption: str) -> Any:
    """Run ``worker`` while showing a live progress bar with elapsed time.

    The bar climbs asymptotically toward 90% while the call runs and jumps to
    100% only on success, so it never falsely claims completion. Returns the
    worker's value, or re-raises its exception.
    """
    progress = st.progress(0)
    status = st.empty()
    start = time.monotonic()
    holder: Dict[str, Any] = {}
    done = threading.Event()

    def _worker_thread():
        try:
            holder["value"] = worker()
        except Exception as exc:
            holder["error"] = exc
        finally:
            done.set()

    thread = threading.Thread(target=_worker_thread, daemon=True)
    thread.start()
    tick = 0
    while not done.is_set():
        tick += 1
        frac = min(0.05 + (1 - math.exp(-tick * 0.04)) * 0.85, 0.9)
        progress.progress(frac)
        status.caption(f"{caption}… {int(time.monotonic() - start)}s elapsed")
        time.sleep(0.2)
    thread.join()
    if "error" in holder:
        status.empty()
        raise holder["error"]
    progress.progress(1.0)
    status.empty()
    return holder["value"]


def subscription_takeoff_page(workspace: Dict[str, Any], session_api_key: str, ai_provider: str = "OpenAI") -> None:
    hero(workspace)
    target_row_id = st.session_state.pop("active_takeoff_row_id", None)
    if target_row_id:
        st.info(f"🎯 Target take-off row #{target_row_id} selected from Review & QA workspace.")
    tabs=st.tabs(["AI plan read","Take-off schedule","Scope registers","Door schedule","Source & basis","AI vs drawn"])
    with tabs[0]:
        pages=ldf("SELECT id,page_label,page_type,image_path,selected FROM pages WHERE workspace_id=? ORDER BY id",(workspace["id"],))
        if pages.empty:
            st.info("Process plan documents first.")
        else:
            options={f"#{int(r.id)} · {r.page_label} · {r.page_type}":int(r.id) for r in pages.itertuples()}
            default=[label for label,rid in options.items() if bool(pages.loc[pages['id'].eq(rid),'selected'].iloc[0])][:6]
            selected=st.multiselect("Pages to analyse",list(options.keys()),default=default)
            model=st.text_input("Model",value=default_ai_model(ai_provider),help="OpenAI uses the Responses API; Google Gemini uses the free-tier generateContent endpoint.")
            st.markdown("<div class='pb-note'>AI is used to organise evidence and draft quantities. It is instructed not to invent dimensions. Every result remains editable and requires estimator review.</div>",unsafe_allow_html=True)
            if st.button("Run subscription-method plan read",type="primary",disabled=not bool(resolve_ai_key(ai_provider,session_api_key))):
                try:
                    data=_run_with_progress(
                        lambda: run_ai_plan_read(int(workspace["id"]),[options[x] for x in selected],resolve_ai_key(ai_provider,session_api_key),model.strip() or default_ai_model(ai_provider),ai_provider),
                        f"Reading {len(selected)} drawing page(s) and building the take-off draft",
                    )
                    st.session_state["latest_ai_result"]=data
                    st.success("AI plan read completed. Review the preview, then import it.")
                except Exception as exc:
                    lexecute("INSERT INTO ai_runs(workspace_id,run_type,model,source_pages,status,response_json,error_message,created_at) VALUES(?,?,?,?,?,?,?,?)",(workspace["id"],"Plan take-off and 3D",model,json.dumps([options[x] for x in selected]),"Failed","",str(exc),now_stamp()))
                    st.error(_ai_error_hint(exc))
            data=st.session_state.get("latest_ai_result")
            if data:
                st.json(data,expanded=False)
                if st.button("Import this AI draft into the workspace",type="primary"):
                    counts=import_ai_result(int(workspace["id"]),data)
                    st.success(f"Imported {counts['takeoff']} take-off rows, {counts['registers']} register items, {counts['masses']} masses and {counts['openings']} openings.")
                    st.session_state.pop("latest_ai_result",None)
                    st.rerun()
            if not resolve_ai_key(ai_provider,session_api_key):
                st.warning("Configure the AI API key (OPENAI_API_KEY or GEMINI_API_KEY) or enter a session key in the sidebar to enable AI plan reading. Manual mapping and 3D modelling still work without it.")
    with tabs[1]:
        takeoff_import_panel(int(workspace["id"]))
        scale_issues=scale_gate_issues(int(workspace["id"]))
        if scale_issues:
            st.warning("**Scale gate:** these page(s) feed the take-off but have no calibrated scale yet — calibrate them in the **Plan Mapper → Scale** tab before treating quantities as measured: " + ", ".join(f"`{i['page_label']}`" for i in scale_issues[:8]))
        takeoff=ldf("SELECT * FROM takeoff_rows WHERE workspace_id=? ORDER BY id",(workspace["id"],))
        editor_cols=["id"]+TAKEOFF_COLUMNS+["row_role"]
        if takeoff.empty:
            takeoff=pd.DataFrame(columns=editor_cols)
        edited=st.data_editor(takeoff[editor_cols],use_container_width=True,hide_index=True,num_rows="dynamic",column_config={
            "id":st.column_config.NumberColumn(disabled=True),
            "substrate":st.column_config.SelectboxColumn(options=SUBSTRATES),
            "finish_system":st.column_config.SelectboxColumn(options=FINISH_SYSTEMS),
            "unit":st.column_config.SelectboxColumn(options=UNIT_OPTIONS),
            "quantity_status":st.column_config.SelectboxColumn(options=STATUS_OPTIONS),
            "inclusion_status":st.column_config.SelectboxColumn(options=INCLUSION_OPTIONS),
            "row_role":st.column_config.SelectboxColumn(options=["", "floor_area"],required=False,help="Mark a row as floor_area to record a per-level internal floor area (m²) measurement that drives floor-m² pricing. Floor-area rows are not priced themselves."),
        },height=560,key="takeoff_editor")
        st.caption("Rows marked **floor_area** hold each level's internal floor area (m²) — they appear in the schedule for reference and drive floor-m² pricing, but are not priced themselves. Set the pricing basis in **Settings → Project-level settings**.")
        gate_confirmed=not scale_issues
        if scale_issues:
            gate_confirmed=st.checkbox("I will calibrate the affected page scales before using these quantities",key=f"scale_gate_{int(workspace['id'])}")
        c1,c2=st.columns(2)
        if c1.button("Save take-off schedule",type="primary",disabled=not gate_confirmed,use_container_width=True):
            lexecute("DELETE FROM takeoff_rows WHERE workspace_id=?",(workspace["id"],))
            for row in edited.to_dict("records"):
                if not any(str(row.get(c) or "").strip() for c in ["section","element","location","source_reference"]):
                    continue
                if not to_float(row.get("rate_per_unit")):
                    row["rate_per_unit"] = default_rate_for(row.get("substrate"),row.get("element"),row.get("finish_system"),row.get("unit"))
                row_role=str(row.get("row_role") or "").strip()
                if row_role not in {"", "floor_area"}:
                    row_role=""
                values=[row.get(col,"") for col in TAKEOFF_COLUMNS]
                lexecute("""INSERT INTO takeoff_rows(workspace_id,section,element,location,substrate,finish_system,quantity,unit,quantity_status,source_page,source_reference,inclusion_status,coats,coverage_m2_per_litre,productivity_m2_per_hour,rate_per_unit,confidence,notes,row_role,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",(workspace["id"],*values,row_role,now_stamp(),now_stamp()))
            st.success("Take-off schedule saved.")
            st.rerun()
        if c2.button("Apply default rates to all rows",use_container_width=True):
            rows=ldf("SELECT id,substrate,element,finish_system,unit,rate_per_unit FROM takeoff_rows WHERE workspace_id=?",(workspace["id"],))
            for r in rows.itertuples():
                if to_float(r.rate_per_unit)>0:
                    continue
                rate=default_rate_for(r.substrate,r.element,r.finish_system,r.unit)
                lexecute("UPDATE takeoff_rows SET rate_per_unit=?,updated_at=? WHERE id=?",(rate,now_stamp(),r.id))
            st.success("Default rates applied. Values are editable in the schedule.")
            st.rerun()
        if c2.button("Add standard empty scope rows",use_container_width=True):
            seeds=[
                ("Internal","Walls","All internal areas","Plasterboard","Low sheen wall system",0,"m²","To measure","","","INCLUSION",3,12,8,default_rate_for("Plasterboard","Walls","Low sheen wall system","m²"),"To review","Net of tiles, glazing and joinery.",""),
                ("Internal","Floor area","All internal areas","Other","N/A",0,"m²","To measure","","","INCLUSION",0,12,0,0.0,"To review","Internal floor area measurement row — drives floor-m² pricing when selected in Settings; not priced itself.","floor_area"),
                ("Internal","Ceilings","Flushset ceilings","Plasterboard","Ceiling flat",0,"m²","To measure","","","INCLUSION",3,12,10,default_rate_for("Plasterboard","Ceilings","Ceiling flat","m²"),"To review","Exclude grid ceiling tiles.",""),
                ("Internal","Doors","Painted door leaves","Timber door","Semi-gloss / enamel",0,"No.","To measure","","","INCLUSION",3,10,1,default_rate_for("Timber door","Doors","Semi-gloss / enamel","No."),"To review","Confirm leaf and frame finishes.",""),
                ("External","External walls / cladding","All elevations","Fibre cement","Exterior acrylic",0,"m²","To measure","","","INCLUSION",3,10,6,default_rate_for("Fibre cement","External walls / cladding","Exterior acrylic","m²"),"To review","Net of glazing, signage and prefinished cladding.",""),
                ("External","Steel / columns","Canopies and exposed steel","Structural steel","Metal primer + topcoats",0,"m²","To measure","","","PROVISIONAL",3,9,4,default_rate_for("Structural steel","Steel / columns","Metal primer + topcoats","m²"),"To review","Confirm site-painted versus factory finish.",""),
            ]
            for seed in seeds:
                lexecute("""INSERT INTO takeoff_rows(workspace_id,section,element,location,substrate,finish_system,quantity,unit,quantity_status,source_page,source_reference,inclusion_status,coats,coverage_m2_per_litre,productivity_m2_per_hour,rate_per_unit,confidence,notes,row_role,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",(workspace["id"],*seed,now_stamp(),now_stamp()))
            st.success("Standard rows added.")
            st.rerun()
        c3, c4 = st.columns(2)
        clear_confirm = c3.checkbox(
            "I understand this deletes every take-off row (quantities, litres, hours) for this job",
            key=f"clear_takeoff_confirm_{int(workspace['id'])}",
        )
        if c4.button("Clear take-off data", type="secondary", disabled=not clear_confirm, use_container_width=True):
            lexecute("DELETE FROM takeoff_rows WHERE workspace_id=?", (workspace["id"],))
            st.success("Take-off data cleared.")
            st.rerun()
        with st.expander("Copy rows to another level", expanded=False):
            st.caption("Duplicates the selected rows for a new level. Locations are rewritten (e.g. 'Level 1 · all walls' → 'Level 2 · all walls'); rows that already exist for that level are skipped.")
            copy_pool=ldf("SELECT id,section,element,location,unit FROM takeoff_rows WHERE workspace_id=? ORDER BY section,element,location",(workspace["id"],))
            if copy_pool.empty:
                st.caption("No take-off rows to copy yet.")
            else:
                copy_options={f"#{int(r.id)} · {r.section} · {r.element} · {r.location} ({r.unit})":int(r.id) for r in copy_pool.itertuples()}
                cc1,cc2,cc3=st.columns([.6,.2,.2])
                picked=cc1.multiselect("Rows to copy",list(copy_options.keys()))
                target_level=cc2.selectbox("Target level",["Level 1","Ground","Level 2","Level 3","Level 4","Level 5","Roof"],key=f"copy_level_{int(workspace['id'])}")
                if cc3.button("Copy rows",type="primary",disabled=not picked):
                    n=copy_takeoff_rows_to_level(int(workspace["id"]),[copy_options[x] for x in picked],str(target_level))
                    st.success(f"Copied {n} row(s) to {target_level}." if n else "No new rows created (they may already exist for that level).")
                    st.rerun()
    with tabs[2]:
        names=["inclusions","exclusions","clarifications","assumptions","rfis","colour_finish_schedule","access_constraints","risks"]
        selected_reg=st.selectbox("Register",names,format_func=lambda x:x.replace("_"," ").title())
        frame=ldf("SELECT id,item_no,title,detail,priority,source_reference,status FROM register_items WHERE workspace_id=? AND register_name=? ORDER BY id",(workspace["id"],selected_reg))
        st.dataframe(frame,use_container_width=True,hide_index=True)
        add_register_item_form(int(workspace["id"]),selected_reg,"scope")
        if not frame.empty:
            choices={f"#{int(r.id)} · {r.title}":int(r.id) for r in frame.itertuples()}
            delete=st.multiselect("Delete selected register items",list(choices.keys()))
            if delete and st.button("Delete register items"):
                for item in delete: lexecute("DELETE FROM register_items WHERE id=?",(choices[item],))
                st.rerun()
    with tabs[3]:
        frame=ldf("SELECT id,item_no,title,detail,priority,source_reference,status FROM register_items WHERE workspace_id=? AND register_name='door_schedule' ORDER BY id",(workspace["id"],))
        st.dataframe(frame,use_container_width=True,hide_index=True)
        add_register_item_form(int(workspace["id"]),"door_schedule","doors")
    with tabs[4]:
        source=ldf("SELECT id,item_no,title,detail,priority,source_reference,status FROM register_items WHERE workspace_id=? AND register_name='source_basis' ORDER BY id",(workspace["id"],))
        st.dataframe(source,use_container_width=True,hide_index=True)
        add_register_item_form(int(workspace["id"]),"source_basis","source")
        if st.button("Generate source/basis rows from drawing register"):
            pages=lquery("""SELECT p.page_label,p.page_type,p.scale_text,p.page_no,d.file_name FROM pages p JOIN documents d ON d.id=p.document_id WHERE p.workspace_id=? ORDER BY d.id,p.page_no""",(workspace["id"],))
            for p in pages:
                ref=f"{p['file_name']} p{p['page_no']}"
                exists=lquery("SELECT id FROM register_items WHERE workspace_id=? AND register_name='source_basis' AND source_reference=?",(workspace["id"],ref))
                if not exists:
                    lexecute("INSERT INTO register_items(workspace_id,register_name,item_no,title,detail,priority,source_reference,status,created_at) VALUES(?,?,?,?,?,?,?,?,?)",(workspace["id"],"source_basis",p.get("page_label",""),p.get("page_type",""),f"Scale: {p.get('scale_text') or 'not confirmed'}","",ref,"Used / review",now_stamp()))
            st.rerun()
    with tabs[5]:
        st.subheader("AI vs drawn reconciliation")
        st.markdown("<div class='pb-note'>Compares AI-drafted quantities with quantities drawn on the plan (and JobHub imports). Rows are grouped by scope line so differences stand out before the take-off is priced. This does not change data — it only reports it.</div>",unsafe_allow_html=True)
        reco=reconcile_ai_vs_drawn(int(workspace["id"]))
        if reco.empty:
            st.info("Build the take-off first, then import an AI draft and/or draw measurements to compare them.")
        else:
            st.dataframe(reco,use_container_width=True,hide_index=True,height=420)
            diffs=reco[reco["status"].isin(["Difference","AI only (not drawn)","Drawn only (no AI draft)"])]
            if diffs.empty:
                st.success("No material differences between AI and drawn quantities.")
            else:
                st.warning(f"**{len(diffs)} line item(s) need review** — check the variance column and confirm which basis to price.")
                st.dataframe(diffs,use_container_width=True,hide_index=True)


def overlay_image(page:Dict[str,Any],zones:List[Dict[str,Any]]) -> Optional[Image.Image]:
    path=Path(str(page.get("image_path") or ""))
    if not path.exists(): return None
    img=Image.open(path).convert("RGBA")
    draw=ImageDraw.Draw(img,"RGBA")
    for zone in zones:
        x,y,w,h=[to_float(zone.get(k)) for k in ["x_px","y_px","w_px","h_px"]]
        draw.rectangle((x,y,x+w,y+h),fill=(215,162,27,55),outline=(179,58,58,255),width=4)
        draw.text((x+5,y+5),str(zone.get("name") or zone.get("id")),fill=(20,20,20,255))
    return img


def plan_mapper_page(workspace:Dict[str,Any]) -> None:
    hero(workspace)
    target_page_id = st.session_state.pop("active_page_id", None)
    if target_page_id:
        st.info(f"🎯 Target drawing page #{target_page_id} selected from Review & QA workspace.")
    pages=ldf("SELECT * FROM pages WHERE workspace_id=? ORDER BY id",(workspace["id"],))
    if pages.empty:
        st.info("Process drawings first.")
        return
    labels=[f"#{int(r.id)} · {r.page_label} · {r.page_type}" for r in pages.itertuples()]
    chosen=st.selectbox("Drawing page",labels)
    page=pages.iloc[labels.index(chosen)].to_dict()
    zones=lquery("SELECT * FROM mapped_zones WHERE page_id=? ORDER BY id",(int(page["id"]),))
    tab0,tab1,tab2,tab3=st.tabs(["Draw measurements","Scale","Map zone","Saved zones"])
    with tab0:
        pxpm=to_float(page.get("px_per_m"))
        st.markdown("### Draw take-off measurements on the plan")
        st.caption("Pick a take-off row, choose **Line** (walls, doors, frames, skirting — click two points) or **Outline** (building footprint, ceilings — click each corner then double-click to close), then click directly on the plan. Drawn lengths and areas are measured from the saved drawing scale.")
        detected=auto_detect_scale(page)
        if not pxpm and detected:
            c1,c2=st.columns([.72,.28])
            c1.write(f"Detected drawing scale **{detected['source']}** → ≈ **{detected['px_per_m']:.1f} px/m**.")
            if c2.button("Use detected scale",key=f"det_scale_{page['id']}"):
                lexecute("UPDATE pages SET px_per_m=? WHERE id=?",(detected["px_per_m"],page["id"]))
                st.rerun()
        elif not pxpm:
            st.warning("No scale for this page yet — calibrate it in the **Scale** tab (or confirm a detected scale above) so line lengths and areas are measured in real units.")
        store_key=f"ml_store_{int(page['id'])}"
        widget_key=f"ml_{int(page['id'])}"
        rev_key=f"mlrev_{int(page['id'])}"
        with st.expander("Quick add a take-off row to draw (e.g. Floor plan by area)",expanded=False):
            qa_cols=st.columns([.2,.26,.22,.18,.14])
            qa_level=qa_cols[0].selectbox("Level",["Level 1","Ground","Level 2","Level 3","Level 4","Roof"],key=f"qa_level_{page['id']}")
            qa_element=qa_cols[1].selectbox("Element",["Floor plan","Walls","Ceilings","Skirtings","Doors","Frames","External walls / cladding","Steel / columns"],key=f"qa_el_{page['id']}")
            qa_loc=qa_cols[2].text_input("Location / description",value=f"{qa_level} · all internal areas",key=f"qa_loc_{page['id']}")
            qa_sub=qa_cols[3].selectbox("Substrate",SUBSTRATES,index=SUBSTRATES.index("Render") if "Render" in SUBSTRATES else 0,key=f"qa_sub_{page['id']}")
            qa_fin=qa_cols[4].selectbox("Finish system",FINISH_SYSTEMS,index=0,key=f"qa_fin_{page['id']}")
            qa_unit = "m²" if qa_element in {"Floor plan","Walls","Ceilings","External walls / cladding"} else "lm"
            qa_section = "Internal" if qa_element not in {"External walls / cladding","Steel / columns"} else "External"
            qa_rate = default_rate_for(qa_sub,qa_element,qa_fin,qa_unit)
            st.caption(f"Will add: **{qa_section} · {qa_element} · {qa_loc}** on **{qa_sub}** with finish **{qa_fin}** ({qa_unit}). Default rate ${qa_rate:.2f}/{qa_unit} — editable in the take-off schedule.")
            if st.button("Add row and select it for drawing",key=f"qa_add_{page['id']}"):
                new_id=lexecute("""INSERT INTO takeoff_rows(workspace_id,section,element,location,substrate,finish_system,quantity,unit,quantity_status,source_page,source_reference,inclusion_status,coats,coverage_m2_per_litre,productivity_m2_per_hour,rate_per_unit,confidence,notes,row_role,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",(int(workspace["id"]),qa_section,qa_element,qa_loc,qa_sub,qa_fin,0,qa_unit,"To measure",page.get("page_label",""),"","INCLUSION",2 if qa_unit=="lm" else 3,12,8,qa_rate,"To review","Quick-added from Draw measurements tab.","",now_stamp(),now_stamp()))
                st.session_state[f"ml_active_{int(page['id'])}"]=int(new_id)
                st.session_state[rev_key]=int(st.session_state.get(rev_key,0))+1
                st.rerun()
        pending=st.session_state.get(store_key)
        if pending is None:
            pending=lquery("SELECT id,takeoff_row_id,label,unit,colour,kind,x1,y1,x2,y2,points,length_m,area_m2,perimeter_m,quantity_status,moved,notes FROM measurement_lines WHERE page_id=? ORDER BY id",(int(page["id"]),))
            if pending:
                pending=[dict(p) for p in pending]
        comp_lines=[]
        for i,ln in enumerate(pending or []):
            comp_lines.append({
                "id": str(ln.get("id") or f"ln{i}"),
                "takeoff_row_id": ln.get("takeoff_row_id"),
                "label": ln.get("label"), "unit": ln.get("unit"), "colour": ln.get("colour"),
                "kind": ln.get("kind") or "line",
                "x1": ln.get("x1"), "y1": ln.get("y1"), "x2": ln.get("x2"), "y2": ln.get("y2"),
                "points": ln.get("points") or [],
                "length_m": ln.get("length_m"), "area_m2": ln.get("area_m2"), "perimeter_m": ln.get("perimeter_m"),
                "quantity_status": ln.get("quantity_status"), "moved": ln.get("moved"), "notes": ln.get("notes"),
            })
        mapper_rows=takeoff_rows_for_mapper(int(workspace["id"]))
        if not mapper_rows:
            st.warning("No drawable take-off rows yet. Create take-off rows with a lineal (m) or area (m²) unit, then come back here to draw them on the plan.")
        else:
            st.caption(f"{len(mapper_rows)} drawable take-off rows: " + "; ".join(f"{r['label']} · {r['unit']}" for r in mapper_rows[:6]) + (" …" if len(mapper_rows) > 6 else ""))
        c1,c2,c3,c4=st.columns(4)
        if c1.button("Clear current drawing (un-saved)",key=f"cleardraw_{page['id']}",use_container_width=True):
            st.session_state.pop(store_key,None)
            st.session_state[rev_key]=int(st.session_state.get(rev_key,0))+1
            st.rerun()
        if c2.button("Delete saved shapes for this page",key=f"delsave_{page['id']}",use_container_width=True):
            lexecute("DELETE FROM measurement_lines WHERE page_id=?",(int(page["id"]),))
            st.session_state.pop(store_key,None)
            st.session_state[rev_key]=int(st.session_state.get(rev_key,0))+1
            st.rerun()
        if c3.button("Add a placeholder shape per row",key=f"auto_{page['id']}",use_container_width=True):
            new_lines=auto_map_measurements(int(workspace["id"]),int(page["id"]),pxpm)
            st.session_state[store_key]=new_lines
            st.session_state[rev_key]=int(st.session_state.get(rev_key,0))+1
            st.rerun()
        if c4.button("Auto-detect building envelope",key=f"autodet_{page['id']}",use_container_width=True):
            det_level=st.session_state.get(f"qa_level_{int(page['id'])}","Level 1")
            auto_lines=auto_detect_envelope_shapes(int(workspace["id"]),page,pxpm,str(det_level))
            if auto_lines:
                st.session_state[store_key]=auto_lines
                st.session_state[rev_key]=int(st.session_state.get(rev_key,0))+1
                st.rerun()
            else:
                st.warning("No building envelope could be detected on this page — draw the external outline manually, or check the page image.")
        st.caption("**Auto-detect** outlines the outer building envelope on this page, creating an **External walls / cladding** box (area = perimeter × 2.7 m wall height) and a **Floor plan · {level}** box (internal floor area, the flat-rate basis for internal works incl. paint). Select a box, then drag a corner to reshape or drag the box to move it — no auto-detect is perfect, so adjust before saving.")
        path=Path(str(page.get("image_path") or ""))
        if plan_line_editor is not None and path.exists():
            active_row_id=st.session_state.get(f"ml_active_{int(page['id'])}")
            result=plan_line_editor(path.read_bytes(),comp_lines,mapper_rows,pxpm,int(st.session_state.get(rev_key,0)),key=widget_key,height=760,active_row_id=active_row_id)
            if result is not None:
                st.session_state[store_key]=result
        else:
            st.error("The interactive line editor is unavailable in this environment.")
        if st.button("Save drawn measurements",type="primary",key=f"save_{page['id']}",disabled=pxpm<=0):
            outcome=save_measurement_lines(int(workspace["id"]),int(page["id"]),st.session_state.get(store_key) or comp_lines)
            st.success(f"Saved {outcome['saved']} drawn shape(s); {outcome['synced']} take-off quantities updated to their mapped measurements.")
            st.rerun()
        if pxpm<=0:
            st.error("**Scale gate:** this page has no calibrated scale, so drawn lengths and areas cannot be converted to real units. Set the scale in the **Scale** tab, then save again.")
        saved_lines=lquery("SELECT * FROM measurement_lines WHERE page_id=? ORDER BY id",(int(page["id"]),))
        if saved_lines:
            st.markdown("#### Saved overlay")
            overlay=render_measurement_overlay(page,saved_lines)
            if overlay: st.image(overlay,use_container_width=True)
        if comp_lines:
            st.markdown("#### Legend")
            lcols=st.columns(2)
            for i,ln in enumerate(comp_lines[:12]):
                with lcols[i%2]:
                    unit_note=f"{ln.get('unit') or ''}" + (f" · {ln.get('length_m')} m" if float(ln.get('length_m') or 0)>0 and (ln.get('kind') or 'line')=='line' else "") + (f" · {ln.get('area_m2')} m²" if float(ln.get('area_m2') or 0)>0 else "")
                    st.markdown(f"<span style='display:inline-block;width:18px;height:4px;background:{ln.get('colour')};vertical-align:middle;margin-right:6px'></span>{ln.get('label') or 'Measurement'} · {unit_note}",unsafe_allow_html=True)
        with st.expander("Substrate coverage check — which rows have been drawn?",expanded=False):
            shape_counts={}
            for ln in (comp_lines or []):
                rid=ln.get("takeoff_row_id")
                if rid is not None:
                    shape_counts[int(rid)]=shape_counts.get(int(rid),0)+1
            if not mapper_rows:
                st.caption("No drawable take-off rows yet.")
            else:
                missing=[r for r in mapper_rows if shape_counts.get(int(r["id"]),0)==0]
                covered=[r for r in mapper_rows if shape_counts.get(int(r["id"]),0)>0]
                if missing:
                    st.warning(f"**{len(missing)} row(s) have nothing drawn on this page yet** — check these substrates before saving.")
                else:
                    st.success(f"All {len(mapper_rows)} drawable row(s) have at least one shape drawn on this page.")
                for r in mapper_rows:
                    n=shape_counts.get(int(r["id"]),0)
                    flag=" <span style='color:#b42318;font-weight:700'>— NOT DRAWN</span>" if n==0 else ""
                    qty=f" · measured {r['quantity']:.2f}" if r["quantity"] else ""
                    st.markdown(f"<span style='display:inline-block;width:12px;height:12px;border-radius:3px;background:{r['colour']};vertical-align:middle;margin-right:6px'></span>{r['label']} · {r['unit']} · {n} shape(s){qty}{flag}",unsafe_allow_html=True)
    with tab1:
        st.write(f"Drawing scale text: **{page.get('scale_text') or 'not entered'}**")
        c1,c2=st.columns(2)
        known_m=c1.number_input("Known drawn distance (metres)",min_value=0.01,value=10.0,step=.1)
        pixel_distance=c2.number_input("Measured pixel distance",min_value=1.0,value=float((page.get("px_per_m") or 100)*known_m),step=1.0)
        if st.button("Save page calibration",type="primary"):
            pxpm=pixel_distance/known_m
            lexecute("UPDATE pages SET px_per_m=? WHERE id=?",(pxpm,page["id"]))
            st.success(f"Saved {pxpm:.2f} pixels per metre.")
            st.rerun()
        st.caption("Use a known dimension line from the drawing. Accurate scale calibration is essential before treating mapped areas as measured.")
    with tab2:
        image_path=Path(str(page.get("image_path") or ""))
        img=Image.open(image_path) if image_path.exists() else None
        pxpm=to_float(page.get("px_per_m"))
        if not img:
            st.error("Rendered page image is missing.")
        else:
            c_form,c_draw=st.columns([.34,.66])
            drawn_rect=None
            with c_draw:
                st.caption("Draw a rectangle around a room footprint, wall/elevation section or paintable zone. Rectangle mapping is intentionally simple and reviewable.")
                if CANVAS_AVAILABLE:
                    max_width=900
                    display_scale=min(1.0,max_width/img.width)
                    display=img.resize((int(img.width*display_scale),int(img.height*display_scale)))
                    canvas=st_canvas(fill_color="rgba(215,162,27,0.25)",stroke_width=3,stroke_color="#B33A3A",background_image=display,update_streamlit=True,height=display.height,width=display.width,drawing_mode="rect",key=f"canvas_{page['id']}")
                    if canvas.json_data and canvas.json_data.get("objects"):
                        obj=canvas.json_data["objects"][-1]
                        drawn_rect=(to_float(obj.get("left"))/display_scale,to_float(obj.get("top"))/display_scale,to_float(obj.get("width"))*to_float(obj.get("scaleX"),1)/display_scale,to_float(obj.get("height"))*to_float(obj.get("scaleY"),1)/display_scale)
                else:
                    overlay=overlay_image(page,zones)
                    st.image(overlay if overlay else img,use_container_width=True)
                    st.info("Drawing canvas is unavailable. Enter rectangle coordinates manually.")
            with c_form:
                name=st.text_input("Zone name",value=f"{page.get('page_label')} zone")
                view_type=st.selectbox("Geometry type",["Building footprint","Room footprint","Floor plan","Elevation area","Ceiling area","Soffit / canopy","Other"])
                if drawn_rect:
                    dx,dy,dw,dh=drawn_rect
                else:
                    dx,dy,dw,dh=0.0,0.0,min(500.0,float(img.width)),min(300.0,float(img.height))
                if pxpm>0:
                    xm=st.number_input("X position (m)",min_value=0.0,value=float(dx/pxpm),step=0.1,help="Distance in real metres from the page's left edge to the rectangle's left edge, converted from the calibrated scale.")
                    ym=st.number_input("Y position (m)",min_value=0.0,value=float(dy/pxpm),step=0.1,help="Distance in real metres from the top of the page to the rectangle's top edge.")
                    wm=st.number_input("Width (m)",min_value=0.01,value=max(0.01,float(dw/pxpm)),step=0.1,help="Rectangle width in real metres.")
                    hm=st.number_input("Height (m)",min_value=0.01,value=max(0.01,float(dh/pxpm)),step=0.1,help="Rectangle height in real metres.")
                    x,y,w,h=xm*pxpm,ym*pxpm,wm*pxpm,hm*pxpm
                    st.caption("Entered in real metres using the saved page calibration. Raw image pixels are stored internally.")
                else:
                    x=st.number_input("X position (pixels)",min_value=0.0,value=float(dx),step=1.0,help="Horizontal position in raw image pixels. Calibrate the scale (tab 1) to enter real metres instead.")
                    y=st.number_input("Y position (pixels)",min_value=0.0,value=float(dy),step=1.0,help="Vertical position in raw image pixels. Calibrate the scale (tab 1) to enter real metres instead.")
                    w=st.number_input("Width (pixels)",min_value=1.0,value=max(1.0,float(dw)),step=1.0,help="Rectangle width in raw image pixels.")
                    h=st.number_input("Height (pixels)",min_value=1.0,value=max(1.0,float(dh)),step=1.0,help="Rectangle height in raw image pixels.")
                    st.caption("Calibrate the page scale (tab 1) to enter and save mapped zones in real metres.")
                wall_height=st.number_input("Wall / extrusion height (m)",min_value=.1,value=2.7,step=.1,help="Height used to extrude this zone into a 3D building mass.")
                substrate=st.selectbox("Substrate",SUBSTRATES,help="Surface being painted; drives the default painting rate.")
                finish=st.selectbox("Finish system",FINISH_SYSTEMS,help="Coating system (render finish, etc.); used with substrate to pick the default rate.")
                qstatus=st.selectbox("Quantity status",STATUS_OPTIONS,index=0 if pxpm>0 else 2,help="Measured (calibrated scale), Estimated, or To review before it enters the take-off.")
                source=st.text_input("Source reference",value=f"{page.get('page_label')} · {page.get('page_type')}")
                if pxpm>0:
                    if view_type in {"Building footprint","Room footprint","Floor plan"}:
                        area=(w/pxpm)*(h/pxpm)
                    else:
                        area=(w/pxpm)*(h/pxpm)
                    st.metric("Calculated rectangle area",f"{area:,.2f} m²")
                else:
                    area=0
                    st.warning("Calibrate the page before saving a measured quantity.")
                if st.button("Save mapped zone",type="primary",use_container_width=True):
                    lexecute("""INSERT INTO mapped_zones(workspace_id,page_id,name,view_type,polygon_json,x_px,y_px,w_px,h_px,px_per_m,wall_height_m,area_m2,substrate,finish_system,quantity_status,source_reference,created_at) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",(workspace["id"],page["id"],name,view_type,json.dumps([[x,y],[x+w,y],[x+w,y+h],[x,y+h]]),x,y,w,h,pxpm,wall_height,area,substrate,finish,qstatus,source,now_stamp()))
                    st.success("Mapped zone saved.")
                    st.rerun()
    with tab3:
        frame=ldf("SELECT id,name,view_type,area_m2,wall_height_m,substrate,finish_system,quantity_status,source_reference FROM mapped_zones WHERE page_id=? ORDER BY id",(page["id"],))
        st.dataframe(frame,use_container_width=True,hide_index=True)
        if pxpm>0:
            st.caption("Area shown in m² from the saved page calibration. Zone positions are stored internally in pixels and converted to metres automatically for take-off and 3D masses.")
        overlay=overlay_image(page,zones)
        if overlay: st.image(overlay,use_container_width=True)
        if not frame.empty:
            choices={f"#{int(r.id)} · {r.name}":int(r.id) for r in frame.itertuples()}
            selected=st.multiselect("Zones",list(choices.keys()))
            c1,c2=st.columns(2)
            if selected and c1.button("Add selected zones to take-off"):
                for label in selected:
                    z=lquery("SELECT * FROM mapped_zones WHERE id=?",(choices[label],))[0]
                    lexecute("""INSERT INTO takeoff_rows(workspace_id,section,element,location,substrate,finish_system,quantity,unit,quantity_status,source_page,source_reference,inclusion_status,coats,coverage_m2_per_litre,productivity_m2_per_hour,rate_per_unit,confidence,notes,row_role,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",(workspace["id"],"Mapped drawing","Mapped zone",z.get("name",""),z.get("substrate",""),z.get("finish_system",""),z.get("area_m2",0),"m²",z.get("quantity_status","Measured"),page.get("page_label",""),z.get("source_reference",""),"INCLUSION",3,12,8,default_rate_for(z.get("substrate",""),"Mapped zone",z.get("finish_system",""),"m²"),"Measured" if pxpm>0 else "To review","Rectangle mapped in PlanReader.","",now_stamp(),now_stamp()))
                st.success("Take-off rows added.")
            if selected and c2.button("Create conceptual 3D masses from selected zones"):
                for label in selected:
                    z=lquery("SELECT * FROM mapped_zones WHERE id=?",(choices[label],))[0]
                    p=to_float(z.get("px_per_m"))
                    if p<=0: continue
                    lexecute("""INSERT INTO model_masses(workspace_id,label,level_name,x,y,z,width,depth,height,finish,source_reference,confidence,notes,created_at) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",(workspace["id"],z.get("name","Zone"),"Ground",to_float(z.get("x_px"))/p,to_float(z.get("y_px"))/p,0,to_float(z.get("w_px"))/p,to_float(z.get("h_px"))/p,to_float(z.get("wall_height_m"),2.7),z.get("finish_system",""),z.get("source_reference",""),"Derived","Derived from calibrated rectangular mapped zone.",now_stamp()))
                st.success("3D masses created.")
                st.rerun()
            if selected and st.button("Delete selected zones"):
                for label in selected: lexecute("DELETE FROM mapped_zones WHERE id=?",(choices[label],))
                st.rerun()


def model_3d_page(workspace:Dict[str,Any], session_api_key: str = "", ai_provider: str = "OpenAI") -> None:
    hero(workspace)
    tabs=st.tabs(["Interactive model","Building masses","Doors & windows","Render / artist's impression","Model exports"])
    with tabs[0]:
        masses=ldf("SELECT * FROM model_masses WHERE workspace_id=?",(workspace["id"],))
        zones=ldf("SELECT * FROM mapped_zones WHERE workspace_id=?",(workspace["id"],))
        if masses.empty and zones.empty:
            st.info("Create measured/derived masses with the Plan Mapper or add a mass manually.")
        else:
            st.plotly_chart(build_3d_figure(int(workspace["id"])),use_container_width=True)
            measured=0 if masses.empty else int(masses["confidence"].astype(str).str.lower().isin(["measured","verified"]).sum())
            assumed=0 if masses.empty else int(masses["confidence"].astype(str).str.lower().eq("assumed").sum())
            c1,c2,c3=st.columns(3)
            c1.metric("Model masses",len(masses))
            c2.metric("Measured / verified",measured)
            c3.metric("Assumed",assumed)
            if assumed:
                st.warning("The model contains assumed geometry. Verify it against dimensions/elevations before relying on the render.")
    with tabs[1]:
        masses=ldf("SELECT id,label,level_name,x,y,z,width,depth,height,finish,source_reference,confidence,notes FROM model_masses WHERE workspace_id=? ORDER BY id",(workspace["id"],))
        if masses.empty: masses=pd.DataFrame(columns=["id","label","level_name","x","y","z","width","depth","height","finish","source_reference","confidence","notes"])
        edited=st.data_editor(masses,use_container_width=True,hide_index=True,num_rows="dynamic",column_config={"id":st.column_config.NumberColumn(disabled=True),"confidence":st.column_config.SelectboxColumn(options=["Measured","Verified","Derived","Assumed","To review"])},height=500)
        if st.button("Save building masses",type="primary"):
            lexecute("DELETE FROM model_masses WHERE workspace_id=?",(workspace["id"],))
            for row in edited.to_dict("records"):
                if not str(row.get("label") or "").strip(): continue
                lexecute("""INSERT INTO model_masses(workspace_id,label,level_name,x,y,z,width,depth,height,finish,source_reference,confidence,notes,created_at) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",(workspace["id"],row.get("label",""),row.get("level_name","Ground"),row.get("x",0),row.get("y",0),row.get("z",0),row.get("width",1),row.get("depth",1),row.get("height",2.7),row.get("finish",""),row.get("source_reference",""),row.get("confidence","To review"),row.get("notes",""),now_stamp()))
            st.success("Building masses saved.")
            st.rerun()
    with tabs[2]:
        masses=lquery("SELECT id,label FROM model_masses WHERE workspace_id=? ORDER BY id",(workspace["id"],))
        mass_options={f"#{m['id']} · {m['label']}":m['id'] for m in masses}
        openings=ldf("SELECT id,mass_id,label,opening_type,face,offset_x,offset_z,width,height,count,notes,source_reference FROM model_openings WHERE workspace_id=? ORDER BY id",(workspace["id"],))
        st.dataframe(openings,use_container_width=True,hide_index=True)
        if masses:
            with st.form("opening_form"):
                mass_label=st.selectbox("Building mass",list(mass_options.keys()),help="The mass the opening is fixed to. Position is measured from this mass's front-left corner.")
                c1,c2=st.columns(2)
                label=c1.text_input("Opening label",value="Door D01")
                opening_type=c2.selectbox("Type",["Door","Window","Glazed opening","Roller door","Louvre","Other"],help="Door and window count toward the take-off schedule when linked to a location.")
                face=c1.selectbox("Face",["Front","Back","Left","Right","North","South","East","West"],help="Which side of the mass the opening sits on.")
                offset_x=c2.number_input("Offset along face (m)",value=0.0,step=.1,help="Horizontal distance from the face's left edge to the opening's left edge.")
                offset_z=c1.number_input("Sill / base height (m)",value=0.0,step=.1,help="Height of the opening's bottom edge above ground / floor level.")
                width=c2.number_input("Width (m)",value=.9,min_value=.05,step=.05,help="Width of the opening.")
                height=c1.number_input("Height (m)",value=2.1,min_value=.05,step=.05,help="Height of the opening.")
                count=c2.number_input("Count",value=1,min_value=1,step=1,help="Repeats (e.g. a bank of identical windows).")
                source=c1.text_input("Source reference")
                notes=st.text_area("Notes")
                save=st.form_submit_button("Add opening")
            if save:
                lexecute("""INSERT INTO model_openings(workspace_id,mass_id,label,opening_type,face,offset_x,offset_z,width,height,count,notes,source_reference,created_at) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?)""",(workspace["id"],mass_options[mass_label],label,opening_type,face,offset_x,offset_z,width,height,count,notes,source,now_stamp()))
                st.rerun()
        else:
            st.info("Add at least one building mass first.")
    with tabs[3]:
        render_pages=ldf("SELECT id,page_label,image_path FROM pages WHERE workspace_id=? AND page_type=? ORDER BY id",(workspace["id"],"Render / artist's impression"))
        if render_pages.empty:
            st.info("No render or artist's impression pages found. Add one in the document processor (page type \u201cRender / artist's impression\u201d) first.")
        else:
            options={f"#{int(r.id)} · {r.page_label}":int(r.id) for r in render_pages.itertuples()}
            selected=st.multiselect("Render pages to read",list(options.keys()),default=list(options.keys()),help="Artist's impressions and renders supply building form, storeys, roof and facade material/colour. Dimensions from a render are always treated as assumed and never override measured plan data.")
            model=st.text_input("Model",value=default_ai_model(ai_provider),help="Vision-capable model used to interpret the render images.")
            if st.button("Read render into 3D model",type="primary",disabled=not bool(resolve_ai_key(ai_provider,session_api_key)) or not selected):
                try:
                    data=_run_with_progress(
                        lambda: run_ai_render_read(int(workspace["id"]),[options[x] for x in selected],resolve_ai_key(ai_provider,session_api_key),model.strip() or default_ai_model(ai_provider),ai_provider),
                        f"Interpreting {len(selected)} render image(s) into the 3D model",
                    )
                    st.session_state["latest_render_result"]=data
                    st.success("Render read completed. Review the preview, then apply it.")
                except Exception as exc:
                    lexecute("INSERT INTO ai_runs(workspace_id,run_type,model,source_pages,status,response_json,error_message,created_at) VALUES(?,?,?,?,?,?,?,?)",(workspace["id"],"Render / artist's impression",model,json.dumps([options[x] for x in selected]),"Failed","",str(exc),now_stamp()))
                    st.error(_ai_error_hint(exc))
            data=st.session_state.get("latest_render_result")
            if data:
                st.json(data,expanded=False)
                mode=st.radio("Apply mode",["Merge with existing model","Replace assumed masses"],index=0,help="Merge keeps existing masses and adds/updates from the render. Replace deletes non-measured masses first, keeping measured/verified masses untouched.")
                if st.button("Apply render to 3D model",type="primary"):
                    counts=apply_render_to_model(int(workspace["id"]),data,mode="replace" if "Replace" in mode else "merge")
                    st.success(f"Applied render: {counts['masses']} masses, {counts['openings']} openings.")
                    st.session_state.pop("latest_render_result",None)
                    st.rerun()
            if not resolve_ai_key(ai_provider,session_api_key):
                st.warning("Configure the AI API key (OPENAI_API_KEY or GEMINI_API_KEY) or enter a session key in the sidebar to read renders with AI.")
    with tabs[4]:
        obj=generate_obj(int(workspace["id"]))
        geometry=json.dumps({"masses":lquery("SELECT * FROM model_masses WHERE workspace_id=?",(workspace["id"],)),"openings":lquery("SELECT * FROM model_openings WHERE workspace_id=?",(workspace["id"],))},indent=2,default=str)
        interactive_html=build_3d_figure(int(workspace["id"])).to_html(full_html=True,include_plotlyjs=True)
        st.download_button("Download interactive 3D render (HTML)",interactive_html,file_name=f"{safe_name(workspace.get('job_no'))}_interactive_3d.html",mime="text/html",use_container_width=True)
        st.download_button("Download OBJ model",obj,file_name=f"{safe_name(workspace.get('job_no'))}_building.obj",mime="text/plain",use_container_width=True)
        st.download_button("Download geometry JSON",geometry,file_name=f"{safe_name(workspace.get('job_no'))}_geometry.json",mime="application/json",use_container_width=True)
        st.caption("The HTML file is the shareable interactive 3D render. OBJ contains the rectangular masses. Openings are included in JSON and shown in the interactive model, but are not boolean-cut from the OBJ geometry.")


def quantity_schedule_page(workspace:Dict[str,Any]) -> None:
    hero(workspace)
    takeoff=dataframe_for_takeoff(int(workspace["id"]))
    if takeoff.empty:
        st.info("Build the take-off schedule first.")
        return
    scale_issues=scale_gate_issues(int(workspace["id"]))
    if scale_issues:
        st.warning("**Scale gate:** these page(s) have take-off rows/zones but no calibrated scale: " + ", ".join(f"`{i['page_label']}`" for i in scale_issues[:8]) + ". Calibrate in **Plan Mapper → Scale** before relying on the totals.")
    c1,c2,c3,c4,c5=st.columns(5)
    work=takeoff_work_rows(takeoff)
    c1.metric("Total m²",f"{work.loc[work['unit'].eq('m²'),'quantity'].sum():,.1f}")
    c2.metric("Lineal metres",f"{work.loc[work['unit'].eq('lm'),'quantity'].sum():,.1f}")
    c3.metric("Paint",f"{work['paint_litres'].sum():,.1f} L")
    c4.metric("Labour",f"{work['labour_hours'].sum():,.1f} hrs")
    c5.metric("Value ex GST",f"${work['value_ex_gst'].sum():,.0f}")
    tabs=st.tabs(["Schedule","Per-level pricing","Reconciliation","Quote export"])
    with tabs[0]:
        st.dataframe(takeoff[["id","section","element","location","substrate","finish_system","quantity","unit","quantity_status","coats","paint_litres","labour_hours","rate_per_unit","value_ex_gst","confidence","source_reference","notes"]],use_container_width=True,hide_index=True,height=560)
        section=work.groupby("section",dropna=False).agg(rows=("id","count"),m2=("quantity",lambda s:float(s[work.loc[s.index,"unit"].eq("m²")].sum())),paint_litres=("paint_litres","sum"),labour_hours=("labour_hours","sum"),value_ex_gst=("value_ex_gst","sum")).reset_index()
        st.subheader("Section summary")
        st.dataframe(section,use_container_width=True,hide_index=True)
    with tabs[1]:
        st.subheader("Priced per-level summary")
        basis=workspace_setting(int(workspace["id"]),"internal_pricing_basis","wall_m2")
        st.caption(f"Levels are read from the row locations (Ground, Level 1 … Roof). Internal wall pricing basis: **{pricing_basis_label(basis)}**. Pricing settings (mark-up, GST) come from Settings. Floor-area rows are not priced.")
        levels=quote_summary_frame(int(workspace["id"]))
        st.dataframe(levels,use_container_width=True,hide_index=True)
        if not levels.empty:
            g1,g2,g3,g4=st.columns(4)
            g1.metric("Value ex GST",f"${levels['value_ex_gst'].sum():,.2f}")
            g2.metric("Mark-up ex GST",f"${(levels['markup_ex_gst']-levels['value_ex_gst']).sum():,.2f}")
            g3.metric("GST",f"${levels['gst'].sum():,.2f}")
            g4.metric("Total inc GST",f"${levels['total_inc_gst'].sum():,.2f}")
            st.download_button("Download per-level summary CSV",per_level_summary_csv(int(workspace["id"])).encode("utf-8"),file_name=f"{safe_name(workspace.get('job_no'))}_per_level_summary.csv",mime="text/csv")
    with tabs[2]:
        st.subheader("AI vs drawn reconciliation")
        reco=reconcile_ai_vs_drawn(int(workspace["id"]))
        if reco.empty:
            st.caption("No rows to reconcile yet.")
        else:
            st.dataframe(reco,use_container_width=True,hide_index=True,height=420)
            diffs=reco[reco["status"].isin(["Difference","AI only (not drawn)","Drawn only (no AI draft)"])]
            if not diffs.empty:
                st.warning(f"**{len(diffs)} line item(s) differ between AI and drawn quantities.**")
    with tabs[3]:
        st.subheader("Quotation export")
        st.markdown("<div class='pb-note'>Priced per-level quotation with the mark-up and GST configured in Settings. PDF uses the same figures.</div>",unsafe_allow_html=True)
        try:
            quote_xlsx=quote_workbook_bytes(int(workspace["id"]))
            quote_pdf=quote_pdf_bytes(int(workspace["id"]))
        except Exception as exc:
            st.exception(exc)
            quote_xlsx, quote_pdf = None, None
        if quote_xlsx:
            st.download_button("Download priced quotation (Excel)",quote_xlsx,file_name=f"{safe_name(workspace.get('job_no'))}_quotation.xlsx",mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",use_container_width=True)
        if quote_pdf:
            st.download_button("Download priced quotation (PDF)",quote_pdf,file_name=f"{safe_name(workspace.get('job_no'))}_quotation.pdf",mime="application/pdf",use_container_width=True)


def export_page(workspace:Dict[str,Any],bridge:Optional[JobHubBridge],user:Dict[str,Any]) -> None:
    hero(workspace)
    st.markdown("<div class='pb-card'>",unsafe_allow_html=True)
    st.subheader("Complete take-off pack")
    excel=excel_export_bytes(int(workspace["id"]))
    package=zip_export_bytes(int(workspace["id"]))
    st.download_button("Download subscription-style Excel take-off pack",excel,file_name=f"{safe_name(workspace.get('job_no'))}_paint_takeoff.xlsx",mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",use_container_width=True)
    st.download_button("Download complete plans + take-off + 3D package",package,file_name=f"{safe_name(workspace.get('job_no'))}_planreader_3d_package.zip",mime="application/zip",use_container_width=True)
    st.markdown("</div>",unsafe_allow_html=True)
    st.markdown("<div class='pb-card'>",unsafe_allow_html=True)
    st.subheader("Send reviewed draft to JobHub")
    if not bridge or not workspace.get("jobhub_job_id"):
        st.info("Link this workspace to a JobHub job to send an approved draft back.")
    else:
        st.markdown("<div class='pb-warning'>This creates a new draft take-off package in JobHub. It does not overwrite previous packages. Review quantities and drawing revision first.</div>",unsafe_allow_html=True)
        confirmed=st.checkbox("I have reviewed the take-off and source references")
        if st.button("Create draft take-off package in JobHub",type="primary",disabled=not confirmed):
            try:
                package_id,line_count=push_takeoff_to_jobhub(int(workspace["id"]),bridge,str(user.get("username") or "PlanReader"))
                st.success(f"Created JobHub take-off package #{package_id} with {line_count} lines.")
            except Exception as exc:
                st.exception(exc)
    st.markdown("</div>",unsafe_allow_html=True)
    st.markdown("<div class='pb-card'>",unsafe_allow_html=True)
    st.subheader("One-file progress marker → JobHub")
    if not bridge or not workspace.get("jobhub_job_id"):
        st.info("Link this workspace to a JobHub job to send the progress marker (3D render + take-off + job documents in one file).")
    else:
        st.markdown("<div class='pb-note'>Builds a single ZIP: the interactive 3D render, OBJ + geometry, the Excel take-off pack, take-off CSVs (including JobHub's shared format), scope registers, source plans and rendered page previews. The ZIP is stored on the job in JobHub as a <b>Progress Marker</b>, take-off rows are synced live into <code>job_takeoff_rows</code>, and a draft take-off package is created alongside.</div>",unsafe_allow_html=True)
        confirmed=st.checkbox("I have reviewed the take-off, drawing revision and 3D geometry")
        if st.button("Build & send progress marker to JobHub",type="primary",disabled=not confirmed):
            try:
                result=push_progress_marker_to_jobhub(int(workspace["id"]),bridge,str(user.get("username") or "PlanReader"))
                st.success(f"Sent {result['file_name']} ({result['size_bytes']/1024:.0f} KB) to JobHub job #{result['job_id']} as a progress marker; {result['takeoff_rows_synced']} take-off rows synced live and draft package #{result['package_id']} created ({result['package_lines']} lines).")
            except Exception as exc:
                st.exception(exc)
        st.markdown("#### Progress markers already on this job")
        try:
            markers=list_jobhub_progress_markers(bridge,int(workspace.get("jobhub_job_id")))
        except Exception:
            markers=[]
        if markers:
            st.dataframe(pd.DataFrame(markers),use_container_width=True,hide_index=True)
        else:
            st.caption("No progress markers sent yet for this job.")
    st.markdown("</div>",unsafe_allow_html=True)
    st.markdown("<div class='pb-card'>",unsafe_allow_html=True)
    st.subheader("Import take-off from JobHub")
    if not bridge or not workspace.get("jobhub_job_id"):
        st.info("Link this workspace to a JobHub job to pull take-off rows already stored in JobHub.")
    else:
        st.markdown("<div class='pb-note'>Pulls the job's take-off rows from the shared JobHub database. Imported rows are tagged and replaced on re-import; existing rows created here are kept.</div>",unsafe_allow_html=True)
        if st.button("Pull take-off rows from JobHub",type="primary"):
            try:
                n=pull_takeoff_from_jobhub(int(workspace["id"]),bridge)
                st.success(f"Imported {n} take-off row(s) from JobHub." if n else "JobHub has no take-off rows for this job.")
                st.rerun()
            except Exception as exc:
                st.exception(exc)
    st.markdown("</div>",unsafe_allow_html=True)
    st.markdown("<div class='pb-card'>",unsafe_allow_html=True)
    st.subheader("Publish final take-off to JobHub")
    if not bridge or not workspace.get("jobhub_job_id"):
        st.info("Link this workspace to a JobHub job to publish the final take-off and quotation.")
    else:
        st.markdown("<div class='pb-warning'>Publishes the final take-off package (status <b>Published</b>), the priced Excel quotation and the progress package, syncs take-off rows live, and marks the shared JobHub job as <b>Published</b>. Only do this when quantities, drawing revision and pricing are final.</div>",unsafe_allow_html=True)
        scale_issues=scale_gate_issues(int(workspace["id"]))
        if scale_issues:
            st.warning("**Scale gate:** " + ", ".join(f"`{i['page_label']}`" for i in scale_issues[:8]) + " has no calibrated scale — calibrate before publishing.")
        publish_confirm=st.checkbox("I confirm this is the final, reviewed and priced take-off for the current drawing issue")
        if st.button("Publish to JobHub",type="primary",disabled=not publish_confirm or bool(scale_issues)):
            try:
                result=publish_job_to_jobhub(int(workspace["id"]),bridge,str(user.get("username") or "PlanReader"))
                st.success(f"Published job #{result['job_id']}: final package #{result['package_id']} ({result['package_lines']} lines), quotation '{result['quotation']}' and progress marker '{result['progress_marker']}' stored; {result['takeoff_rows_synced']} take-off rows synced. JobHub status: {result['job_status']}.")
            except Exception as exc:
                st.exception(exc)
    st.markdown("</div>",unsafe_allow_html=True)


# ---------------------------------------------------------------------------
# Offline Plan Reader (no AI)
# ---------------------------------------------------------------------------

def offline_plan_reader_page(workspace: Dict[str, Any]) -> None:
    """Offline plan reading page - no AI required. Uses OCR, vector extraction, and pattern matching."""
    st.markdown('<div class="pb-header"><h2>Offline Plan Reader</h2><p>No AI required - uses OCR, vector extraction, and pattern matching</p></div>', unsafe_allow_html=True)
    
    if not OFFLINE_READER_AVAILABLE:
        st.error("Offline reader module not available. Please ensure pb_planreader_offline.py is installed.")
        st.info("The offline reader uses PyMuPDF4LLM for OCR, PyMuPDF for vector extraction, and OpenCV for image processing.")
        return
    
    # Get documents in workspace
    docs = lquery("SELECT * FROM documents WHERE workspace_id=? ORDER BY uploaded_at DESC", (workspace["id"],))
    
    if not docs:
        st.warning("No documents uploaded yet. Go to **Job & Documents** to upload a plan.")
        return
    
    # Document selector
    doc_options = {f"{d['file_name']} (p{d.get('page_count', '?')})": d["id"] for d in docs}
    selected_label = st.selectbox("Select document", list(doc_options.keys()))
    doc_id = doc_options[selected_label]
    doc = [d for d in docs if d["id"] == doc_id][0]
    doc_path = Path(str(doc["path"]))
    
    if not doc_path.exists():
        st.error("Document file not found on disk.")
        return
    
    # Page selector
    page_count = doc.get("page_count", 1) or 1
    page_range = st.multiselect(
        "Select pages to analyze",
        list(range(1, page_count + 1)),
        default=[1] if page_count >= 1 else [],
        help="Select which pages to analyze",
    )
    
    # OCR toggle
    use_ocr = st.checkbox("Enable OCR (for scanned plans)", value=True, help="Uses PyMuPDF4LLM to extract text from scanned/image-based PDFs")
    
    # Analysis options
    st.subheader("Analysis Options")
    col1, col2, col3 = st.columns(3)
    with col1:
        analyze_dims = st.checkbox("Detect dimensions", value=True)
    with col2:
        analyze_walls = st.checkbox("Detect walls", value=True)
    with col3:
        analyze_rooms = st.checkbox("Detect rooms/areas", value=True)
    
    col4, col5, col6 = st.columns(3)
    with col4:
        analyze_materials = st.checkbox("Detect materials", value=True)
    with col5:
        analyze_colours = st.checkbox("Detect colours", value=True)
    with col6:
        analyze_grid = st.checkbox("Detect grid lines", value=True)
    
    # Run analysis button
    if st.button("Analyze Pages", type="primary", use_container_width=True):
        if not page_range:
            st.error("Please select at least one page.")
            return
        
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        results = []
        takeoff_rows = []
        
        for i, page_no in enumerate(page_range):
            status_text.text(f"Analyzing page {page_no}...")
            progress_bar.progress((i + 1) / len(page_range))
            
            try:
                # Run offline analysis
                analysis = analyze_page_offline(
                    str(doc_path), page_no, use_ocr=use_ocr
                )
                results.append(analysis)
                
                # Generate takeoff rows based on analysis
                page_type = analysis.get("page_type", "Other")
                label = analysis.get("label", f"Page {page_no}")
                scale = analysis.get("scale")
                scale_text = scale.get("text", "") if scale else ""
                
                # Floor Plan takeoff
                if "Floor Plan" in page_type:
                    walls = analysis.get("walls", [])
                    if walls:
                        # Convert wall lengths: PDF pts → page mm → real metres
                        real_m = wall_length_real_m(
                            sum(w["length"] for w in walls), scale
                        )
                        if real_m is not None:
                            total_length = round(real_m, 3)
                            scale_note = f"Scale {scale_text}"
                            status = "Auto-detected"
                        else:
                            # No scale: do NOT produce a real-world lm
                            # quantity that could be summed as metres.
                            total_length = None
                            scale_note = "NO SCALE — real-world quantity unavailable"
                            status = "Uncalibrated"

                        takeoff_rows.append({
                            "page_no": page_no,
                            "drawing": label,
                            "section": "Internal",
                            "element": "Walls",
                            "unit": "lm" if total_length is not None else "",
                            "quantity": total_length,
                            "scale": scale_text,
                            "notes": f"{len(walls)} wall segments detected. {scale_note}",
                            "status": status,
                        })
                    
                    # Rooms
                    for room in analysis.get("rooms", []):
                        if room.get("type") == "text_label":
                            takeoff_rows.append({
                                "page_no": page_no,
                                "drawing": label,
                                "section": "Internal",
                                "element": room["label"],
                                "unit": "m2",
                                "quantity": 0,
                                "scale": scale_text,
                                "notes": "Room label detected - area to measure",
                                "status": "To measure",
                            })
                
                # Elevation takeoff
                elif "Elevation" in page_type:
                    walls = analysis.get("walls", [])
                    if walls:
                        takeoff_rows.append({
                            "page_no": page_no,
                            "drawing": label,
                            "section": "External",
                            "element": "Walls",
                            "unit": "m2",
                            "quantity": 0,
                            "scale": scale_text,
                            "notes": f"{len(walls)} wall lines detected",
                            "status": "To measure",
                        })
                    
                    # Materials
                    for mat in analysis.get("materials", []):
                        takeoff_rows.append({
                            "page_no": page_no,
                            "drawing": label,
                            "section": "External",
                            "element": mat["keyword"].title(),
                            "unit": "m2",
                            "quantity": 0,
                            "scale": scale_text,
                            "notes": f"Material: {mat['text']}",
                            "status": "To measure",
                        })
                
                # Section takeoff
                elif "Section" in page_type:
                    for dim in analysis.get("dimensions", []):
                        if dim["unit"] == "mm" and 100 <= dim["value"] <= 10000:
                            takeoff_rows.append({
                                "page_no": page_no,
                                "drawing": label,
                                "section": "Section",
                                "element": "Dimension",
                                "unit": "mm",
                                "quantity": dim["value"],
                                "scale": scale_text,
                                "notes": f"Auto: {dim['text']}",
                                "status": "Auto-detected",
                            })
                
                # Schedule/Specification takeoff
                elif "Schedule" in page_type or "Specification" in page_type:
                    for word in analysis.get("words", []):
                        text_lower = word.get("text", "").strip().lower()
                        if any(kw in text_lower for kw in ["paint", "coat", "primer", "sealer", "finish"]):
                            takeoff_rows.append({
                                "page_no": page_no,
                                "drawing": label,
                                "section": "Schedule",
                                "element": word["text"].strip(),
                                "unit": "item",
                                "quantity": 1,
                                "scale": scale_text,
                                "notes": "Schedule item",
                                "status": "To verify",
                            })
            
            except Exception as e:
                st.error(f"Error analyzing page {page_no}: {e}")
        
        progress_bar.empty()
        status_text.empty()
        
        # Store results in session state
        st.session_state["offline_results"] = results
        st.session_state["offline_takeoff"] = takeoff_rows
        st.session_state["offline_doc_name"] = doc["file_name"]
        
        st.success(f"Analysis complete! Analyzed {len(results)} pages, found {len(takeoff_rows)} takeoff items.")
    
    # Display results if available
    if "offline_results" in st.session_state and st.session_state.get("offline_doc_name") == doc["file_name"]:
        results = st.session_state["offline_results"]
        takeoff_rows = st.session_state["offline_takeoff"]
        
        # Tabs for different views
        tab1, tab2, tab3, tab4 = st.tabs(["Takeoff", "Page Analysis", "Raw Data", "Export"])
        
        with tab1:
            st.subheader("Auto-Generated Takeoff")
            
            if takeoff_rows:
                takeoff_df = pd.DataFrame(takeoff_rows)
                
                # Summary stats
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    st.metric("Total Items", len(takeoff_rows))
                with col2:
                    st.metric("Auto-detected", len([r for r in takeoff_rows if r["status"] == "Auto-detected"]))
                with col3:
                    st.metric("To measure", len([r for r in takeoff_rows if r["status"] == "To measure"]))
                with col4:
                    st.metric("To verify", len([r for r in takeoff_rows if r["status"] == "To verify"]))
                
                st.dataframe(takeoff_df, use_container_width=True)
                
                # Export button
                csv = takeoff_df.to_csv(index=False)
                st.download_button(
                    "Download Takeoff CSV",
                    csv,
                    f"{doc['file_name']}_offline_takeoff.csv",
                    "text/csv",
                    use_container_width=True,
                )
            else:
                st.info("No takeoff items generated. Try selecting different pages or adjusting analysis options.")
        
        with tab2:
            st.subheader("Page-by-Page Analysis")
            
            for analysis in results:
                with st.expander(f"Page {analysis['page_no']} - {analysis['page_type']} ({analysis['label']})", expanded=False):
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        st.markdown("**Classification:**")
                        st.write(f"- Type: {analysis['page_type']}")
                        st.write(f"- Label: {analysis['label']}")
                        
                        if analysis.get("scale"):
                            st.write(f"- Scale: {analysis['scale']['text']}")
                        
                        st.markdown("**Vector Graphics:**")
                        st.write(f"- Lines: {analysis['line_count']}")
                        st.write(f"- Rectangles: {analysis['rect_count']}")
                        st.write(f"- Walls detected: {analysis['wall_count']}")
                        st.write(f"- Grid detected: {'Yes' if analysis['grid']['grid_detected'] else 'No'}")
                    
                    with col2:
                        st.markdown("**Extracted Data:**")
                        st.write(f"- Words: {analysis['word_count']}")
                        st.write(f"- Dimensions: {len(analysis['dimensions'])}")
                        st.write(f"- Rooms: {len(analysis['rooms'])}")
                        st.write(f"- Materials: {len(analysis['materials'])}")
                        st.write(f"- Colours: {len(analysis['colours'])}")
                    
                    if analysis["dimensions"]:
                        st.markdown("**Dimensions Found:**")
                        dims = [{"Text": d["text"], "Value": f"{d['value']:.0f} {d['unit']}", "Confidence": "High" if d.get("confidence") != "low" else "Low"} for d in analysis["dimensions"][:10]]
                        st.dataframe(pd.DataFrame(dims), use_container_width=True)
                    
                    if analysis["rooms"]:
                        st.markdown("**Rooms/Areas:**")
                        rooms = [{"Label": r["label"], "Type": r["type"]} for r in analysis["rooms"]]
                        st.dataframe(pd.DataFrame(rooms), use_container_width=True)
                    
                    if analysis["materials"]:
                        st.markdown("**Materials:**")
                        mats = [{"Keyword": m["keyword"], "Text": m["text"]} for m in analysis["materials"]]
                        st.dataframe(pd.DataFrame(mats), use_container_width=True)
        
        with tab3:
            st.subheader("Raw Extracted Data")
            
            selected_page = st.selectbox("Select page for raw data", [r["page_no"] for r in results])
            analysis = [r for r in results if r["page_no"] == selected_page][0]
            
            # Text preview
            st.markdown("**Extracted Text (first 2000 chars):**")
            st.text_area("Text", analysis["text"][:2000], height=200, disabled=True)
            
            # Walls
            if analysis["walls"]:
                st.markdown("**Detected Walls:**")
                walls_df = pd.DataFrame([
                    {
                        "Start": f"({w['start'][0]:.0f}, {w['start'][1]:.0f})",
                        "End": f"({w['end'][0]:.0f}, {w['end'][1]:.0f})",
                        "Length": f"{w['length']:.0f} pts",
                        "Width": f"{w['width']:.2f}",
                    }
                    for w in analysis["walls"][:20]
                ])
                st.dataframe(walls_df, use_container_width=True)
            
            # Dimensions
            if analysis["dimensions"]:
                st.markdown("**Detected Dimensions:**")
                dims_df = pd.DataFrame([
                    {
                        "Text": d["text"],
                        "Value": f"{d['value']:.0f}",
                        "Unit": d["unit"],
                        "Position": f"({d['position']['x0']:.0f}, {d['position']['y0']:.0f})",
                    }
                    for d in analysis["dimensions"]
                ])
                st.dataframe(dims_df, use_container_width=True)
        
        with tab4:
            st.subheader("Export Options")
            
            if results:
                # Full report
                full_report = ""
                for analysis in results:
                    full_report += generate_report(analysis)
                    full_report += "\n\n" + "=" * 80 + "\n\n"
                
                st.download_button(
                    "Download Full Analysis Report",
                    full_report,
                    f"{doc['file_name']}_offline_report.txt",
                    "text/plain",
                    use_container_width=True,
                )
            
            if takeoff_rows:
                takeoff_df = pd.DataFrame(takeoff_rows)
                
                # Excel export
                excel_buffer = io.BytesIO()
                with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
                    takeoff_df.to_excel(writer, index=False, sheet_name='Takeoff')
                
                st.download_button(
                    "Download Takeoff Excel",
                    excel_buffer.getvalue(),
                    f"{doc['file_name']}_offline_takeoff.xlsx",
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True,
                )
            
            st.info("Export includes all auto-detected items. Review 'To measure' items manually.")
    
    # Instructions
    with st.expander("How the Offline Reader Works", expanded=False):
        st.markdown("""
        ### Offline Plan Reader Features
        
        **No AI Required** - Everything runs locally on the server.
        
        #### Text Extraction
        - Uses PyMuPDF4LLM for OCR on scanned plans
        - Falls back to PyMuPDF for text-based PDFs
        - Extracts text with position coordinates
        
        #### Vector Graphics Analysis
        - Detects wall lines (dark, medium-width lines)
        - Finds grid lines (dashed, light lines)
        - Extracts all line segments and rectangles
        
        #### Dimension Detection
        - Recognizes: 1200, 3.5m, 10'-6", 1:100, etc.
        - Filters reasonable ranges (10mm to 100m)
        - Associates dimensions with positions
        
        #### Room/Area Detection
        - Identifies room labels (bedroom, kitchen, etc.)
        - Detects enclosed rectangles as potential rooms
        
        #### Material/Colour Detection
        - Recognizes: render, brick, timber, concrete, etc.
        - Detects: white, charcoal, cream, etc.
        
        #### Page Classification
        - Floor Plans, Elevations, Sections, Schedules
        - Uses text content and visual features
        - Detects scale indicators
        
        #### Auto Takeoff Generation
        - Creates takeoff rows from detected elements
        - Marks items as auto-detected or to-measure
        - Exports to CSV/Excel for further use
        """)
    
    with st.expander("Tips for Best Results", expanded=False):
        st.markdown("""
        #### For Scanned Plans
        - Enable OCR (checked by default)
        - Ensure scans are clear and straight
        - Minimum 300 DPI recommended
        
        #### For Vector PDFs
        - Most architectural PDFs have vector text
        - Wall detection works best with standard line weights
        - Grid lines are detected automatically
        
        #### Dimension Detection
        - Looks for numbers with units (mm, m, etc.)
        - Also finds bare numbers in reasonable ranges
        - Scale ratios (1:100) are detected automatically
        
        #### After Analysis
        - Review "To measure" items manually
        - Verify "Auto-detected" quantities
        - Export to Excel for editing
        """)


def settings_page(workspace:Dict[str,Any],bridge:Optional[JobHubBridge],session_api_key:str,ai_provider:str="OpenAI") -> None:
    hero(workspace)
    st.write(f"App version: `{APP_VERSION}`")
    st.write(f"PlanReader data folder: `{DATA_DIR}`")
    st.write(f"Workspace folder: `{workspace_path(int(workspace['id']))}`")
    st.write(f"Drawing canvas: `{'available' if CANVAS_AVAILABLE else 'manual coordinate fallback'}`")
    st.write(f"AI provider: `{ai_provider}`")
    st.write(f"OpenAI key: `{'configured' if resolve_openai_key(session_api_key) else 'not configured'}`")
    st.write(f"Gemini key: `{'configured' if resolve_gemini_key(session_api_key) else 'not configured'}`")
    st.write(f"JobHub bridge: `{'connected' if bridge else 'not configured'}`")
    if bridge:
        try:
            st.write(f"JobHub tables detected: `{', '.join(sorted(bridge.table_names())[:30])}`")
        except Exception as exc:
            st.error(f"JobHub bridge error: {exc}")
    st.markdown("<div class='pb-card'>",unsafe_allow_html=True)
    st.subheader("Project-level settings")
    st.caption("Defaults used by the take-off and quotation for this job only. Rows keep their own values where already set; these are the fallbacks and quote pricing inputs.")
    wid=int(workspace["id"])
    with st.form("workspace_settings_form"):
        s1,s2,s3=st.columns(3)
        coverage=s1.number_input("Default coverage (m²/L)",min_value=1.0,value=float(workspace_setting(wid,"default_coverage_m2_per_litre",12.0)),step=0.5)
        productivity=s2.number_input("Default productivity (m²/hr)",min_value=0.5,value=float(workspace_setting(wid,"default_productivity_m2_per_hour",8.0)),step=0.5)
        wall_height=s3.number_input("Default wall height (m)",min_value=0.5,value=float(workspace_setting(wid,"default_wall_height_m",2.7)),step=0.1)
        p1,p2=st.columns(2)
        margin=p1.number_input("Pricing mark-up (%)",min_value=0.0,value=float(workspace_setting(wid,"pricing_margin_pct",0.0)),step=0.5)
        gst=p2.number_input("GST rate (%)",min_value=0.0,value=float(workspace_setting(wid,"gst_rate_pct",10.0)),step=0.5)
        basis=st.radio("Internal wall pricing basis",options=["wall_m2","floor_m2"],format_func=lambda o:"Wall m² (rate × measured wall area)" if o=="wall_m2" else "Floor m² (rate × each level's floor area)",index=0 if workspace_setting(wid,"internal_pricing_basis","wall_m2")!="floor_m2" else 1,help="Wall m² prices internal wall rows from their measured wall area. Floor m² prices internal wall rows from each level's internal floor area (rate × floor m²), with measured wall m² still shown for reference in the schedule and quotation. Floor-area rows are never priced themselves.")
        header=st.text_input("Quote header / company name",value=str(workspace_setting(wid,"quote_header","Premier Brushworks Pty Ltd")))
        footer=st.text_area("Quote footer",value=str(workspace_setting(wid,"quote_footer","Valid for 30 days. All quantities require estimator review against the current issued drawings and specification before pricing or construction use.")),height=90)
        if st.form_submit_button("Save project settings",type="primary"):
            for key,val in [("default_coverage_m2_per_litre",coverage),("default_productivity_m2_per_hour",productivity),("default_wall_height_m",wall_height),("pricing_margin_pct",margin),("gst_rate_pct",gst),("internal_pricing_basis",basis),("quote_header",header),("quote_footer",footer)]:
                set_workspace_setting(wid,key,val)
            st.success("Project settings saved.")
            st.rerun()
    st.markdown("</div>",unsafe_allow_html=True)
    st.markdown("<div class='pb-note'>For permanent Render storage, attach a persistent disk and set PLANREADER_DATA_DIR to its mount path. Separate Render services cannot read each other's local disk paths; use shared PostgreSQL for metadata and object storage or PlanReader uploads for file bytes.</div>",unsafe_allow_html=True)
    if st.checkbox("Show destructive controls"):
        confirm=st.text_input("Type DELETE to remove this PlanReader workspace")
        if st.button("Delete workspace",type="secondary",disabled=confirm!="DELETE"):
            shutil.rmtree(workspace_path(int(workspace["id"])),ignore_errors=True)
            lexecute("DELETE FROM workspaces WHERE id=?",(workspace["id"],))
            st.session_state.pop("workspace_id",None)
            st.rerun()


def main() -> None:
    st.set_page_config(page_title=APP_NAME,page_icon="🏗️",layout="wide")
    if "reset" in st.query_params:
        st.query_params.clear()
        user = st.session_state.get("planreader_user") or {"username": "Estimator", "role": "Estimator"}
        st.session_state.clear()
        st.session_state["planreader_user"] = user
        st.session_state["workspace_id"] = 1
    app_css()
    init_local_db()
    bridge=get_jobhub_bridge()
    user=st.session_state.get("planreader_user")
    if not user:
        login_screen(bridge)
    workspace_id=sidebar_workspace_selector(bridge)
    ai_provider=st.sidebar.selectbox("AI provider",AI_PROVIDERS,index=AI_PROVIDERS.index(resolve_ai_provider()) if resolve_ai_provider() in AI_PROVIDERS else 0,help="OpenAI is pay-per-use. Google Gemini has a generous free tier and needs a GEMINI_API_KEY from AI Studio.")
    session_api_key=st.sidebar.text_input("AI API key (session only)",type="password",help="Leave blank when OPENAI_API_KEY or GEMINI_API_KEY is configured in Render or Windows.")
    if resolve_ai_key(ai_provider,session_api_key):
        st.sidebar.success("AI plan reading ready")
    else:
        st.sidebar.caption("Manual take-off and 3D modelling work without AI.")
    if not workspace_id:
        hero()
        st.info("Open a JobHub job or create a standalone workspace from the sidebar.")
        return
    workspace=current_workspace()
    if not workspace:
        st.session_state.pop("workspace_id",None)
        st.rerun()
    menu_options = [
        "Dashboard",
        "Job & Documents",
        "Drawing Register",
        "Review & QA",
        "Subscription Take-off",
        "Plan Mapper",
        "3D Building Model",
        "Quantity Schedule",
        "Offline Plan Reader",
        "Export / JobHub",
        "Settings",
    ]
    nav_target = st.session_state.get("_pb_nav_target")
    nav_payload = st.session_state.get("_pb_nav_payload")

    valid_nav_target = None
    active_ws_id = int(workspace["id"]) if isinstance(workspace, dict) and workspace.get("id") else None

    if nav_target and isinstance(nav_payload, dict):
        payload_ws_id = nav_payload.get("workspace_id")
        if payload_ws_id == active_ws_id:
            valid_nav_target = nav_target
            if nav_target == "takeoff":
                if nav_payload.get("takeoff_row_id"):
                    st.session_state["active_takeoff_row_id"] = nav_payload["takeoff_row_id"]
            elif nav_target in ("drawing", "page"):
                if nav_payload.get("page_id"):
                    st.session_state["active_page_id"] = nav_payload["page_id"]
            elif nav_target == "register":
                if nav_payload.get("register_item_id"):
                    st.session_state["active_register_item_id"] = nav_payload["register_item_id"]

    # Pop consumed/rejected navigation signals cleanly
    st.session_state.pop("_pb_nav_target", None)
    st.session_state.pop("_pb_nav_payload", None)

    default_index = 0
    if valid_nav_target == "takeoff":
        default_index = menu_options.index("Subscription Take-off")
    elif valid_nav_target in ("drawing", "page"):
        default_index = menu_options.index("Plan Mapper")
    elif valid_nav_target == "register":
        default_index = menu_options.index("Drawing Register")
    elif valid_nav_target == "model":
        default_index = menu_options.index("3D Building Model")
    elif valid_nav_target == "review":
        default_index = menu_options.index("Review & QA")
    else:
        saved_menu = st.session_state.get("_pb_current_menu")
        if saved_menu in menu_options:
            default_index = menu_options.index(saved_menu)

    menu = st.sidebar.radio("Menu", menu_options, index=default_index)
    st.session_state["_pb_current_menu"] = menu

    if menu == "Dashboard": dashboard_page(workspace)
    elif menu == "Job & Documents": project_documents_page(workspace, bridge, user)
    elif menu == "Drawing Register": drawing_register_page(workspace)
    elif menu == "Review & QA":
        hero(workspace)
        import sys
        from pb_commercial_review_v161 import render_commercial_review_workspace
        render_commercial_review_workspace(sys.modules[__name__], workspace)
    elif menu == "Subscription Take-off": subscription_takeoff_page(workspace, session_api_key, ai_provider)
    elif menu == "Plan Mapper": plan_mapper_page(workspace)
    elif menu == "3D Building Model": model_3d_page(workspace, session_api_key, ai_provider)
    elif menu == "Quantity Schedule": quantity_schedule_page(workspace)
    elif menu == "Offline Plan Reader": offline_plan_reader_page(workspace)
    elif menu == "Export / JobHub": export_page(workspace, bridge, user)
    else: settings_page(workspace, bridge, session_api_key, ai_provider)


if __name__ == "__main__":
    main()
